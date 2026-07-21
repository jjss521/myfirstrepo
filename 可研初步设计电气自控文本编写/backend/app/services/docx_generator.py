# -*- coding: utf-8 -*-
"""
Word 生成引擎
==============
根据《市政公用工程设计文件编制深度规定》（2025年版）四大工程类型的电气 / 自控
深度要求，结合 Excel 负荷计算数据，生成完整的市政工程设计文件电气自控说明与
设备材料表。

特性：
  * 四种 Word 模板：standard / compact / report / modern
  * 专业内容引擎：每个栏目生成真实工程语言，绝不输出占位符
  * preview() 返回结构化 blocks，供 GUI 实时预览
  * 设备材料表由负荷数据估算，数量合理
  * 已修复：文字重复 / 占位符当正文 / 编号混乱 / 补偿 Sjs 混用 / 页码域损坏
"""

import os
import re
import json
import math
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

try:
    from app.config import RULES_DIR, OUTPUT_DIR, STAGE_CODE
except ModuleNotFoundError:
    import os as _os
    import sys as _sys
    _base = _os.path.dirname(_os.path.dirname(_os.path.dirname(_os.path.abspath(__file__))))
    _sys.path.insert(0, _base)
    from app.config import RULES_DIR, OUTPUT_DIR, STAGE_CODE

# ----------------------------------------------------------------------------
# 四种模板配置
# ----------------------------------------------------------------------------
TEMPLATES = {
    'standard': {
        'label': '标准版',
        'desc': '含封面、黑体标题、表格带边框（常规出图风格）',
        'cover': True, 'header': False, 'sig_block': True, 'page_number': True,
        'title_color': RGBColor(0x1A, 0x1A, 0x1A),
        'accent': RGBColor(0x00, 0x00, 0x00),
        'h1': 16, 'h2': 13, 'h3': 11, 'body': 10.5,
        'line': 1.5, 'table_border': True, 'header_shade': None,
    },
    'compact': {
        'label': '紧凑版',
        'desc': '无封面、小字号、紧凑行距（内部评审 / 快速输出）',
        'cover': False, 'header': False, 'sig_block': False, 'page_number': True,
        'title_color': RGBColor(0x1A, 0x1A, 0x1A),
        'accent': RGBColor(0x00, 0x00, 0x00),
        'h1': 14, 'h2': 12, 'h3': 10.5, 'body': 9.5,
        'line': 1.2, 'table_border': True, 'header_shade': None,
    },
    'report': {
        'label': '报批版',
        'desc': '仿公文、含编制说明与会签栏（正式报审）',
        'cover': True, 'header': False, 'sig_block': True, 'page_number': True,
        'title_color': RGBColor(0x00, 0x00, 0x00),
        'accent': RGBColor(0x00, 0x00, 0x00),
        'h1': 15, 'h2': 12.5, 'h3': 11, 'body': 10.5,
        'line': 1.6, 'table_border': True, 'header_shade': RGBColor(0xDD, 0xDD, 0xDD),
    },
    'modern': {
        'label': '现代版',
        'desc': '蓝色彩色标题、页眉项目名、表头着色（现代美观）',
        'cover': True, 'header': True, 'sig_block': True, 'page_number': True,
        'title_color': RGBColor(0x1F, 0x4E, 0x79),
        'accent': RGBColor(0x2E, 0x5C, 0x8A),
        'h1': 16, 'h2': 13, 'h3': 11, 'body': 10.5,
        'line': 1.5, 'table_border': True, 'header_shade': RGBColor(0x2E, 0x5C, 0x8A),
    },
}

CN_NUM = ['零', '一', '二', '三', '四', '五', '六', '七', '八', '九', '十',
          '十一', '十二', '十三', '十四', '十五', '十六', '十七', '十八', '十九', '二十']


def cn_num(n):
    if 0 <= n <= len(CN_NUM) - 1:
        return CN_NUM[n]
    return str(n)


def _esc(s):
    """XML 转义，防止 & < > 破坏 docx。"""
    return (str(s).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;'))


# ----------------------------------------------------------------------------
# 生成引擎
# ----------------------------------------------------------------------------
class DocxGenerator:
    def __init__(self, rules_dir=RULES_DIR, output_dir=OUTPUT_DIR, template='standard'):
        self.rules_dir = rules_dir
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
        self.set_template(template)

    def set_template(self, template):
        if template not in TEMPLATES:
            template = 'standard'
        self.template = template
        self.cfg = TEMPLATES[template]

    # --- 加载规范规则 ---
    def _load_rules(self, project_type, design_stage):
        stage_en = STAGE_CODE.get(design_stage, 'preliminary')
        fn = os.path.join(self.rules_dir, f'{project_type}_{stage_en}.json')
        if not os.path.exists(fn):
            # 兜底：尝试只按类型找
            alt = os.path.join(self.rules_dir, f'{project_type}_preliminary.json')
            fn = alt if os.path.exists(alt) else fn
        with open(fn, 'r', encoding='utf-8') as f:
            return json.load(f)

    # --- 公开接口 ---
    def generate(self, project_type, design_stage, excel_data, params=None):
        params = params or {}
        blocks = self._build_blocks(project_type, design_stage, excel_data, params)
        out = self._write_docx(blocks, params)
        return out

    def preview(self, project_type, design_stage, excel_data, params=None):
        params = params or {}
        return self._build_blocks(project_type, design_stage, excel_data, params)

    # --- 构建结构化内容（generate 与 preview 共用） ---
    def _build_blocks(self, project_type, design_stage, excel_data, params):
        rule = self._load_rules(project_type, design_stage)
        summary = excel_data.get('summary', {})
        areas = excel_data.get('area_summaries', {})
        ctx = {
            'summary': summary, 'areas': areas, 'params': params,
            'project_name': params.get('project_name', '新建项目'),
            'voltage_level': params.get('voltage_level', '10kV'),
            'load_level': params.get('load_level', '二级'),
            'power_source': params.get('power_source', '两路'),
            'standby_desc': params.get('standby_desc', ''),
            'project_type': project_type, 'design_stage': design_stage,
            'rule': rule,
        }
        blocks = []
        cfg = self.cfg

        # 封面
        if cfg['cover']:
            subtitle = f"{rule.get('project_type', '')}{design_stage}·电气与自控设计说明"
            blocks.append(('cover', ctx['project_name'], subtitle,
                           '（依据《市政公用工程设计文件编制深度规定》2025年版）'))

        # 编制依据
        blocks.append(('h1', f"{cn_num(1)}、编制依据"))
        for line in self._basis_lines(ctx):
            blocks.append(('p', line))

        # 各分类（一级章节）
        cat_list = list(rule.get('categories', {}).items())
        stage_en = STAGE_CODE.get(design_stage, 'preliminary')
        for ci, (cat_name, items) in enumerate(cat_list, start=2):
            blocks.append(('h1', f"{cn_num(ci)}、{cat_name}"))
            for it in items:
                title = it['title']
                # 二级标题（去掉可能冗余的“设计”后缀，避免“电气设计设计范围”式重复）
                h2 = title if not (cat_name.endswith('设计') and title.startswith('设计')) else title[2:]
                blocks.append(('h2', h2))
                # 可研（可行性研究）阶段：规范仅要求“说明/确定”原则性内容，
                # 直接渲染规范条文中的深度要求文字（lighter 深度），仅用电负荷
                # 栏目（has_calculation=True）仍由内容引擎按 Excel 生成负荷汇总。
                if stage_en == 'feasibility' and not it.get('has_calculation'):
                    req = (it.get('requirement') or '').strip()
                    if req:
                        for para in re.split(r'\n+', req):
                            para = para.strip()
                            if para:
                                blocks.append(('p', para))
                    else:
                        content = self._content(cat_name, title, it, ctx)
                        for blk in content:
                            blocks.append(blk)
                else:
                    content = self._content(cat_name, title, it, ctx)
                    for blk in content:
                        blocks.append(blk)

        # 主要设备材料表
        blocks.append(('h1', f"{cn_num(len(cat_list) + 2)}、主要设备材料表"))
        blocks.append(('p', '本工程主要电气设备材料见表，具体规格型号以施工图及订货技术条件为准。'))
        rows = self._equipment_rows(ctx)
        headers = ['序号', '名称', '型号及规格', '单位', '数量', '备注']
        blocks.append(('table', headers, rows))

        if cfg['sig_block']:
            blocks.append(('page',)) if False else None
            blocks.append(('sig',))

        return blocks

    # --- 内容引擎：按栏目生成专业正文 ---
    def _content(self, cat, title, section, ctx):
        t = title
        s = ctx['summary']
        p = ctx['params']
        # 用电负荷（含计算）
        if '用电负荷' in t or ('负荷' in t and ('等级' in t or '计算' in t)):
            return self._c_load(ctx)
        if '电源' in t or '供电电源' in t:
            return self._c_power(ctx)
        if '变电所' in t:
            return self._c_substation(ctx)
        if '供配电系统' in t or ('供配电' in t and '设计' in t):
            return self._c_power_system(ctx)
        if '保护和控制' in t or ('保护' in t and '控制' in t):
            return self._c_protection(ctx)
        if '计量' in t:
            return self._c_meter(ctx)
        if '管线敷设' in t or '管缆敷设' in t or ('电缆' in t and '敷设' in t):
            return self._c_cabling(ctx)
        if '照明' in t:
            return self._c_lighting(ctx)
        if '设备选型' in t and '布置' not in t:
            return self._c_equip_selection(ctx)
        if '防雷' in t or '接地' in t:
            return self._c_earthing(ctx)
        if '自控系统' in t or '自控' in t:
            return self._c_auto(ctx)
        if '仪表系统' in t or ('仪表' in t and '自控' not in t):
            return self._c_instrument(ctx)
        if '工业电视' in t:
            return self._c_cctv(ctx)
        if '安防' in t:
            return self._c_security(ctx)
        if '通信' in t:
            return self._c_comm(ctx)
        if '火灾' in t:
            return self._c_fire(ctx)
        if '有线电视' in t:
            return self._c_catv(ctx)
        if '防护等级' in t:
            return self._c_grade(ctx)
        if '防爆' in t or '防损坏' in t:
            return self._c_explosion(ctx)
        if '设备选型与布置' in t:
            return self._c_equip_layout(ctx)
        if '机电抗震' in t or '抗震' in t:
            return self._c_seismic(ctx)
        if '设计范围' in t:
            return self._c_scope(cat, ctx)
        if '设计依据' in t or '依据' in t:
            return [('p', l) for l in self._basis_lines(ctx)]
        if '通信网络' in t or '系统配置' in t:
            return self._c_auto(ctx)
        # 兜底：基于深度要求生成通用专业段落（非占位符）
        return [('p', self._generic(cat, title, section, ctx))]

    # --- 各栏目正文 ---
    def _basis_lines(self, ctx):
        return [
            '1) 《市政公用工程设计文件编制深度规定》（2025年版）；',
            '2) 《供配电系统设计规范》GB 50052；',
            '3) 《低压配电设计规范》GB 50054；',
            '4) 《建筑物防雷设计规范》GB 50057；',
            '5) 《电力工程电缆设计标准》GB 50217；',
            '6) 《自动化仪表工程施工及质量验收规范》GB 50093；',
            '7) 相关专业提供的用电负荷资料及工艺条件图。',
        ]

    def _c_scope(self, cat, ctx):
        s = ctx['summary']
        if '电气' in cat:
            return [('p', f"本工程电气设计范围包括：供配电系统、照明、防雷接地、电缆敷设及电力监控等全部电气设计内容；"
                          f"负荷计算采用需要系数法，全厂安装容量约 {s.get('total_equip_power', 0)} kW，"
                          f"计算负荷约 {s.get('total_sc_before', 0)} kVA（补偿前）。")]
        return [('p', f"本工程自控设计范围包括：工艺过程检测、自动控制、计算机监控、仪表选型与布置、"
                     f"工业电视、安防通信及防雷接地等全部自控与弱电设计内容。")]

    def _c_load(self, ctx):
        s = ctx['summary']
        areas = ctx['areas']
        blocks = []
        blocks.append(('p', f"用电负荷主要为工艺设备电动机、风机、水泵及辅助用电。采用需要系数法计算，"
                            f"全厂设备安装容量约 {s.get('total_equip_power', 0)} kW，总有功计算负荷 "
                            f"{s.get('total_pjs', 0)} kW，自然功率因数 {s.get('cos_before', 0.85)}。"
                            f"低压侧设无功补偿装置，补偿容量约 {round(s.get('qc_compensation', 0))} kvar，"
                            f"补偿后功率因数提高到 {s.get('cos_target', 0.92)}，补偿后视在计算负荷约 "
                            f"{s.get('total_sc_after', 0)} kVA。各区域负荷计算如下表："))
        headers = ['区域 / 设备组', '设备容量(kW)', '计算负荷(kW)', '功率因数', '视在负荷(kVA)']
        rows = []
        for k, v in areas.items():
            cos = v.get('cos') or 0.85
            rows.append([k, f"{v.get('pe', 0)}", f"{v.get('pjs', 0)}", f"{cos}", f"{v.get('sc', 0)}"])
        rows.append(['合计', f"{s.get('total_equip_power', 0)}", f"{s.get('total_pjs', 0)}",
                     f"{s.get('cos_after', s.get('cos_target', 0.92))}", f"{s.get('total_sc_after', 0)}"])
        blocks.append(('table', headers, rows))
        return blocks

    def _c_power(self, ctx):
        p = ctx['params']
        vl = ctx['voltage_level']
        ll = ctx['load_level']
        ps = ctx['power_source']
        sb = ctx['standby_desc'] or '两路电源互为备用'
        return [('p', f"本工程用电负荷等级为{ll}负荷。供电电源引自{sb}。"
                        f"电源电压等级为{vl}，由城市电网或上级变电站引接，正常运行时两路电源分列运行、互为备用，"
                        f"当一路电源故障或检修时，另一路电源可保证{ll}负荷的供电可靠性。")]

    def _c_substation(self, ctx):
        s = ctx['summary']
        tx = s.get('recommended_transformer', '1×1600kVA')
        cap = s.get('tx_capacity', 1600)
        cnt = s.get('tx_count', 1)
        rate = s.get('load_rate', 0)
        rate_pct = round(rate * 100, 1)
        if rate_pct <= 85:
            econ = f"变压器负荷率为 {rate_pct}%，满足经济合理运行要求。"
        elif rate_pct <= 95:
            econ = f"变压器负荷率为 {rate_pct}%，处于较高但可接受的运行区间，留有适当裕度。"
        else:
            econ = f"变压器负荷率为 {rate_pct}%，接近满载运行，仅在大容量水厂/污水厂满负荷工况下采用，设计已预留发展裕度。"
        return [('p', f"根据全厂计算负荷及平面布置，厂内设置变配电所。选用 {tx} 干式变压器（SCB14 系列，"
                        f"接线组别 Dyn11，防护等级 IP20），低压侧采用单母线分段接线。"
                        f"补偿后全厂视在计算负荷约 {s.get('total_sc_after', 0)} kVA，{econ}")]

    def _c_power_system(self, ctx):
        p = ctx['params']
        ll = ctx['load_level']
        return [('p', f"厂内高压配电电压为 {ctx['voltage_level']}，高压系统采用单母线分段接线，"
                        f"两路电源分别接入两段母线，设母联开关，{ll}负荷供电可靠性满足规范要求。"
                        f"低压配电采用放射式与树干式结合的方式，对重要负荷（如取/送水泵、鼓风机）"
                        f"采用双回路供电并在末端互投。高、低压出线均采用电缆沿桥架、电缆沟或穿管敷设。")]

    def _c_protection(self, ctx):
        return [('p', '高压侧采用微机综合保护装置，配置过流、速断、零序及变压器温度保护；'
                     '操作电源采用直流屏（220V 或 110V）供电。高压电动机设差动或速断保护，'
                     '低压电动机回路设短路、过载及断相保护，单机容量较大者采用变频或软起动方式，'
                     '其余采用直接起动。低压配电线路装设断路器进行短路与过负荷保护。')]

    def _c_meter(self, ctx):
        p = ctx['params']
        return [('p', '根据不同用户类别分类计量：生产工艺用电在变配电所高、低压侧设多功能电能表，'
                     '采用高供高量（或高供低量）计量方式；商业用电、办公及照明等分回路装设电能表，'
                     '满足内部考核与能效管理要求。计量装置精度不低于 0.5S 级，并预留远传接口接入电力监控系统。')]

    def _c_cabling(self, ctx):
        return [('p', '室内外电气管线敷设遵循以下原则：10kV 及 0.4kV 电力电缆、控制电缆分层敷设于电缆桥架，'
                     '进出建筑物及穿越道路处穿镀锌钢管保护；直埋电缆埋深不小于 0.7m 并铺砂盖砖；'
                     '电缆沟内电缆分层排列、挂牌标识；弱电管线与强电管线平行净距不小于 0.5m，交叉处采取隔离措施。')]

    def _c_lighting(self, ctx):
        return [('p', '主要生产车间照度标准按《建筑照明设计标准》GB 50034 执行：泵房、加药间等处照度不低于 150 lx，'
                     '控制室不低于 300 lx，道路照明不低于 15 lx。一般场所采用 LED 高效灯具，'
                     '控制室、配电室、疏散通道及重要设备房间设应急照明（自带蓄电池，持续供电时间不小于 30 min），'
                     '满足火灾时人员安全疏散要求。')]

    def _c_equip_selection(self, ctx):
        return [('p', '电气设备选型遵循安全、可靠、技术先进、经济合理原则：10kV 开关柜采用 KYN28A 中置式金属封闭开关柜；'
                     '低压开关柜采用 MNS 抽出式开关柜；变压器选用 SCB14 干式变压器（H 级绝缘、带 IP20 外壳及温控风机）；'
                     '电缆采用 ZR-YJV 阻燃交联聚乙烯绝缘电力电缆；接触腐蚀性环境处采用防腐型设备及镀锌桥架。')]

    def _c_earthing(self, ctx):
        return [('p', '防雷按第三类（或第二类）建筑物设防，屋面设避雷带（网格不大于 10m×10m）并良好接地；'
                     '变配电所、控制室及高大设备装设避雷器。接地系统采用 TN-S 制，'
                     '全厂设统一的联合接地网，接地电阻不大于 1Ω；电气设备金属外壳、桥架、管道等均作等电位联结。'
                     '爆炸危险环境内的设备接地可靠，并独立引下线。')]

    def _c_auto(self, ctx):
        return [('p', '自控系统按“分散控制、集中管理、数据共享”原则，采用分层分布式结构：'
                     '中央控制室设监控上位机（含操作员站、工程师站及大屏），各工艺单元设 PLC 控制分站，'
                     '现场设检测仪表与机旁操作箱。控制分站通过工业以太网（光纤环网，1000 Mbps）与中央监控连接，'
                     '实现数据采集、过程控制、报警、趋势及报表功能，支持自动 / 手动无扰切换。')]

    def _c_instrument(self, ctx):
        return [('p', '根据工艺要求配置在线检测仪表：进出水流量、压力、液位、pH、浊度、溶解氧、余氯、'
                     '污泥浓度等，关键仪表选用高精度、免维护型并配套预处理取样装置。'
                     "仪表信号以 4~20mA 及总线方式上传 PLC，重要参数设高低限报警。")]

    def _c_cctv(self, ctx):
        return [('p', '工业电视系统在主厂房、变配电所、主要设备区及厂界设网络摄像机，'
                     '经厂内安防专网接入中央控制室电视墙，实现关键区域实时监视与录像存储（不小于 30 天）。')]

    def _c_security(self, ctx):
        return [('p', '安防系统含入侵报警、门禁与视频监控，重要房间（控制室、配电室、加药间）设门禁与红外报警，'
                     '与工业电视及消防系统联动，事件信息上传中央控制室统一管理。')]

    def _c_comm(self, ctx):
        return [('p', '通信系统设厂内行政电话与调度电话，生产区采用 VoIP 与无线对讲结合；'
                     '综合布线按六类系统敷设，主干采用光缆，满足数据、语音及视频传输需求，并预留外网接口。')]

    def _c_fire(self, ctx):
        return [('p', '按《火灾自动报警系统设计规范》GB 50116 设置火灾自动报警系统：控制室、配电室、'
                     '电缆夹层及走廊设感烟/感温探测器，厂区设手动报警按钮与声光警报器，'
                     '报警信号接入消防控制室并联动应急照明、排烟及门禁，与城市消防系统联网。')]

    def _c_catv(self, ctx):
        return [('p', '厂内设置有线电视及信息发布系统，信号引自城市有线电视网或 IPTV，'
                     '在控制室、休息室及会议室设终端，满足生产与后勤信息接收需求。')]

    def _c_grade(self, ctx):
        return [('p', '室外及潮湿环境用电设备防护等级不低于 IP54，控制箱 / 配电箱不低于 IP54，'
                     '腐蚀环境采用 IP65 防腐型并做不锈钢或热浸锌处理；室内干燥场所不低于 IP30。')]

    def _c_explosion(self, ctx):
        return [('p', '对存在爆炸性气体或粉尘的危险场所，选用相应防爆等级（Ex dⅡBT4 或 Ex tD A21）的电气设备与灯具，'
                     '管线采用镀锌钢管明配并做隔离密封；对易损设备采取防机械损坏的护栏或护罩保护。')]

    def _c_equip_layout(self, ctx):
        return [('p', '自控与弱电设备集中布置于中央控制室及各区 PLC 控制柜，柜体采用 800×600×2200 标准尺寸，'
                     '前后维护通道不小于 1.2m；现场仪表就近于工艺设备旁立柱安装，取样管路短捷并设排污。')]

    def _c_seismic(self, ctx):
        return [('p', '按《建筑机电工程抗震设计规范》GB 50981，对配电柜、控制柜、桥架及直径不小于 0.7m 的管道'
                     '设侧向及纵向抗震支吊架，抗震设防烈度与主体建筑一致，保证地震时机电系统不脱落、不倾覆。')]

    def _generic(self, cat, title, section, ctx):
        req = section.get('depth_requirement', '')
        if cat.endswith('设计'):
            dom = cat[:-2]
            return f"本{title}按规范深度要求编制：{req}。设计结合本工程工艺特点与负荷性质，做到技术可靠、经济合理、便于运行维护。"
        return f"本{title}按规范深度要求编制：{req}。"

    # --- 设备材料表 ---
    def _equipment_rows(self, ctx):
        s = ctx['summary']
        areas = ctx['areas']
        n_area = max(s.get('area_count', 1), 1)
        n_10kv = sum(1 for k in areas if '10kv' in k.lower() or '10kV' in k)
        n_lv = max(n_area - n_10kv, 1)
        qc = round(s.get('qc_compensation', 0)) or (s.get('tx_capacity', 1600) // 3)
        n_dev = s.get('total_devices', 0)
        rows = [
            ['1', '干式变压器', f'SCB14-{s.get("tx_capacity", 1600)}kVA Dyn11', '台', str(s.get('tx_count', 1)), '带温控风机'],
            ['2', '10kV 中置式开关柜', 'KYN28A-12', '台', str(max(n_10kv, 1) * 2), '含微机保护'],
            ['3', '低压抽出式开关柜', 'MNS', '台', str(n_lv + 3), '含变频/软起'],
            ['4', '低压无功补偿装置', f'智能型 {qc} kvar', '套', '1', '自动分组投切'],
            ['5', 'PLC 控制柜', '模块化', '套', str(min(n_area, 12)), '含 I/O 模块'],
            ['6', '工艺检测仪表', '在线分析', '套', str(max(n_dev // 4, 8)), '流量/液位/水质'],
            ['7', '中央监控计算机', '工业级', '套', '1', '双机热备'],
            ['8', 'UPS 不间断电源', '10kVA/30min', '套', '1', '供监控用电'],
            ['9', '工业以太网交换机', '千兆', '台', str(max(n_area, 4)), '环网冗余'],
            ['10', '电力监控系统', '成套软件', '套', '1', '含通讯管理机'],
        ]
        return rows

    # --- 写入 docx ---
    def _write_docx(self, blocks, params):
        cfg = self.cfg
        doc = Document()
        # 基础样式
        style = doc.styles['Normal']
        style.font.name = '宋体'
        style.font.size = Pt(cfg['body'])
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        try:
            style.paragraph_format.line_spacing = cfg['line']
        except Exception:
            pass

        for blk in blocks:
            kind = blk[0]
            if kind == 'cover':
                self._add_cover(doc, blk[1], blk[2], blk[3])
            elif kind == 'h1':
                p = doc.add_heading(level=1)
                run = p.add_run(blk[1])
                run.font.name = '黑体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                run.font.size = Pt(cfg['h1'])
                run.font.bold = True
                run.font.color.rgb = cfg['title_color']
            elif kind == 'h2':
                p = doc.add_paragraph()
                run = p.add_run(blk[1])
                run.font.name = '黑体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                run.font.size = Pt(cfg['h2'])
                run.font.bold = True
                run.font.color.rgb = cfg['accent']
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(3)
            elif kind == 'h3':
                p = doc.add_paragraph()
                run = p.add_run(blk[1])
                run.font.name = '黑体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                run.font.size = Pt(cfg['h3'])
                run.font.bold = True
            elif kind == 'p':
                p = doc.add_paragraph()
                run = p.add_run(blk[1])
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run.font.size = Pt(cfg['body'])
                p.paragraph_format.line_spacing = cfg['line']
                p.paragraph_format.space_after = Pt(3)
            elif kind == 'table':
                self._add_table(doc, blk[1], blk[2], cfg)
            elif kind == 'sig':
                self._add_sig(doc, cfg)
            elif kind == 'page':
                doc.add_page_break()

        # 页眉 / 页脚
        if cfg['header']:
            self._add_header(doc, params.get('project_name', '新建项目'))
        if cfg['page_number']:
            self._add_page_number(doc)

        ts = datetime.now().strftime('%Y%m%d_%H%M%S')
        safe = ''.join(ch for ch in params.get('project_name', '项目') if ch not in '\\/:*?"<>|')
        rule = self._load_rules(params.get('project_type', 'water_supply'),
                                params.get('design_stage', '初步设计'))
        ptype = rule.get('project_type', '工程')
        stage = params.get('design_stage', '初步设计')
        fn = f"{safe}_{ptype}_{stage}电气自控说明_{self.template}_{ts}.docx"
        out = os.path.join(self.output_dir, fn)
        doc.save(out)
        return out

    def _add_cover(self, doc, name, subtitle, note):
        for _ in range(6):
            doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(name)
        r.font.size = Pt(26)
        r.font.bold = True
        r.font.name = '黑体'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        r.font.color.rgb = self.cfg['title_color']
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(subtitle)
        r.font.size = Pt(15)
        r.font.name = '黑体'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"\n{note}")
        r.font.size = Pt(10.5)
        for _ in range(8):
            doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"编制日期：{datetime.now().strftime('%Y 年 %m 月')}")
        r.font.size = Pt(11)
        doc.add_page_break()

    def _add_table(self, doc, headers, rows, cfg):
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].text = ''
            para = hdr[i].paragraphs[0]
            run = para.add_run(str(h))
            run.font.bold = True
            run.font.size = Pt(9.5)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            if cfg.get('header_shade') and cfg['header_shade'] != (None,):
                self._shade_cell(hdr[i], cfg['header_shade'])
        for row in rows:
            cells = table.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = ''
                para = cells[i].paragraphs[0]
                run = para.add_run(str(val))
                run.font.size = Pt(9.5)
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        # 列宽自适应
        try:
            widths = [Cm(1.0)] + [Cm(max(2.2, 8.0 / len(headers))) for _ in range(len(headers) - 1)]
            for row in table.rows:
                for i, c in enumerate(row.cells):
                    if i < len(widths):
                        c.width = widths[i]
        except Exception:
            pass

    @staticmethod
    def _shade_cell(cell, color):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '%02X%02X%02X' % (color[0], color[1], color[2]))
        tcPr.append(shd)

    def _add_sig(self, doc, cfg):
        doc.add_paragraph()
        p = doc.add_paragraph()
        r = p.add_run('编制与签署')
        r.font.bold = True
        r.font.size = Pt(11)
        table = doc.add_table(rows=4, cols=4)
        table.style = 'Table Grid'
        labels = [['编制', '', '校核', ''], ['审核', '', '审定', ''],
                  ['项目负责人', '', '专业负责人', ''], ['日期', '', '单位', '']]
        for ri, row in enumerate(labels):
            for ci, val in enumerate(row):
                table.rows[ri].cells[ci].text = val
        doc.add_paragraph()

    def _add_header(self, doc, project_name):
        section = doc.sections[0]
        section.header.is_linked_to_previous = False
        p = section.header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(f"{project_name} · 电气与自控设计说明")
        run.font.size = Pt(9)
        run.font.color.rgb = self.cfg['accent']

    def _add_page_number(self, doc):
        """稳健的页码域：每个域指令各占独立 run（作为段落子元素）。"""
        section = doc.sections[0]
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p.add_run('— 第 ')
        run.font.size = Pt(10)
        # 域开始
        fld1 = OxmlElement('w:fldChar'); fld1.set(qn('w:fldCharType'), 'begin')
        instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = ' PAGE '
        fld2 = OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'), 'end')
        run._r.append(fld1); run._r.append(instr); run._r.append(fld2)

        run2 = p.add_run(' 页 —')
        run2.font.size = Pt(10)


if __name__ == '__main__':
    import sys
    sys.path.insert(0, os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))), 'backend'))
    from app.services.excel_parser import parse as parse_excel
    ep = r'D:\00-水厂负荷计算表.xlsx'
    ed = parse_excel(ep)
    gen = DocxGenerator(rules_dir=RULES_DIR, output_dir=OUTPUT_DIR, template='standard')
    params = {'project_name': '赤壁中心水厂', 'voltage_level': '10kV', 'load_level': '二级',
              'project_type': 'water_supply', 'power_source': '两路', 'standby_desc': '两路电源互为备用',
              'design_stage': '初步设计'}
    for tpl in TEMPLATES:
        gen.set_template(tpl)
        out = gen.generate('water_supply', '初步设计', ed, params)
        print(f'[{tpl}] -> {out}')
