"""Extract all design documents to Markdown."""
import os
import sys
import traceback
import io

SRC_DIR = r"D:\qoderwork\可研初步设计电气文本TESTB\source"
OUT_DIR = r"D:\qoderwork\可研初步设计电气文本TESTB\output\extracted"
os.makedirs(OUT_DIR, exist_ok=True)

def safe(msg):
    """Print with safe ASCII-only fallback."""
    try:
        print(msg)
    except UnicodeEncodeError:
        print(msg.encode('ascii', errors='replace').decode('ascii'))


def repair_docx(path):
    """Remove problematic entries (like word/NULL) from a DOCX zip."""
    import zipfile
    import shutil
    tmp = path + ".repaired.tmp"
    found_bad = False
    with zipfile.ZipFile(path, 'r') as zin:
        with zipfile.ZipFile(tmp, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                if item.filename == 'word/NULL':
                    safe(f"  [!] Skipping corrupt entry: {item.filename}")
                    found_bad = True
                    continue
                data = zin.read(item.filename)
                zout.writestr(item, data)
    if found_bad:
        shutil.move(tmp, path)
        safe("  [OK] Repaired DOCX (removed word/NULL entry)")
    else:
        os.unlink(tmp)
    return path


# ──────────────────────────────────────────────
# 1. python-docx extract
# ──────────────────────────────────────────────
def extract_docx(path, out_path):
    from docx import Document
    try:
        doc = Document(path)
    except Exception as e:
        safe(f"  [!] docx open failed: {e}, attempting repair...")
        path = repair_docx(path)
        doc = Document(path)

    lines = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        is_heading = False
        if p.style and p.style.name and p.style.name.startswith("Heading"):
            is_heading = True
        if not is_heading and p.runs:
            is_heading = any(r.bold for r in p.runs if r.bold)
        prefix = "## " if is_heading else ""
        lines.append(prefix + text)

    # Also extract tables
    table_text = []
    for ti, table in enumerate(doc.tables, 1):
        table_text.append(f"### 表格 {ti}")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            table_text.append(" | ".join(cells))

    out_text = "\n\n".join(lines)
    if table_text:
        out_text += "\n\n---\n\n" + "\n\n".join(table_text)

    with open(out_path, "w", encoding="utf-8") as f:
        f.write(out_text)
    safe(f"  [OK] {os.path.basename(path)} -> {len(lines)} paragraphs + {len(doc.tables)} tables")
    return out_path


# ──────────────────────────────────────────────
# 2. .doc -> convert via Word COM -> docx -> extract
# ──────────────────────────────────────────────
def extract_doc_via_word(path, out_path):
    import comtypes.client
    import tempfile
    import pythoncom
    word = None
    tmp_docx = None
    try:
        # Initialize COM for this thread
        pythoncom.CoInitialize()
        word = comtypes.client.CreateObject("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False

        abs_path = os.path.abspath(path)
        safe(f"  [..] Opening Word document (may take a while for 26MB)...")
        doc = word.Documents.Open(abs_path)

        tmp_docx = tempfile.NamedTemporaryFile(suffix=".docx", delete=False).name
        doc.SaveAs(tmp_docx, FileFormat=16)  # wdFormatDocumentDefault = .docx
        doc.Close()

        size_mb = os.path.getsize(tmp_docx) / 1024 / 1024
        safe(f"  [OK] Converted .doc -> .docx ({size_mb:.1f} MB)")

        result = extract_docx(tmp_docx, out_path)
        return result
    except Exception as e:
        safe(f"  [FAIL] Word COM: {e}")
        traceback.print_exc()
        try:
            return extract_doc_raw_ole(path, out_path)
        except Exception as e2:
            safe(f"  [FAIL] OLE fallback: {e2}")
            return None
    finally:
        if word:
            try:
                word.Quit()
            except Exception:
                pass
        if tmp_docx and os.path.exists(tmp_docx):
            try:
                os.unlink(tmp_docx)
            except Exception:
                pass
        try:
            pythoncom.CoUninitialize()
        except Exception:
            pass


# ──────────────────────────────────────────────
# 3. Fallback: raw OLE text extraction
# ──────────────────────────────────────────────
def extract_doc_raw_ole(path, out_path):
    """Minimal OLE2 text extractor for old .doc files via UTF-16LE decode."""
    safe("  [..] Trying raw OLE text extraction...")
    with open(path, "rb") as f:
        data = f.read()
    text = data.decode("utf-16le", errors="ignore")
    lines = []
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue
        printable = sum(1 for c in line if c.isprintable()
                        and (c.isascii()
                             or '\u4e00' <= c <= '\u9fff'
                             or '\u3000' <= c <= '\u303f'))
        if printable > len(line) * 0.5 and len(line) > 2:
            lines.append(line)
    if lines:
        out_text = "\n\n".join(lines)
        with open(out_path, "w", encoding="utf-8") as f:
            f.write(out_text)
        safe(f"  [WARN] OLE raw extract: {len(lines)} lines -> {out_path}")
        return out_path
    return None


# ──────────────────────────────────────────────
# 4. pdfplumber extract
# ──────────────────────────────────────────────
def extract_pdf(path, out_path):
    import pdfplumber
    with pdfplumber.open(path) as pdf:
        pages_text = []
        for i, page in enumerate(pdf.pages, 1):
            t = page.extract_text()
            if t:
                pages_text.append(f"--- Page {i} ---\n{t}")
            else:
                pages_text.append(f"--- Page {i} (no text) ---")
    out_text = "\n\n".join(pages_text)
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(out_text)
    safe(f"  [OK] PDF: {len(pdf.pages)} pages -> {out_path}")
    return out_path


# ──────────────────────────────────────────────
# Main
# ──────────────────────────────────────────────
files = [
    ("taochong.docx",     "陶冲污水处理厂三期工程项目-初设文本0408.md", extract_docx),
    ("chibi.docx",        "赤壁中心水厂技改可研电气部分.md", extract_docx),
    ("nanchang.doc",      "南昌取水口优化调整可研报告2021.md", extract_doc_via_word),
    ("编制深度规定.pdf",  "市政公用工程设计文件编制深度规定2025.md", extract_pdf),
]

results = {}

for src_name, out_name, extractor in files:
    src_path = os.path.join(SRC_DIR, src_name)
    out_path = os.path.join(OUT_DIR, out_name)
    safe("\n" + "=" * 60)
    size_mb = os.path.getsize(src_path) / 1024 / 1024
    safe(f"Extracting: {src_name} ({size_mb:.1f} MB)")
    safe("=" * 60)
    if not os.path.exists(src_path):
        safe(f"  [FAIL] FILE NOT FOUND: {src_path}")
        continue
    try:
        result = extractor(src_path, out_path)
        if result:
            results[out_name] = out_path
        else:
            safe(f"  [FAIL] Extraction returned no result")
    except Exception as e:
        safe(f"  [FAIL] Error: {e}")
        traceback.print_exc()

safe("\n" + "=" * 60)
safe("EXTRACTION COMPLETE")
safe("=" * 60)
for name, path in results.items():
    size_kb = os.path.getsize(path) / 1024
    safe(f"  {name}: {size_kb:.0f} KB")
