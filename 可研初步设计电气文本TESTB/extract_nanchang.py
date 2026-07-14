"""Extract nanchang.doc via Word COM (without pythoncom)."""
import os
import sys
import comtypes.client
import tempfile

src = r"D:\qoderwork\可研初步设计电气文本TESTB\source\nanchang.doc"
out = r"D:\qoderwork\可研初步设计电气文本TESTB\output\extracted\南昌取水口优化调整可研报告2021.md"

# Initialize COM for this thread
import ctypes
ole32 = ctypes.OleDLL('ole32')
ole32.CoInitialize(None)

word = None
tmp_docx = None
try:
    word = comtypes.client.CreateObject("Word.Application")
    word.Visible = False
    word.DisplayAlerts = False

    abs_path = os.path.abspath(src)
    print(f"Opening Word document: {src}")
    doc = word.Documents.Open(abs_path)

    tmp_docx = tempfile.NamedTemporaryFile(suffix=".docx", delete=False).name
    doc.SaveAs(tmp_docx, FileFormat=16)
    doc.Close()
    print(f"Saved as DOCX: {tmp_docx}")

    # Now extract with python-docx
    from docx import Document
    doc2 = Document(tmp_docx)
    lines = []
    for p in doc2.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        is_heading = False
        if p.style and p.style.name and p.style.name.startswith("Heading"):
            is_heading = True
        if not is_heading and p.runs:
            is_heading = any(r.bold for r in p.runs if r.bold)
        prefix = "## " if is_heading else ""
        lines.append(prefix + text)

    # Tables
    table_text = []
    for ti, table in enumerate(doc2.tables, 1):
        table_text.append(f"### 表格 {ti}")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            table_text.append(" | ".join(cells))

    out_text = "\n\n".join(lines)
    if table_text:
        out_text += "\n\n---\n\n" + "\n\n".join(table_text)

    with open(out, "w", encoding="utf-8") as f:
        f.write(out_text)

    print(f"\nOutput: {out}")
    print(f"Paragraphs: {len(lines)}, Tables: {len(doc2.tables)}")

except Exception as e:
    print(f"Error: {e}")
    import traceback
    traceback.print_exc()
finally:
    if word:
        try:
            word.Quit()
        except:
            pass
    if tmp_docx and os.path.exists(tmp_docx):
        try:
            os.unlink(tmp_docx)
        except:
            pass
    try:
        ole32.CoUninitialize()
    except:
        pass
