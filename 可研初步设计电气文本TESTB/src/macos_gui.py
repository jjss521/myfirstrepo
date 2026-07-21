#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
macOS Sonoma-style customtkinter GUI for 市政工程电气自控设计说明书生成器.

Replaces the ttkbootstrap GUI with a modern, sidebar-navigated desktop app.
All UI code lives in this single file — zero backend changes required.

Launch:  python src/macos_gui.py
"""

import os
import sys
import json
import threading
import traceback
from datetime import datetime
from typing import Optional, Dict, Any, List, Tuple
from pathlib import Path

# ── Path setup (mirrors main.py) ───────────────────────────────────────────
_SRC_DIR = os.path.dirname(os.path.abspath(__file__))
if _SRC_DIR not in sys.path:
    sys.path.insert(0, _SRC_DIR)

# Allow importing backend modules (for DocxGenerator preview)
_PROJECT_ROOT = os.path.dirname(_SRC_DIR)
_BACKEND_DIR = os.path.join(_PROJECT_ROOT, 'backend')
if _BACKEND_DIR not in sys.path:
    sys.path.insert(0, _BACKEND_DIR)

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext

import customtkinter as ctk

from core.engine import GenerateEngine

# ── Optional: tkinterdnd2 for drag-drop ────────────────────────────────────
_TKINTERDND2_AVAILABLE = False
try:
    from tkinterdnd2 import DND_FILES  # noqa: F401
    _TKINTERDND2_AVAILABLE = True
except ImportError:
    pass

# ── Backend imports for preview ────────────────────────────────────────────
_DOCX_PREVIEW_AVAILABLE = False
DocxGenerator = None  # type: ignore[assignment]
try:
    from app.services.docx_generator import DocxGenerator  # noqa: F811
    _DOCX_PREVIEW_AVAILABLE = True
except ImportError:
    pass

from app.services.equipment_table_generator import EquipmentTableGenerator
from app.services.historical_template_importer import HistoricalTemplateImporter


# ═══════════════════════════════════════════════════════════════════════════
# 1. DESIGN SYSTEM CONSTANTS
# ═══════════════════════════════════════════════════════════════════════════

COLORS: Dict[str, str] = {
    "bg_primary":    "#F5F5F7",
    "bg_sidebar":    "#E8E8ED",
    "bg_card":       "#FFFFFF",
    "bg_hover":      "#E8E8ED",
    "bg_selected":   "#D2D2D7",
    "accent_blue":   "#007AFF",
    "accent_green":  "#34C759",
    "accent_orange": "#FF9500",
    "accent_red":    "#FF3B30",
    "text_primary":  "#1D1D1F",
    "text_secondary":"#86868B",
    "text_tertiary": "#C7C7CC",
    "separator":     "#D2D2D7",
}

FONTS: Dict[str, Dict[str, Any]] = {
    "header":       {"family": "微软雅黑", "size": 20, "weight": "bold"},
    "section":      {"family": "微软雅黑", "size": 15, "weight": "bold"},
    "body":         {"family": "微软雅黑", "size": 13},
    "caption":      {"family": "微软雅黑", "size": 11},
    "mono":         {"family": "Consolas", "size": 11},
    "title":        {"family": "微软雅黑", "size": 24, "weight": "bold"},
    "sidebar_item": {"family": "微软雅黑", "size": 13},
}


def get_ctk_font(key: str) -> ctk.CTkFont:
    """Create a CTkFont from the FONTS params dict. Safe to call after CTk root exists."""
    params = dict(FONTS[key])
    return ctk.CTkFont(**params)

APP_TITLE    = "市政工程电气自控设计说明书生成器"
APP_VERSION  = "1.0.0"
PROJECT_TYPES: List[tuple] = [
    ("water_supply", "给水工程"),
    ("drainage",     "排水工程"),
    ("road",         "道路工程"),
    ("sanitation",   "环卫工程"),
]
DESIGN_STAGES  = ["可行性研究", "初步设计"]
VOLTAGE_LEVELS = ["10kV", "35kV", "0.4kV", "6kV"]
LOAD_LEVELS    = ["一级", "二级", "三级"]

DEFAULT_PARAMS: Dict[str, Any] = {
    "project_name":   "",
    "project_type":   "water_supply",
    "design_stage":   "初步设计",
    "voltage_level":  "10kV",
    "load_level":     "二级",
    "power_source":   "一路",
    "standby_desc":   "",
    "tx_config":      "",
    "tx_count":       "1",
    "tx_location":    "变配电间内",
    "project_date":   datetime.now().strftime("%Y年%m月%d日"),
}

_RECENT_DIR = os.path.join(os.environ.get("APPDATA", os.path.expanduser("~")),
                           "电气自控生成器")
_RECENT_FILE = os.path.join(_RECENT_DIR, "recent.json")

# ═══════════════════════════════════════════════════════════════════════════
# 2. EngineContext
# ═══════════════════════════════════════════════════════════════════════════

class EngineContext:
    """Holds engine instance, current data, and cached state."""

    def __init__(self) -> None:
        project_root = os.path.dirname(_SRC_DIR)
        self.engine = GenerateEngine(project_root)
        self.excel_data: Optional[Dict[str, Any]] = None
        self.excel_path: str = ""
        self.last_output_path: str = ""
        self.project_params: Dict[str, Any] = dict(DEFAULT_PARAMS)

    def health_status(self) -> Dict[str, Any]:
        return self.engine.health_check()

    def is_healthy(self) -> bool:
        return self.engine.health_check()["all_pass"]


# ═══════════════════════════════════════════════════════════════════════════
# 3. PAGE: ProjectPage
# ═══════════════════════════════════════════════════════════════════════════

class ProjectPage(ctk.CTkFrame):
    """Card-based project info form."""

    def __init__(self, master, ctx: EngineContext, **kwargs) -> None:
        super().__init__(master, fg_color="transparent", **kwargs)
        self.ctx = ctx
        self.params: Dict[str, Any] = dict(DEFAULT_PARAMS)
        self._build()

    # ── UI ─────────────────────────────────────────────────────────────────

    def _build(self) -> None:
        # Header
        hdr = ctk.CTkFrame(self, fg_color="transparent")
        hdr.pack(fill="x", pady=(10, 20))
        ctk.CTkLabel(hdr, text="项目信息", font=get_ctk_font("header"),
                     text_color=COLORS["text_primary"]).pack(anchor="w")
        ctk.CTkLabel(hdr, text="配置工程基本参数用于文档生成",
                     font=get_ctk_font("caption"),
                     text_color=COLORS["text_secondary"]).pack(anchor="w")

        # Two-column card layout
        cols = ctk.CTkFrame(self, fg_color="transparent")
        cols.pack(fill="both", expand=True)

        # Left card: 基本信息
        left_card = ctk.CTkFrame(cols, fg_color=COLORS["bg_card"],
                                 corner_radius=12)
        left_card.pack(side="left", fill="both", expand=True, padx=(0, 8))
        self._build_left_card(left_card)

        # Right card: 电气配置
        right_card = ctk.CTkFrame(cols, fg_color=COLORS["bg_card"],
                                  corner_radius=12)
        right_card.pack(side="right", fill="both", expand=True, padx=(8, 0))
        self._build_right_card(right_card)

        # Bottom save button
        ctk.CTkButton(
            self, text="💾 保存项目信息", font=get_ctk_font("body"),
            fg_color=COLORS["accent_blue"], hover_color="#0066D6",
            corner_radius=10, height=40,
            command=self.save_params,
        ).pack(pady=(15, 0))

    def _build_left_card(self, card: ctk.CTkFrame) -> None:
        pad = {"padx": 20, "pady": (12, 4)}
        epad = {"padx": 20, "pady": (2, 15)}

        ctk.CTkLabel(card, text="基本信息", font=get_ctk_font("section"),
                     text_color=COLORS["text_primary"]).pack(anchor="w", **pad)

        # 工程名称
        ctk.CTkLabel(card, text="工程名称 *", font=get_ctk_font("caption"),
                     text_color=COLORS["text_secondary"]).pack(anchor="w", **epad)
        self.entry_name = ctk.CTkEntry(card, font=get_ctk_font("body"),
                                       placeholder_text="示例工程",
                                       corner_radius=8)
        self.entry_name.pack(fill="x", padx=20, pady=(0, 15))
        self.entry_name.insert(0, "示例工程")

        # 工程类型
        ctk.CTkLabel(card, text="工程类型 *", font=get_ctk_font("caption"),
                     text_color=COLORS["text_secondary"]).pack(anchor="w", **epad)
        self.type_var = ctk.StringVar(value="给水工程")
        self.type_combo = ctk.CTkOptionMenu(
            card, values=[lbl for _, lbl in PROJECT_TYPES],
            variable=self.type_var, font=get_ctk_font("body"),
            corner_radius=8, fg_color=COLORS["bg_primary"],
            button_color=COLORS["accent_blue"],
            command=self._on_type_change,
        )
        self.type_combo.pack(fill="x", padx=20, pady=(0, 15))

        # 设计阶段
        ctk.CTkLabel(card, text="设计阶段 *", font=get_ctk_font("caption"),
                     text_color=COLORS["text_secondary"]).pack(anchor="w", **epad)
        self.stage_var = ctk.StringVar(value="初步设计")
        ctk.CTkOptionMenu(
            card, values=DESIGN_STAGES, variable=self.stage_var,
            font=get_ctk_font("body"), corner_radius=8,
            fg_color=COLORS["bg_primary"],
            button_color=COLORS["accent_blue"],
        ).pack(fill="x", padx=20, pady=(0, 15))

        # 编制日期
        ctk.CTkLabel(card, text="编制日期", font=get_ctk_font("caption"),
                     text_color=COLORS["text_secondary"]).pack(anchor="w", **epad)
        date_row = ctk.CTkFrame(card, fg_color="transparent")
        date_row.pack(fill="x", padx=20, pady=(0, 15))
        today_str = datetime.now().strftime("%Y年%m月%d日")
        self.date_var = ctk.StringVar(value=today_str)
        ctk.CTkEntry(date_row, textvariable=self.date_var, font=get_ctk_font("body"),
                     corner_radius=8).pack(side="left", fill="x", expand=True)
        ctk.CTkButton(date_row, text="今天", width=60, font=get_ctk_font("caption"),
                      fg_color=COLORS["text_tertiary"],
                      hover_color=COLORS["bg_selected"],
                      text_color=COLORS["text_primary"],
                      command=lambda: self.date_var.set(today_str),
                      ).pack(side="left", padx=(8, 0))

    def _build_right_card(self, card: ctk.CTkFrame) -> None:
        pad = {"padx": 20, "pady": (12, 4)}
        epad = {"padx": 20, "pady": (2, 15)}

        ctk.CTkLabel(card, text="电气配置", font=get_ctk_font("section"),
                     text_color=COLORS["text_primary"]).pack(anchor="w", **pad)

        # 供电电压等级
        ctk.CTkLabel(card, text="供电电压等级", font=get_ctk_font("caption"),
                     text_color=COLORS["text_secondary"]).pack(anchor="w", **epad)
        self.voltage_var = ctk.StringVar(value="10kV")
        ctk.CTkOptionMenu(
            card, values=VOLTAGE_LEVELS, variable=self.voltage_var,
            font=get_ctk_font("body"), corner_radius=8,
            fg_color=COLORS["bg_primary"],
            button_color=COLORS["accent_blue"],
        ).pack(fill="x", padx=20, pady=(0, 15))

        # 用电负荷等级
        ctk.CTkLabel(card, text="用电负荷等级", font=get_ctk_font("caption"),
                     text_color=COLORS["text_secondary"]).pack(anchor="w", **epad)
        self.load_var = ctk.StringVar(value="二级")
        ctk.CTkOptionMenu(
            card, values=LOAD_LEVELS, variable=self.load_var,
            font=get_ctk_font("body"), corner_radius=8,
            fg_color=COLORS["bg_primary"],
            button_color=COLORS["accent_blue"],
        ).pack(fill="x", padx=20, pady=(0, 15))

        # 电源回路数
        ctk.CTkLabel(card, text="电源回路数", font=get_ctk_font("caption"),
                     text_color=COLORS["text_secondary"]).pack(anchor="w", **epad)
        self.ps_var = ctk.StringVar(value="一路")
        ctk.CTkSegmentedButton(
            card, values=["一路", "两路"], variable=self.ps_var,
            font=get_ctk_font("body"),
        ).pack(fill="x", padx=20, pady=(0, 15))

        # 备用电源说明
        ctk.CTkLabel(card, text="备用电源说明", font=get_ctk_font("caption"),
                     text_color=COLORS["text_secondary"]).pack(anchor="w", **epad)
        self.standby_text = ctk.CTkTextbox(card, height=60, font=get_ctk_font("caption"),
                                           corner_radius=8,
                                           fg_color=COLORS["bg_primary"],
                                           wrap="word")
        self.standby_text.pack(fill="x", padx=20, pady=(0, 10))
        self.standby_text.insert("1.0", "当一路电源故障时，另一路可承担全部负荷。")

        # 变配电间位置
        ctk.CTkLabel(card, text="变配电间位置", font=get_ctk_font("caption"),
                     text_color=COLORS["text_secondary"]).pack(anchor="w", **epad)
        self.tx_loc_var = ctk.StringVar(value="变配电间内")
        ctk.CTkEntry(card, textvariable=self.tx_loc_var, font=get_ctk_font("body"),
                     corner_radius=8).pack(fill="x", padx=20, pady=(0, 10))

    # ── Logic ──────────────────────────────────────────────────────────────

    def _on_type_change(self, value: str) -> None:
        """When project type changes, load rule and update design stage."""
        for code, lbl in PROJECT_TYPES:
            if lbl == value:
                rule = self.ctx.engine.load_rule(code)
                if rule:
                    stage = rule.get("design_stage", "初步设计")
                    self.stage_var.set(stage)
                break

    def save_params(self) -> None:
        """Persist form values into context."""
        label = self.type_var.get()
        project_type = "water_supply"
        for code, lbl in PROJECT_TYPES:
            if lbl == label:
                project_type = code
                break

        self.params.update({
            "project_name":   self.entry_name.get().strip() or "示例工程",
            "project_type":   project_type,
            "design_stage":   self.stage_var.get(),
            "voltage_level":  self.voltage_var.get(),
            "load_level":     self.load_var.get(),
            "power_source":   self.ps_var.get(),
            "standby_desc":   self.standby_text.get("1.0", "end-1c").strip(),
            "tx_location":    self.tx_loc_var.get().strip(),
            "project_date":   self.date_var.get(),
        })
        self.ctx.project_params = dict(self.params)
        messagebox.showinfo("保存成功", "项目信息已保存")

    def get_params(self) -> Dict[str, Any]:
        """Return current params dict, saving first."""
        self.save_params()
        return self.params


# ═══════════════════════════════════════════════════════════════════════════
# 4. PAGE: ExcelPage
# ═══════════════════════════════════════════════════════════════════════════

class ExcelPage(ctk.CTkFrame):
    """Excel import page with drag-drop zone and results table."""

    def __init__(self, master, ctx: EngineContext, **kwargs) -> None:
        super().__init__(master, fg_color="transparent", **kwargs)
        self.ctx = ctx
        self._build()

    # ── UI ─────────────────────────────────────────────────────────────────

    def _build(self) -> None:
        # Header
        hdr = ctk.CTkFrame(self, fg_color="transparent")
        hdr.pack(fill="x", pady=(10, 15))
        ctk.CTkLabel(hdr, text="导入负荷计算表", font=get_ctk_font("header"),
                     text_color=COLORS["text_primary"]).pack(side="left")
        ctk.CTkButton(hdr, text="📂 选择文件", font=get_ctk_font("body"),
                      fg_color=COLORS["accent_blue"], corner_radius=8,
                      command=self._select_file).pack(side="right", padx=(8, 0))
        self.parse_btn = ctk.CTkButton(
            hdr, text="🔄 解析并预览", font=get_ctk_font("body"),
            fg_color=COLORS["bg_selected"],
            text_color=COLORS["text_primary"],
            hover_color=COLORS["bg_hover"],
            corner_radius=8, command=self._parse_excel,
        )
        self.parse_btn.pack(side="right", padx=(8, 0))

        # Drag-drop zone
        self.drop_zone = ctk.CTkFrame(
            self, fg_color=COLORS["bg_primary"],
            border_width=2, border_color=COLORS["text_tertiary"],
            corner_radius=12, height=120,
        )
        self.drop_zone.pack(fill="x", pady=(0, 10))
        self.drop_zone.pack_propagate(False)

        inner = ctk.CTkFrame(self.drop_zone, fg_color="transparent")
        inner.pack(expand=True)
        drop_hint = "拖拽 Excel 文件到这里" if _TKINTERDND2_AVAILABLE else "📂 点击此区域选择 Excel 文件"
        ctk.CTkLabel(inner, text=drop_hint, font=get_ctk_font("section"),
                     text_color=COLORS["text_secondary"]).pack()
        ctk.CTkLabel(inner, text="支持 .xlsx / .xls 格式",
                     font=get_ctk_font("caption"),
                     text_color=COLORS["text_tertiary"]).pack(pady=(4, 0))

        # Make drop zone clickable
        self.drop_zone.bind("<Button-1>", lambda e: self._select_file())
        for child in inner.winfo_children():
            child.bind("<Button-1>", lambda e: self._select_file())
        inner.bind("<Button-1>", lambda e: self._select_file())

        # Hover effects for drop zone
        self.drop_zone.bind("<Enter>", lambda e: self._on_drop_enter())
        self.drop_zone.bind("<Leave>", lambda e: self._on_drop_leave())

        # File label
        self.file_label = ctk.CTkLabel(
            self, text="未选择文件", font=get_ctk_font("caption"),
            text_color=COLORS["text_secondary"],
        )
        self.file_label.pack(anchor="w", pady=(0, 10))

        # Results area (initially hidden)
        self.results_frame = ctk.CTkFrame(self, fg_color="transparent")

    def _on_drop_enter(self) -> None:
        self.drop_zone.configure(border_color=COLORS["accent_blue"])

    def _on_drop_leave(self) -> None:
        self.drop_zone.configure(border_color=COLORS["text_tertiary"])

    # ── Logic ──────────────────────────────────────────────────────────────

    def _select_file(self) -> None:
        path = filedialog.askopenfilename(
            title="选择负荷计算Excel文件",
            filetypes=[("Excel文件", "*.xlsx *.xls"), ("所有文件", "*.*")],
        )
        if path:
            self.ctx.excel_path = path
            self.file_label.configure(
                text=f"📄 {path}", text_color=COLORS["text_primary"])
            self.ctx.excel_data = None
            self._parse_excel()

    def _parse_excel(self) -> None:
        if not self.ctx.excel_path:
            messagebox.showwarning("提示", "请先选择Excel文件")
            return
        self.parse_btn.configure(text="⏳ 解析中...", state="disabled")
        threading.Thread(target=self._do_parse, daemon=True).start()

    def _do_parse(self) -> None:
        try:
            data = self.ctx.engine.parse_excel(self.ctx.excel_path)
            self.ctx.excel_data = data
            self.after(0, self._display_result, data)
        except Exception as exc:
            self.after(0, lambda: messagebox.showerror("解析失败", str(exc)))
            self.after(0, lambda: self.parse_btn.configure(
                text="🔄 解析并预览", state="normal"))

    def _display_result(self, data: Dict[str, Any]) -> None:
        self.parse_btn.configure(text="🔄 解析并预览", state="normal")
        summary = data.get("summary", {})

        # Hide old results
        self.results_frame.pack_forget()
        self.results_frame.destroy()
        self.results_frame = ctk.CTkFrame(self, fg_color="transparent")
        self.results_frame.pack(fill="both", expand=True, pady=(10, 0))

        # ── Summary card ──
        card = ctk.CTkFrame(self.results_frame, fg_color=COLORS["bg_card"],
                            corner_radius=12)
        card.pack(fill="x", pady=(0, 12))

        stats = [
            ("设备总数",   summary.get("total_devices", 0), "台"),
            ("安装容量",   f'{summary.get("total_equip_power", 0):.1f}', "kW"),
            ("计算负荷",   f'{summary.get("total_sc_k", 0):.1f}', "kVA"),
            ("补偿容量",   f'{summary.get("qc_compensation", 0):.1f}', "kvar"),
            ("补偿前cosφ", f'{summary.get("cos_before", 0):.4f}', ""),
            ("推荐变压器", summary.get("recommended_transformer", "-"), ""),
        ]
        stat_grid = ctk.CTkFrame(card, fg_color="transparent")
        stat_grid.pack(padx=20, pady=15, fill="x")
        for i, (label, value, unit) in enumerate(stats):
            col = i % 3
            row = i // 3
            cell = ctk.CTkFrame(stat_grid, fg_color="transparent")
            cell.grid(row=row, column=col, padx=15, pady=8, sticky="w")
            display = f"{value}{unit}"
            ctk.CTkLabel(cell, text=display, font=get_ctk_font("section"),
                         text_color=COLORS["accent_blue"]).pack(anchor="w")
            ctk.CTkLabel(cell, text=label, font=get_ctk_font("caption"),
                         text_color=COLORS["text_secondary"]).pack(anchor="w")
        # Make grid columns stretch
        for c in range(3):
            stat_grid.grid_columnconfigure(c, weight=1)

        # ── Area detail table ──
        table_card = ctk.CTkFrame(self.results_frame, fg_color=COLORS["bg_card"],
                                  corner_radius=12)
        table_card.pack(fill="both", expand=True)

        ctk.CTkLabel(table_card, text="各区域负荷明细", font=get_ctk_font("section"),
                     text_color=COLORS["text_primary"]).pack(anchor="w", padx=20, pady=(15, 5))

        # Treeview for detail
        tree_frame = ctk.CTkFrame(table_card, fg_color="transparent")
        tree_frame.pack(fill="both", expand=True, padx=15, pady=(0, 15))

        columns = ("area", "devices", "power", "pc", "qc", "sc")
        self.tree = ttk.Treeview(tree_frame, columns=columns, show="headings", height=8)
        self.tree.heading("area", text="区域名称")
        self.tree.heading("devices", text="设备数")
        self.tree.heading("power", text="设备容量(kW)")
        self.tree.heading("pc", text="有功(kW)")
        self.tree.heading("qc", text="无功(kvar)")
        self.tree.heading("sc", text="视在(kVA)")
        self.tree.column("area", width=150)
        self.tree.column("devices", width=70, anchor="center")
        for c in ("power", "pc", "qc", "sc"):
            self.tree.column(c, width=110, anchor="e")

        vsb = ttk.Scrollbar(tree_frame, orient="vertical", command=self.tree.yview)
        self.tree.configure(yscrollcommand=vsb.set)
        self.tree.pack(side="left", fill="both", expand=True)
        vsb.pack(side="right", fill="y")

        # Populate
        for area, info in data.get("area_summaries", {}).items():
            self.tree.insert("", "end", values=(
                area,
                info.get("device_count", 0),
                f'{info.get("equip_power", 0):.1f}',
                f'{info.get("pc", 0):.1f}',
                f'{info.get("qc", 0):.1f}',
                f'{info.get("sc", 0):.1f}',
            ))
        # 合计 row
        self.tree.insert("", "end", values=(
            "合计",
            summary.get("total_devices", 0),
            f'{summary.get("total_equip_power", 0):.1f}',
            f'{summary.get("total_pc", 0):.1f}',
            f'{summary.get("total_qc", 0):.1f}',
            f'{summary.get("total_sc", 0):.1f}',
        ), tags=("summary",))
        self.tree.tag_configure("summary", font=("微软雅黑", 11, "bold"))
        self.after(0, self.tree.yview_moveto, 1.0)


# ═══════════════════════════════════════════════════════════════════════════
# 5. PAGE: GeneratePage
# ═══════════════════════════════════════════════════════════════════════════

class GeneratePage(ctk.CTkFrame):
    """Document generation page with config summary, progress, log, and preview."""

    def __init__(self, master, ctx: EngineContext, get_params_cb,
                 **kwargs) -> None:
        super().__init__(master, fg_color="transparent", **kwargs)
        self.ctx = ctx
        self.get_params_cb = get_params_cb  # callable → Dict[str, Any]
        self._preview_expanded = False
        self._build()

    # ── UI ─────────────────────────────────────────────────────────────────

    def _build(self) -> None:
        # Header
        hdr = ctk.CTkFrame(self, fg_color="transparent")
        hdr.pack(fill="x", pady=(10, 15))
        ctk.CTkLabel(hdr, text="生成设计说明书", font=get_ctk_font("header"),
                     text_color=COLORS["text_primary"]).pack(side="left")
        ctk.CTkButton(
            hdr, text="🔄 刷新配置", font=get_ctk_font("caption"),
            fg_color=COLORS["bg_selected"],
            text_color=COLORS["text_primary"],
            corner_radius=8, command=self._refresh_config,
        ).pack(side="right")

        # Three-column layout
        main = ctk.CTkFrame(self, fg_color="transparent")
        main.pack(fill="both", expand=True)

        # ── Left: config summary ──
        left_card = ctk.CTkFrame(main, fg_color=COLORS["bg_card"],
                                 corner_radius=12, width=260)
        left_card.pack(side="left", fill="both", padx=(0, 8))
        left_card.pack_propagate(False)
        ctk.CTkLabel(left_card, text="当前配置", font=get_ctk_font("section"),
                     text_color=COLORS["text_primary"]).pack(anchor="w", padx=15, pady=(15, 5))
        self.config_text = ctk.CTkTextbox(
            left_card, font=get_ctk_font("mono"), fg_color=COLORS["bg_card"],
            corner_radius=8, wrap="word",
        )
        self.config_text.pack(fill="both", expand=True, padx=10, pady=(0, 10))
        self.config_text.configure(state="disabled")

        # ── Center: actions + log ──
        center = ctk.CTkFrame(main, fg_color="transparent")
        center.pack(side="left", fill="both", expand=True, padx=(8, 8))

        # Action buttons
        btn_row = ctk.CTkFrame(center, fg_color="transparent")
        btn_row.pack(fill="x", pady=(0, 10))
        self.gen_btn = ctk.CTkButton(
            btn_row, text="📝 生成Word文档", font=get_ctk_font("body"),
            fg_color=COLORS["accent_blue"], hover_color="#0066D6",
            corner_radius=12, height=40, command=self._generate,
        )
        self.gen_btn.pack(side="left", fill="x", expand=True, padx=(0, 5))
        ctk.CTkButton(
            btn_row, text="📂 打开输出文件夹", font=get_ctk_font("body"),
            fg_color="transparent", border_width=1,
            border_color=COLORS["accent_blue"],
            text_color=COLORS["accent_blue"],
            corner_radius=8, command=self._open_output,
        ).pack(side="left", padx=(5, 0))

        # Progress bar
        self.progress = ctk.CTkProgressBar(
            center, mode="indeterminate", height=6,
            fg_color=COLORS["bg_primary"],
            progress_color=COLORS["accent_blue"],
            corner_radius=3,
        )
        self.progress.pack(fill="x", pady=(0, 10))

        # Log output
        log_card = ctk.CTkFrame(center, fg_color=COLORS["bg_card"],
                                corner_radius=12)
        log_card.pack(fill="both", expand=True)
        ctk.CTkLabel(log_card, text="生成日志", font=get_ctk_font("caption"),
                     text_color=COLORS["text_secondary"]).pack(anchor="w", padx=15, pady=(10, 3))
        self.log_text = ctk.CTkTextbox(
            log_card, font=get_ctk_font("mono"), fg_color=COLORS["bg_card"],
            corner_radius=8, wrap="word",
        )
        self.log_text.pack(fill="both", expand=True, padx=10, pady=(0, 10))
        self.log_text.configure(state="disabled")

        # Output path label
        self.path_var = ctk.StringVar(value="")
        self.path_label = ctk.CTkLabel(
            center, textvariable=self.path_var,
            font=get_ctk_font("caption"), text_color=COLORS["accent_green"],
        )
        self.path_label.pack(anchor="w", pady=(5, 0))

        # ── Right: steps ──
        right_card = ctk.CTkFrame(main, fg_color=COLORS["bg_card"],
                                  corner_radius=12, width=200)
        right_card.pack(side="right", fill="both", padx=(8, 0))
        right_card.pack_propagate(False)
        ctk.CTkLabel(right_card, text="生成步骤", font=get_ctk_font("section"),
                     text_color=COLORS["text_primary"]).pack(anchor="w", padx=15, pady=(15, 10))
        steps = [
            "1️⃣ 填写项目信息",
            "2️⃣ 导入Excel",
            "3️⃣ 生成文档",
            "4️⃣ 查看输出",
        ]
        for s in steps:
            ctk.CTkLabel(right_card, text=s, font=get_ctk_font("caption"),
                         text_color=COLORS["text_secondary"]).pack(
                anchor="w", padx=15, pady=2)

        # ── Preview section (below center) ──
        self.preview_toggle = ctk.CTkButton(
            center, text="📋 实时预览  ▸", font=get_ctk_font("caption"),
            fg_color="transparent", text_color=COLORS["accent_blue"],
            hover_color=COLORS["bg_hover"], corner_radius=8,
            command=self._toggle_preview,
        )
        self.preview_toggle.pack(anchor="w", pady=(8, 2))

        self.preview_frame = ctk.CTkFrame(center, fg_color=COLORS["bg_card"],
                                          corner_radius=12, height=180)
        # Hidden by default

    # ── Preview ────────────────────────────────────────────────────────────

    def _toggle_preview(self) -> None:
        if self._preview_expanded:
            self.preview_frame.pack_forget()
            self.preview_toggle.configure(text="📋 实时预览  ▸")
            self._preview_expanded = False
        else:
            if not self.preview_frame.winfo_ismapped():
                self._build_preview_content()
                self.preview_frame.pack(fill="both", expand=True, pady=(5, 0))
            self.preview_toggle.configure(text="📋 实时预览  ▾")
            self._preview_expanded = True
            self._refresh_preview()

    def _build_preview_content(self) -> None:
        btn_row = ctk.CTkFrame(self.preview_frame, fg_color="transparent")
        btn_row.pack(fill="x", padx=10, pady=(8, 2))
        ctk.CTkLabel(btn_row, text="文档大纲预览", font=get_ctk_font("caption"),
                     text_color=COLORS["text_secondary"]).pack(side="left")
        ctk.CTkButton(
            btn_row, text="🔄 刷新预览", font=get_ctk_font("caption"),
            fg_color=COLORS["bg_selected"],
            text_color=COLORS["text_primary"],
            corner_radius=6, width=90,
            command=self._refresh_preview,
        ).pack(side="right")

        self.preview_text = ctk.CTkTextbox(
            self.preview_frame, font=get_ctk_font("caption"),
            fg_color=COLORS["bg_card"], corner_radius=8, wrap="word",
        )
        self.preview_text.pack(fill="both", expand=True, padx=10, pady=(0, 8))
        self.preview_text.configure(state="disabled")

    def _refresh_preview(self) -> None:
        if not self._preview_expanded:
            return

        if not _DOCX_PREVIEW_AVAILABLE:
            self._set_preview("预览功能需后端支持 (DocxGenerator 导入失败)")
            return

        if not self.ctx.excel_data:
            self._set_preview("请先在「Excel导入」页面解析负荷计算表")
            return

        params = self.get_params_cb()
        project_type = params.get("project_type", "water_supply")
        design_stage = params.get("design_stage", "初步设计")

        rules_dir = os.path.join(_PROJECT_ROOT, "backend", "data", "rules")
        output_dir = os.path.join(_PROJECT_ROOT, "output")

        def _task() -> None:
            try:
                from app.services.docx_generator import DocxGenerator as DG
                gen = DG(rules_dir=rules_dir, output_dir=output_dir,
                         template="standard")
                blocks = gen.preview(project_type, design_stage,
                                     self.ctx.excel_data, params)
                self.after(0, self._render_preview_blocks, blocks)
            except Exception as exc:
                self.after(0, self._set_preview, f"预览生成失败: {exc}")

        threading.Thread(target=_task, daemon=True).start()

    def _render_preview_blocks(self, blocks: list) -> None:
        lines: List[str] = []
        for blk in blocks:
            kind = blk[0]
            if kind == "cover":
                _, name, subtitle, info_list = blk
                lines.append(f"📕 {name}")
                lines.append(f"   {subtitle}")
                for line in info_list:
                    lines.append(f"   {line}")
                lines.append("")
            elif kind == "pagebreak":
                lines.append("─" * 48 + "  分页  " + "─" * 48)
            elif kind in ("h1", "h2", "h3"):
                prefix = {"h1": "# ", "h2": "## ", "h3": "### "}[kind]
                lines.append(f"{prefix}{blk[1]}")
            elif kind == "p":
                lines.append(f"  {blk[1]}")
            elif kind == "b":
                lines.append(f"  • {blk[1]}")
            elif kind == "table":
                _, _, headers, rows = blk
                lines.append(f"  [表格] {' | '.join(headers)}")
                for row in rows[:5]:
                    lines.append(f"    {' | '.join(str(c) for c in row)}")
                if len(rows) > 5:
                    lines.append(f"    ... (共 {len(rows)} 行)")
                lines.append("")
            elif kind == "sig":
                lines.append("  [签署栏] 编制 / 校核 / 审核 / 审定")
        self._set_preview("\n".join(lines))

    def _set_preview(self, text: str) -> None:
        if hasattr(self, "preview_text"):
            self.preview_text.configure(state="normal")
            self.preview_text.delete("1.0", "end")
            self.preview_text.insert("1.0", text)
            self.preview_text.configure(state="disabled")

    # ── Config Refresh ─────────────────────────────────────────────────────

    def _refresh_config(self) -> None:
        params = self.get_params_cb()
        self.config_text.configure(state="normal")
        self.config_text.delete("1.0", "end")
        text = (
            f'项目名称: {params.get("project_name", "-")}\n'
            f'工程类型: {params.get("project_type", "-")}\n'
            f'设计阶段: {params.get("design_stage", "-")}\n'
            f'电压等级: {params.get("voltage_level", "-")}\n'
            f'负荷等级: {params.get("load_level", "-")}\n'
            f'电源回路: {params.get("power_source", "-")}\n'
            f'备用说明: {params.get("standby_desc", "-")}\n'
            f'变配电间: {params.get("tx_location", "-")}\n'
            f'\n--- Excel数据 ---\n'
            f'文件: {"✅ 已加载" if self.ctx.excel_data else "❌ 未选择"}\n'
        )
        if self.ctx.excel_data:
            s = self.ctx.excel_data.get("summary", {})
            text += (
                f'设备总数: {s.get("total_devices", 0)}\n'
                f'总容量: {s.get("total_equip_power", 0):.1f} kW\n'
                f'计算负荷: {s.get("total_sc_k", 0):.1f} kVA\n'
                f'变压器: {s.get("recommended_transformer", "-")}\n'
            )
        self.config_text.insert("1.0", text)
        self.config_text.configure(state="disabled")

    # ── Generation ─────────────────────────────────────────────────────────

    def _log(self, msg: str) -> None:
        ts = datetime.now().strftime("%H:%M:%S")
        self.log_text.configure(state="normal")
        self.log_text.insert("end", f"[{ts}] {msg}\n")
        self.log_text.see("end")
        self.log_text.configure(state="disabled")

    def _generate(self) -> None:
        if not self.ctx.excel_data:
            messagebox.showwarning("缺少数据",
                                   "请先在「Excel导入」页面解析负荷计算表")
            return

        params = self.get_params_cb()
        self.gen_btn.configure(text="⏳ 生成中...", state="disabled")
        self.progress.start()
        self._log("🚀 开始生成文档...")
        self._log(f'   工程类型: {params.get("project_type", "-")}')
        self._log(f'   设计阶段: {params.get("design_stage", "-")}')
        self._log(f'   项目名称: {params.get("project_name", "-")}')

        def _task() -> None:
            try:
                edata = self.ctx.excel_data
                assert edata is not None, "excel_data must be set"
                output_path = self.ctx.engine.generate(
                    project_type=params.get("project_type", "water_supply"),
                    design_stage=params.get("design_stage", "初步设计"),
                    excel_data=edata,
                    params=params,
                )
                self.ctx.last_output_path = output_path
                self.after(0, self._on_success, output_path)
            except Exception as exc:
                tb = traceback.format_exc()
                self.after(0, self._on_error, str(exc), tb)

        threading.Thread(target=_task, daemon=True).start()

    def _on_success(self, path: str) -> None:
        self.progress.stop()
        self.gen_btn.configure(text="📝 生成Word文档", state="normal")
        self._log("✅ 文档生成成功！")
        self._log(f'   输出路径: {path}')
        self.path_var.set(f"✅ 文档已生成: {path}")
        messagebox.showinfo("生成成功", f"文档已生成：\n{path}")

    def _on_error(self, err: str, tb: str) -> None:
        self.progress.stop()
        self.gen_btn.configure(text="📝 生成Word文档", state="normal")
        self._log(f"❌ 生成失败: {err}")
        messagebox.showerror("生成失败", err)

    @staticmethod
    def _open_output() -> None:
        output_dir = os.path.join(_PROJECT_ROOT, "output")
        os.makedirs(output_dir, exist_ok=True)
        os.startfile(output_dir)


# ═══════════════════════════════════════════════════════════════════════════
# 5b. RICH TEXT EDITOR WIDGET
# ═══════════════════════════════════════════════════════════════════════════

class RichTextFrame(ctk.CTkFrame):
    """A frame containing a formatting toolbar + tk.Text widget for rich text editing."""

    _TOOLBAR_BTNS = [
        ("B", "bold", "加粗"),
        ("I", "italic", "斜体"),
        ("U", "underline", "下划线"),
        ("•", "bullet", "项目符号"),
        ("1.", "numbered", "编号列表"),
        ("⇥", "indent", "增加缩进"),
        ("⇤", "outdent", "减少缩进"),
    ]

    def __init__(self, master, insert_callback: Optional[callable] = None, **kwargs) -> None:
        super().__init__(master, fg_color="transparent", **kwargs)
        self._active_tags: set[str] = set()
        self._bullet_counter: int = 0
        self._numbered_counter: int = 0
        self._insert_callback = insert_callback
        self._build_toolbar()
        self._build_editor()

    # ── Build ────────────────────────────────────────────────────────────

    def _build_toolbar(self) -> None:
        self.toolbar = ctk.CTkFrame(self, fg_color=COLORS["bg_primary"],
                                    corner_radius=6, height=38)
        self.toolbar.pack(fill="x", padx=0, pady=(0, 0))
        self.toolbar.pack_propagate(False)

        self._tb_buttons: Dict[str, ctk.CTkButton] = {}
        for label, action, tooltip in self._TOOLBAR_BTNS:
            btn = ctk.CTkButton(
                self.toolbar, text=label, width=32, height=28,
                corner_radius=4,
                fg_color="transparent",
                hover_color=COLORS["bg_hover"],
                text_color=COLORS["text_primary"],
                font=ctk.CTkFont(family="微软雅黑", size=12,
                                 weight="bold" if label in ("B",) else "normal"),
                command=lambda a=action: self._toolbar_action(a),
            )
            btn.pack(side="left", padx=1, pady=4)
            self._tb_buttons[action] = btn

        # Separator
        sep = ctk.CTkFrame(self.toolbar, width=1, fg_color=COLORS["separator"])
        sep.pack(side="left", fill="y", padx=4, pady=4)

        # Insert Template button
        self.insert_tpl_btn = ctk.CTkButton(
            self.toolbar, text="[ ]", width=32, height=28,
            font=("微软雅黑", 11),
            fg_color="transparent", text_color=COLORS["text_primary"],
            hover_color=COLORS["bg_hover"], corner_radius=4,
            command=self._insert_callback if self._insert_callback else self._insert_template_placeholder,
        )
        self.insert_tpl_btn.pack(side="left", padx=2, pady=4)

    def _build_editor(self) -> None:
        # Thin divider line between toolbar and editor
        divider = ctk.CTkFrame(self, height=1, fg_color=COLORS["separator"])
        divider.pack(fill="x", padx=0, pady=0)

        # The actual text widget
        self.text = tk.Text(
            self,
            font=("Consolas", 11),
            bg=COLORS["bg_card"],
            fg=COLORS["text_primary"],
            insertbackground=COLORS["text_primary"],
            selectbackground=COLORS["accent_blue"],
            selectforeground="white",
            relief="flat",
            bd=0,
            padx=10,
            pady=8,
            wrap="word",
            undo=True,
            maxundo=-1,
            spacing1=1,
            spacing3=1,
        )
        self.text.pack(fill="both", expand=True, padx=1, pady=(0, 1))

        # Configure tags
        self.text.tag_configure("bold", font=("Consolas", 11, "bold"))
        self.text.tag_configure("italic", font=("Consolas", 11, "italic"))
        self.text.tag_configure("underline", font=("Consolas", 11, "underline"))
        self.text.tag_configure("bold_italic", font=("Consolas", 11, "bold italic"))
        self.text.tag_configure("bold_underline", font=("Consolas", 11, "bold underline"))
        self.text.tag_configure("italic_underline", font=("Consolas", 11, "italic underline"))
        self.text.tag_configure("bold_italic_underline", font=("Consolas", 11, "bold italic underline"))
        self.text.tag_configure("bullet", lmargin1=24, lmargin2=40)
        self.text.tag_configure("numbered", lmargin1=24, lmargin2=40)
        self.text.tag_configure("indent1", lmargin1=40, lmargin2=56)
        self.text.tag_configure("indent2", lmargin1=56, lmargin2=72)

        # Bind selection change to update toolbar button states
        self.text.bind("<<Selection>>", self._on_selection_change)

    # ── Public API ───────────────────────────────────────────────────────

    def get_content(self) -> str:
        """Return the full text content."""
        return self.text.get("1.0", "end-1c")

    def set_content(self, text: str) -> None:
        """Replace all content with the given text."""
        self.text.delete("1.0", "end")
        self.text.insert("1.0", text)

    def clear(self) -> None:
        """Clear all content."""
        self.text.delete("1.0", "end")

    def insert_at_cursor(self, text: str) -> None:
        """Insert text at the current cursor position."""
        try:
            # Ensure widget is editable
            if self.text.cget("state") == "disabled":
                self.text.configure(state="normal")
            cursor_pos = self.text.index("insert")
            self.text.insert(cursor_pos, text)
            self.text.focus_set()
        except Exception:
            # Fallback: append at end
            self.text.insert("end", "\n" + text)
        self.text.see("insert")

    def configure_readonly(self, readonly: bool = True) -> None:
        """Set the editor to read-only or editable mode."""
        state = "disabled" if readonly else "normal"
        self.text.configure(state=state)
        # Disable/enable toolbar buttons
        for btn in self._tb_buttons.values():
            btn.configure(state="disabled" if readonly else "normal")
        self.insert_tpl_btn.configure(state="disabled" if readonly else "normal")

    # ── Toolbar actions ──────────────────────────────────────────────────

    def _toolbar_action(self, action: str) -> None:
        if action == "bold":
            self._toggle_tag("bold")
        elif action == "italic":
            self._toggle_tag("italic")
        elif action == "underline":
            self._toggle_tag("underline")
        elif action == "bullet":
            self._insert_bullet()
        elif action == "numbered":
            self._insert_numbered()
        elif action == "indent":
            self._change_indent(1)
        elif action == "outdent":
            self._change_indent(-1)

    def _toggle_tag(self, tag_name: str) -> None:
        """Toggle a formatting tag on the selected text, or set it for future typing."""
        sel_range = self._get_sel_range()
        if sel_range:
            start, end = sel_range
            existing_tags = self.text.tag_names(start)
            if tag_name in existing_tags:
                self.text.tag_remove(tag_name, start, end)
            else:
                self.text.tag_add(tag_name, start, end)
        else:
            # No selection — toggle for future typing at cursor
            cursor = self.text.index("insert")
            existing_tags = self.text.tag_names(cursor)
            if tag_name in existing_tags:
                self.text.tag_remove(tag_name, cursor)
                self._active_tags.discard(tag_name)
            else:
                self.text.tag_add(tag_name, cursor)
                self._active_tags.add(tag_name)
        self._update_toolbar_highlights()

    def _insert_bullet(self) -> None:
        """Insert a bullet list item at the current line."""
        self.text.focus_set()
        cursor = self.text.index("insert")
        line_start = self.text.index(f"{cursor} linestart")
        line_text = self.text.get(line_start, f"{line_start} lineend")

        if line_text.strip():
            # Current line has content — insert bullet before it
            self.text.insert(line_start, "• ")
            self.text.tag_add("bullet", line_start, f"{line_start} lineend")
        else:
            # Empty line — insert bullet and move cursor after it
            self.text.insert(line_start, "• ")
            self.text.tag_add("bullet", line_start, f"{line_start} lineend")
            self.text.mark_set("insert", f"{line_start}+2c")

    def _insert_numbered(self) -> None:
        """Insert a numbered list item at the current line."""
        self.text.focus_set()
        cursor = self.text.index("insert")
        line_start = self.text.index(f"{cursor} linestart")
        line_text = self.text.get(line_start, f"{line_start} lineend")

        # Count existing numbered lines above to determine next number
        num = 1
        line_num = int(cursor.split(".")[0])
        check = line_num - 1
        while check >= 1:
            prev_start = self.text.index(f"{check}.0")
            prev_text = self.text.get(prev_start, f"{prev_start} lineend")
            if prev_text.strip() and prev_text.strip()[0].isdigit():
                num += 1
                check -= 1
            else:
                break

        if line_text.strip():
            self.text.insert(line_start, f"{num}. ")
            self.text.tag_add("numbered", line_start, f"{line_start} lineend")
        else:
            self.text.insert(line_start, f"{num}. ")
            self.text.tag_add("numbered", line_start, f"{line_start} lineend")
            prefix_len = len(f"{num}. ")
            self.text.mark_set("insert", f"{line_start}+{prefix_len}c")

    def _change_indent(self, direction: int) -> None:
        """Increase or decrease indent on selected lines or current line."""
        sel_range = self._get_sel_range()
        if sel_range:
            start_line = int(self.text.index(sel_range[0]).split(".")[0])
            end_line = int(self.text.index(sel_range[1]).split(".")[0])
        else:
            cursor = self.text.index("insert")
            start_line = end_line = int(cursor.split(".")[0])

        for line_num in range(start_line, end_line + 1):
            line_start = f"{line_num}.0"
            line_text = self.text.get(line_start, f"{line_start} lineend")
            if direction > 0:
                self.text.insert(line_start, "    ")
            else:
                if line_text.startswith("    "):
                    self.text.delete(line_start, f"{line_start}+4c")
                elif line_text.startswith("\t"):
                    self.text.delete(line_start, f"{line_start}+1c")

    def _insert_template_placeholder(self) -> None:
        """Placeholder for Batch 7 template insertion."""
        messagebox.showinfo("插入模板", "待实现 — 此功能将在 Batch 7 中完成。")

    # ── Internal helpers ─────────────────────────────────────────────────

    def _get_sel_range(self) -> Optional[Tuple[str, str]]:
        """Return (start, end) of text selection, or None if empty."""
        try:
            sel = self.text.tag_ranges("sel")
            if sel:
                return (str(sel[0]), str(sel[1]))
        except Exception:
            pass
        return None

    def _on_selection_change(self, _event=None) -> None:
        """Update toolbar button highlights when selection changes."""
        self._update_toolbar_highlights()

    def _update_toolbar_highlights(self) -> None:
        """Highlight toolbar buttons whose tag is active at cursor/selection."""
        cursor = self.text.index("insert")
        active = set(self.text.tag_names(cursor))

        for action, btn in self._tb_buttons.items():
            if action in ("bold", "italic", "underline"):
                if action in active:
                    btn.configure(fg_color=COLORS["accent_blue"],
                                  text_color="white")
                else:
                    btn.configure(fg_color="transparent",
                                  text_color=COLORS["text_primary"])


# ═══════════════════════════════════════════════════════════════════════════
# 5c. TEMPLATE INSERT DIALOG
# ═══════════════════════════════════════════════════════════════════════════

class TemplateInsertDialog(ctk.CTkToplevel):
    """Dialog to browse and insert historical templates into the editor."""

    # Project type display mapping
    _PT_LABELS = {
        "water_supply": "💧 给水工程",
        "drainage":     "💧 排水工程",
        "road":         "🛣️ 道路工程",
        "sanitation":   "🗑️ 环卫工程",
        "unclassified": "❓ 未分类",
    }
    _STAGE_LABELS = {
        "可行性研究": "📋 可行性研究",
        "初步设计":   "📋 初步设计",
        "unclassified": "❓ 未分类",
    }
    _CAT_LABELS = {
        "电气":   "⚡ 电气",
        "自控":   "🔧 自控",
        "unclassified": "❓ 未分类",
    }

    def __init__(self, master, importer: HistoricalTemplateImporter,
                 insert_callback: callable, **kwargs) -> None:
        super().__init__(master, **kwargs)
        self.importer = importer
        self._insert_callback = insert_callback
        self._selected_content: str = ""
        self._all_templates: List[dict] = []
        self._filtered_templates: List[dict] = []
        self._selected_card_id: Optional[str] = None

        self.title("插入历史模板")
        self.geometry("750x520")
        self.resizable(True, True)
        self.minsize(600, 400)

        # macOS Sonoma style
        self.configure(fg_color=COLORS["bg_card"])

        self._build()
        self._load_templates()

        # Bind Escape to close
        self.bind("<Escape>", lambda e: self.destroy())

        # Make modal — must happen after the window is mapped; use after_idle
        # so that CTkToplevel has finished its initial geometry work.
        try:
            self.after(50, self._make_modal)
        except Exception:
            pass

    def _make_modal(self) -> None:
        """Acquire grab + lift/focus to enforce modal behavior."""
        try:
            self.grab_set()
        except Exception:
            pass
        try:
            self.lift()
            self.focus_force()
        except Exception:
            pass

    def _build(self) -> None:
        """Build the dialog layout."""
        # ── Top bar: search ──────────────────────────────────────────────
        top_bar = ctk.CTkFrame(self, fg_color=COLORS["bg_primary"], corner_radius=0, height=48)
        top_bar.pack(fill="x", padx=0, pady=0)
        top_bar.pack_propagate(False)

        self._search_var = tk.StringVar()
        search_entry = ctk.CTkEntry(
            top_bar, textvariable=self._search_var,
            placeholder_text="搜索模板标题或内容…",
            font=get_ctk_font("body"), corner_radius=8,
            fg_color=COLORS["bg_card"], border_color=COLORS["separator"],
        )
        search_entry.pack(side="left", fill="x", expand=True, padx=12, pady=8)
        search_entry.bind("<KeyRelease>", self._on_search_changed)

        search_btn = ctk.CTkButton(
            top_bar, text="🔍", width=36, height=32,
            font=("微软雅黑", 14), corner_radius=8,
            fg_color=COLORS["accent_blue"], text_color="white",
            hover_color="#0066DD",
            command=self._on_search_changed,
        )
        search_btn.pack(side="left", padx=(0, 12), pady=8)

        # ── Main content: left filter + right preview ────────────────────
        content_frame = ctk.CTkFrame(self, fg_color="transparent")
        content_frame.pack(fill="both", expand=True, padx=0, pady=0)
        content_frame.grid_columnconfigure(1, weight=1)
        content_frame.grid_rowconfigure(0, weight=1)

        # Left: filter tree (200px)
        left_frame = ctk.CTkFrame(
            content_frame, fg_color=COLORS["bg_primary"],
            corner_radius=0, width=200,
        )
        left_frame.grid(row=0, column=0, sticky="nsew")
        left_frame.grid_propagate(False)

        ctk.CTkLabel(
            left_frame, text="分类筛选", font=get_ctk_font("caption"),
            text_color=COLORS["text_secondary"],
        ).pack(anchor="w", padx=12, pady=(10, 4))

        # Filter tree using tkinter Treeview (native look)
        tree_frame = ctk.CTkFrame(left_frame, fg_color="transparent")
        tree_frame.pack(fill="both", expand=True, padx=8, pady=(0, 8))

        self._filter_tree = ttk.Treeview(
            tree_frame, show="tree", selectmode="browse",
            style="Custom.Treeview",
        )
        self._filter_tree.pack(fill="both", expand=True)

        # Configure tree style
        style = ttk.Style()
        style.configure("Custom.Treeview",
                        background=COLORS["bg_primary"],
                        foreground=COLORS["text_primary"],
                        fieldbackground=COLORS["bg_primary"],
                        font=("微软雅黑", 10),
                        rowheight=24)
        style.configure("Custom.Treeview.Heading",
                        background=COLORS["bg_primary"],
                        foreground=COLORS["text_primary"],
                        font=("微软雅黑", 10, "bold"))

        self._filter_tree.bind("<<TreeviewSelect>>", self._on_filter_changed)

        # Right: template list + preview
        right_frame = ctk.CTkFrame(content_frame, fg_color="transparent")
        right_frame.grid(row=0, column=1, sticky="nsew", padx=(1, 0))
        right_frame.grid_rowconfigure(0, weight=1)
        right_frame.grid_rowconfigure(1, weight=1)
        right_frame.grid_columnconfigure(0, weight=1)

        # Template list (top half)
        list_frame = ctk.CTkFrame(right_frame, fg_color=COLORS["bg_card"], corner_radius=0)
        list_frame.grid(row=0, column=0, sticky="nsew", padx=0, pady=(0, 1))

        ctk.CTkLabel(
            list_frame, text="模板列表", font=get_ctk_font("caption"),
            text_color=COLORS["text_secondary"],
        ).pack(anchor="w", padx=12, pady=(8, 2))

        self._template_list = ctk.CTkScrollableFrame(
            list_frame, fg_color="transparent",
        )
        self._template_list.pack(fill="both", expand=True, padx=8, pady=(0, 8))

        # Content preview (bottom half)
        preview_frame = ctk.CTkFrame(right_frame, fg_color=COLORS["bg_card"], corner_radius=0)
        preview_frame.grid(row=1, column=0, sticky="nsew", padx=0, pady=(1, 0))

        ctk.CTkLabel(
            preview_frame, text="内容预览", font=get_ctk_font("caption"),
            text_color=COLORS["text_secondary"],
        ).pack(anchor="w", padx=12, pady=(8, 2))

        self._preview_text = ctk.CTkTextbox(
            preview_frame, font=("Consolas", 11),
            fg_color=COLORS["bg_primary"],
            text_color=COLORS["text_primary"],
            corner_radius=6, state="disabled",
        )
        self._preview_text.pack(fill="both", expand=True, padx=8, pady=(0, 8))

        # ── Bottom bar: action buttons ───────────────────────────────────
        bottom_bar = ctk.CTkFrame(self, fg_color=COLORS["bg_primary"], corner_radius=0, height=48)
        bottom_bar.pack(fill="x", padx=0, pady=0)
        bottom_bar.pack_propagate(False)

        self._insert_btn = ctk.CTkButton(
            bottom_bar, text="📋 插入", font=get_ctk_font("body"),
            fg_color=COLORS["accent_green"], text_color="white",
            hover_color="#2DB84E", corner_radius=8,
            width=100, height=32, command=self._on_insert,
            state="disabled",
        )
        self._insert_btn.pack(side="right", padx=(0, 12), pady=8)

        cancel_btn = ctk.CTkButton(
            bottom_bar, text="取消", font=get_ctk_font("body"),
            fg_color=COLORS["bg_hover"], text_color=COLORS["text_primary"],
            hover_color=COLORS["bg_selected"], corner_radius=8,
            width=80, height=32, command=self.destroy,
        )
        cancel_btn.pack(side="right", padx=(0, 8), pady=8)

        # Template count label
        self._count_label = ctk.CTkLabel(
            bottom_bar, text="", font=get_ctk_font("caption"),
            text_color=COLORS["text_secondary"],
        )
        self._count_label.pack(side="left", padx=12, pady=8)

    def _load_templates(self, project_type: Optional[str] = None,
                        design_stage: Optional[str] = None,
                        category: Optional[str] = None,
                        search: Optional[str] = None) -> None:
        """Load templates from importer and populate the list."""
        try:
            self._all_templates = self.importer.list_templates()
        except Exception as e:
            self._all_templates = []
            messagebox.showwarning("警告", f"加载模板失败: {e}", parent=self)

        # Apply filters
        self._filtered_templates = self._all_templates.copy()
        if project_type:
            self._filtered_templates = [
                t for t in self._filtered_templates
                if t.get("project_type") == project_type
            ]
        if design_stage:
            self._filtered_templates = [
                t for t in self._filtered_templates
                if t.get("design_stage") == design_stage
            ]
        if category:
            self._filtered_templates = [
                t for t in self._filtered_templates
                if t.get("category") == category
            ]
        if search:
            query = search.lower()
            self._filtered_templates = [
                t for t in self._filtered_templates
                if query in t.get("title", "").lower()
                or query in t.get("content", "").lower()
            ]

        # Build filter tree
        self._build_filter_tree()

        # Populate template list
        self._populate_template_list()

        # Update count
        total = len(self._all_templates)
        filtered = len(self._filtered_templates)
        if filtered == total:
            self._count_label.configure(text=f"共 {total} 个模板")
        else:
            self._count_label.configure(text=f"显示 {filtered}/{total} 个模板")

    def _build_filter_tree(self) -> None:
        """Build the filter tree from loaded templates."""
        # Clear existing
        for item in self._filter_tree.get_children():
            self._filter_tree.delete(item)

        # Count by category
        pt_counts: Dict[str, int] = {}
        ds_counts: Dict[str, int] = {}
        cat_counts: Dict[str, int] = {}

        for t in self._all_templates:
            pt = t.get("project_type", "unclassified")
            ds = t.get("design_stage", "unclassified")
            cat = t.get("category", "unclassified")
            pt_counts[pt] = pt_counts.get(pt, 0) + 1
            ds_counts[ds] = ds_counts.get(ds, 0) + 1
            cat_counts[cat] = cat_counts.get(cat, 0) + 1

        # Add "All" node
        all_id = self._filter_tree.insert(
            "", "end", text=f"📁 全部 ({len(self._all_templates)})",
            values=("all",),
        )

        # Add project type nodes
        for pt_key, count in sorted(pt_counts.items()):
            label = self._PT_LABELS.get(pt_key, pt_key)
            pt_id = self._filter_tree.insert(
                all_id, "end", text=f"{label} ({count})",
                values=("project_type", pt_key),
            )

            # Add design stage sub-nodes
            for ds_key in sorted(ds_counts.keys()):
                ds_count = sum(
                    1 for t in self._all_templates
                    if t.get("project_type") == pt_key and t.get("design_stage") == ds_key
                )
                if ds_count > 0:
                    ds_label = self._STAGE_LABELS.get(ds_key, ds_key)
                    ds_id = self._filter_tree.insert(
                        pt_id, "end", text=f"{ds_label} ({ds_count})",
                        values=("design_stage", pt_key, ds_key),
                    )

                    # Add category sub-nodes
                    for cat_key in sorted(cat_counts.keys()):
                        cat_count = sum(
                            1 for t in self._all_templates
                            if t.get("project_type") == pt_key
                            and t.get("design_stage") == ds_key
                            and t.get("category") == cat_key
                        )
                        if cat_count > 0:
                            cat_label = self._CAT_LABELS.get(cat_key, cat_key)
                            self._filter_tree.insert(
                                ds_id, "end",
                                text=f"{cat_label} ({cat_count})",
                                values=("category", pt_key, ds_key, cat_key),
                            )

        # Expand all
        for item in self._filter_tree.get_children():
            self._filter_tree.item(item, open=True)
            for child in self._filter_tree.get_children(item):
                self._filter_tree.item(child, open=True)

        # Select "All" by default
        self._filter_tree.selection_set(all_id)
        self._filter_tree.focus(all_id)

    def _populate_template_list(self) -> None:
        """Populate the template list with filtered templates."""
        # Clear existing cards
        for widget in self._template_list.winfo_children():
            widget.destroy()

        if not self._filtered_templates:
            empty_label = ctk.CTkLabel(
                self._template_list,
                text="没有找到历史模板。\n请先在历史模板库页面导入文件。",
                font=get_ctk_font("body"),
                text_color=COLORS["text_tertiary"],
                justify="center",
            )
            empty_label.pack(fill="both", expand=True, pady=40)
            return

        for tpl in self._filtered_templates:
            self._create_template_card(tpl)

    def _create_template_card(self, tpl: dict) -> None:
        """Create a clickable card for a template."""
        card = ctk.CTkFrame(
            self._template_list,
            fg_color=COLORS["bg_primary"],
            corner_radius=6,
        )
        card.pack(fill="x", padx=4, pady=2)

        # Store template data on the card
        card._template_data = tpl

        # Title
        title = tpl.get("title", "无标题")
        source = tpl.get("source_file", "")
        content_preview = tpl.get("content", "")[:100]

        title_label = ctk.CTkLabel(
            card, text=title,
            font=get_ctk_font("body"),
            text_color=COLORS["text_primary"],
            anchor="w",
        )
        title_label.pack(fill="x", padx=10, pady=(6, 0))

        # Source file
        if source:
            source_label = ctk.CTkLabel(
                card, text=f"📄 {source}",
                font=get_ctk_font("caption"),
                text_color=COLORS["text_secondary"],
                anchor="w",
            )
            source_label.pack(fill="x", padx=10, pady=(0, 2))

        # Content preview
        if content_preview:
            preview_label = ctk.CTkLabel(
                card, text=content_preview + "…",
                font=get_ctk_font("caption"),
                text_color=COLORS["text_tertiary"],
                anchor="w",
                wraplength=500,
            )
            preview_label.pack(fill="x", padx=10, pady=(0, 6))

        # Bind click events
        for widget in [card, title_label, source_label, preview_label]:
            widget.bind("<Button-1>", lambda e, t=tpl: self._on_template_selected(t))

    def _on_template_selected(self, template: dict) -> None:
        """Show full content preview for selected template."""
        self._selected_content = template.get("content", "")
        self._selected_card_id = template.get("id")

        # Update preview
        self._preview_text.configure(state="normal")
        self._preview_text.delete("1.0", "end")
        self._preview_text.insert("1.0", self._selected_content)
        self._preview_text.configure(state="disabled")

        # Enable insert button
        self._insert_btn.configure(state="normal")

        # Highlight selected card
        for card in self._template_list.winfo_children():
            if hasattr(card, "_template_data"):
                if card._template_data.get("id") == self._selected_card_id:
                    card.configure(fg_color=COLORS["accent_blue"])
                    for child in card.winfo_children():
                        if isinstance(child, ctk.CTkLabel):
                            child.configure(text_color="white")
                else:
                    card.configure(fg_color=COLORS["bg_primary"])
                    for child in card.winfo_children():
                        if isinstance(child, ctk.CTkLabel):
                            if child.cget("text").startswith("📄"):
                                child.configure(text_color=COLORS["text_secondary"])
                            elif child.cget("text").endswith("…"):
                                child.configure(text_color=COLORS["text_tertiary"])
                            else:
                                child.configure(text_color=COLORS["text_primary"])

    def _on_filter_changed(self, _event=None) -> None:
        """Handle filter tree selection change."""
        selection = self._filter_tree.selection()
        if not selection:
            return

        item = selection[0]
        values = self._filter_tree.item(item, "values")

        if not values:
            return

        filter_type = values[0]

        if filter_type == "all":
            self._filtered_templates = self._all_templates.copy()
        elif filter_type == "project_type":
            pt_key = values[1]
            self._filtered_templates = [
                t for t in self._all_templates
                if t.get("project_type") == pt_key
            ]
        elif filter_type == "design_stage":
            pt_key, ds_key = values[1], values[2]
            self._filtered_templates = [
                t for t in self._all_templates
                if t.get("project_type") == pt_key
                and t.get("design_stage") == ds_key
            ]
        elif filter_type == "category":
            pt_key, ds_key, cat_key = values[1], values[2], values[3]
            self._filtered_templates = [
                t for t in self._all_templates
                if t.get("project_type") == pt_key
                and t.get("design_stage") == ds_key
                and t.get("category") == cat_key
            ]

        # Apply search filter if active
        search_query = self._search_var.get().strip()
        if search_query:
            query = search_query.lower()
            self._filtered_templates = [
                t for t in self._filtered_templates
                if query in t.get("title", "").lower()
                or query in t.get("content", "").lower()
            ]

        self._populate_template_list()

        # Update count
        total = len(self._all_templates)
        filtered = len(self._filtered_templates)
        if filtered == total:
            self._count_label.configure(text=f"共 {total} 个模板")
        else:
            self._count_label.configure(text=f"显示 {filtered}/{total} 个模板")

        # Clear preview
        self._selected_content = ""
        self._selected_card_id = None
        self._preview_text.configure(state="normal")
        self._preview_text.delete("1.0", "end")
        self._preview_text.configure(state="disabled")
        self._insert_btn.configure(state="disabled")

    def _on_search_changed(self, _event=None) -> None:
        """Handle search input change."""
        # Re-apply current filter with search
        self._on_filter_changed()

    def _on_insert(self) -> None:
        """Call the insert callback with selected content and close."""
        if self._selected_content:
            self._insert_callback(self._selected_content)
        self.destroy()


# ═══════════════════════════════════════════════════════════════════════════
# 6. PAGE: RulesPage
# ═══════════════════════════════════════════════════════════════════════════

class RulesPage(ctk.CTkFrame):
    """Knowledge base viewer / editor — 5-level hierarchical knowledge tree with detail panel."""

    # Level constants
    L1_PROJECT = 1
    L2_STAGE = 2
    L3_CATEGORY = 3
    L4_ITEM = 4
    L5_SUBMODULE = 5

    def __init__(self, master, ctx: EngineContext, **kwargs) -> None:
        super().__init__(master, fg_color="transparent", **kwargs)
        self.ctx = ctx
        self.edit_mode = False
        self._sel_iid: Optional[str] = None
        self._node_meta: Dict[str, Dict[str, Any]] = {}
        self._l5_original_content: str = ""
        self._variant_search_results: List[dict] = []
        self._editing_iid: Optional[str] = None
        self._edit_widget: Optional[tk.Entry] = None
        self._original_title: str = ""
        # Drag-drop state
        self._drag_iid: Optional[str] = None
        self._drag_started: bool = False
        self._drag_x: int = 0
        self._drag_y: int = 0
        self._drag_threshold: int = 5  # pixels before drag is recognized
        self._build()

    # ── UI ─────────────────────────────────────────────────────────────────

    def _build(self) -> None:
        # Configure treeview style
        style = ttk.Style()
        style.configure(
            "KnowledgeTree.Treeview",
            rowheight=28,
            font=("微软雅黑", 11),
        )
        style.configure(
            "KnowledgeTree.Treeview.Heading",
            font=("微软雅黑", 11, "bold"),
        )

        # Header
        hdr = ctk.CTkFrame(self, fg_color="transparent")
        hdr.pack(fill="x", pady=(10, 15))
        ctk.CTkLabel(hdr, text="规范深度要求 - 知识库", font=get_ctk_font("header"),
                     text_color=COLORS["text_primary"]).pack(side="left")

        # Edit mode toggle
        self.edit_btn = ctk.CTkButton(
            hdr, text="✏️ 编辑模式", font=get_ctk_font("caption"),
            fg_color=COLORS["bg_primary"], text_color=COLORS["text_primary"],
            hover_color="#0066D6", corner_radius=8,
            width=100, command=self._toggle_edit_mode,
        )
        self.edit_btn.pack(side="right", padx=(6, 0))

        self.filter_var = ctk.StringVar(value="全部")
        ctk.CTkOptionMenu(
            hdr, values=["全部"] + [lbl for _, lbl in PROJECT_TYPES],
            variable=self.filter_var, font=get_ctk_font("body"),
            corner_radius=8, fg_color=COLORS["bg_primary"],
            button_color=COLORS["accent_blue"],
            command=lambda _: self.refresh(),
        ).pack(side="right")
        ctk.CTkLabel(hdr, text="筛选工程类型：", font=get_ctk_font("caption"),
                     text_color=COLORS["text_secondary"]).pack(side="right", padx=(0, 6))

        # Numbering config section
        num_frame = ctk.CTkFrame(self, fg_color=COLORS["bg_primary"], corner_radius=8)
        num_frame.pack(fill="x", pady=(0, 10))
        ctk.CTkLabel(num_frame, text="编号设置", font=get_ctk_font("caption"),
                     text_color=COLORS["text_secondary"]).pack(side="left", padx=(12, 8), pady=6)
        self.show_numbering_var = ctk.BooleanVar(
            value=self.ctx.engine.get_numbering_config().get("show_in_tree", True))
        self.show_numbering_check = ctk.CTkCheckBox(
            num_frame, text="显示编号", variable=self.show_numbering_var,
            onvalue=True, offvalue=False,
            font=get_ctk_font("caption"), command=self._on_numbering_toggle)
        self.show_numbering_check.pack(side="left", padx=(0, 12), pady=6)

        # Split layout
        main = ctk.CTkFrame(self, fg_color="transparent")
        main.pack(fill="both", expand=True)

        # Left: knowledge tree
        left_card = ctk.CTkFrame(main, fg_color=COLORS["bg_card"],
                                 corner_radius=12, width=420)
        left_card.pack(side="left", fill="both", padx=(0, 8))
        left_card.pack_propagate(False)

        tree_frame = ctk.CTkFrame(left_card, fg_color="transparent")
        tree_frame.pack(fill="both", expand=True, padx=10, pady=10)

        self.tree = ttk.Treeview(
            tree_frame, columns=(), show="tree",
            height=20, style="KnowledgeTree.Treeview",
        )
        self.tree.heading("#0", text="规范深度要求 - 知识库", anchor="w")
        self.tree.column("#0", width=380)
        self.tree.bind("<<TreeviewSelect>>", self._on_select)
        self.tree.bind("<Double-1>", self._on_double_click)
        # Right-click: on Windows use <Button-3>, on macOS use <Button-2>
        self.tree.bind("<Button-3>", self._on_right_click)
        if sys.platform == "darwin":
            self.tree.bind("<Button-2>", self._on_right_click)
            self.tree.bind("<Control-Button-1>", self._on_right_click)
        # Drag-drop reordering for Level 4 items
        self.tree.bind("<ButtonPress-1>", self._on_tree_drag_start)
        self.tree.bind("<B1-Motion>", self._on_tree_drag_motion)
        self.tree.bind("<ButtonRelease-1>", self._on_tree_drag_release)

        vsb = ttk.Scrollbar(tree_frame, orient="vertical",
                            command=self.tree.yview)
        self.tree.configure(yscrollcommand=vsb.set)
        self.tree.pack(side="left", fill="both", expand=True)
        vsb.pack(side="right", fill="y")

        # Right: detail panel
        right_card = ctk.CTkFrame(main, fg_color=COLORS["bg_card"],
                                  corner_radius=12)
        right_card.pack(side="right", fill="both", expand=True, padx=(8, 0))

        # ── Browse mode widgets ──
        self.browse_frame = ctk.CTkFrame(right_card, fg_color="transparent")
        self.browse_frame.pack(fill="both", expand=True)

        # Breadcrumb label
        self.breadcrumb_label = ctk.CTkLabel(
            self.browse_frame, text="", font=get_ctk_font("caption"),
            text_color=COLORS["text_secondary"],
        )
        self.breadcrumb_label.pack(anchor="w", padx=20, pady=(12, 0))

        self.detail_title = ctk.CTkLabel(
            self.browse_frame, text="选择一个节点查看详情", font=get_ctk_font("body"),
            text_color=COLORS["text_primary"],
        )
        self.detail_title.pack(anchor="w", padx=20, pady=(5, 5))

        self.detail_text = ctk.CTkTextbox(
            self.browse_frame, font=get_ctk_font("body"),
            fg_color=COLORS["bg_card"], corner_radius=8,
            wrap="word",
        )
        self.detail_text.pack(fill="both", expand=True, padx=15, pady=(5, 15))
        self.detail_text.configure(state="disabled")

        # ── Edit mode widgets ──
        self.edit_frame = ctk.CTkFrame(right_card, fg_color="transparent")

        # CTkTabview with 3 tabs: 基本信息 / 模板内容 / 子模块
        self.edit_tabview = ctk.CTkTabview(
            self.edit_frame,
            fg_color=COLORS["bg_card"],
            segmented_button_fg_color=COLORS["bg_primary"],
            segmented_button_selected_color=COLORS["accent_blue"],
            segmented_button_unselected_color=COLORS["bg_primary"],
        )
        self.edit_tabview.pack(fill="both", expand=True, padx=10, pady=10)

        # ── Tab 1: 基本信息 ──
        self.tab_basic = self.edit_tabview.add("基本信息")

        # -- Level 4 edit sub-frame (item editing) --
        self.l4_edit_frame = ctk.CTkFrame(self.tab_basic, fg_color="transparent")

        ctk.CTkLabel(self.l4_edit_frame, text="条目名称", font=get_ctk_font("caption"),
                     text_color=COLORS["text_secondary"]).pack(anchor="w", padx=15, pady=(15, 2))
        self.edit_title = ctk.CTkEntry(
            self.l4_edit_frame, font=get_ctk_font("body"),
            fg_color=COLORS["bg_primary"], corner_radius=8,
        )
        self.edit_title.pack(fill="x", padx=15, pady=(0, 8))

        ctk.CTkLabel(self.l4_edit_frame, text="要求描述", font=get_ctk_font("caption"),
                     text_color=COLORS["text_secondary"]).pack(anchor="w", padx=15, pady=(8, 2))
        self.edit_requirement = ctk.CTkTextbox(
            self.l4_edit_frame, font=get_ctk_font("body"),
            fg_color=COLORS["bg_primary"], corner_radius=8,
            height=120, wrap="word",
        )
        self.edit_requirement.pack(fill="x", padx=15, pady=(0, 8))

        sw_frame = ctk.CTkFrame(self.l4_edit_frame, fg_color="transparent")
        sw_frame.pack(fill="x", padx=15, pady=8)

        self.edit_calc_var = ctk.BooleanVar(value=False)
        ctk.CTkSwitch(sw_frame, text="含计算", variable=self.edit_calc_var,
                      font=get_ctk_font("caption")).pack(side="left", padx=(0, 15))
        self.edit_table_var = ctk.BooleanVar(value=False)
        ctk.CTkSwitch(sw_frame, text="需要表格", variable=self.edit_table_var,
                      font=get_ctk_font("caption")).pack(side="left", padx=(0, 15))
        self.edit_excel_var = ctk.BooleanVar(value=False)
        ctk.CTkSwitch(sw_frame, text="来自Excel", variable=self.edit_excel_var,
                      font=get_ctk_font("caption")).pack(side="left")

        self.l4_save_btn = ctk.CTkButton(
            self.l4_edit_frame, text="💾 保存修改", font=get_ctk_font("body"),
            fg_color=COLORS["accent_blue"], text_color="white",
            corner_radius=8, height=36, command=self._save_changes,
        )
        self.l4_save_btn.pack(fill="x", padx=15, pady=(15, 15))

        # -- Summary sub-frame (L1-L3 / L5 read-only in 基本信息 tab) --
        self.summary_edit_frame = ctk.CTkFrame(self.tab_basic, fg_color="transparent")
        self.summary_edit_title = ctk.CTkLabel(
            self.summary_edit_frame, text="选择一个条目或子模块进行编辑",
            font=get_ctk_font("body"), text_color=COLORS["text_secondary"],
        )
        self.summary_edit_title.pack(anchor="w", padx=15, pady=(15, 5))
        self.summary_edit_text = ctk.CTkTextbox(
            self.summary_edit_frame, font=get_ctk_font("body"),
            fg_color=COLORS["bg_primary"], corner_radius=8, wrap="word",
        )
        self.summary_edit_text.pack(fill="both", expand=True, padx=15, pady=(5, 15))
        self.summary_edit_text.configure(state="disabled")

        # ── Tab 2: 模板内容 ──
        self.tab_template = self.edit_tabview.add("模板内容")

        # -- Level 5 edit sub-frame (submodule rich-text editing) --
        self.l5_edit_frame = ctk.CTkFrame(self.tab_template, fg_color="transparent")

        self.l5_name_label = ctk.CTkLabel(
            self.l5_edit_frame, text="子模块", font=get_ctk_font("section"),
            text_color=COLORS["text_primary"],
        )
        self.l5_name_label.pack(anchor="w", padx=15, pady=(15, 5))

        ctk.CTkLabel(self.l5_edit_frame, text="模板内容", font=get_ctk_font("caption"),
                     text_color=COLORS["text_secondary"]).pack(anchor="w", padx=15, pady=(8, 2))

        # Rich text editor with toolbar (replaces plain CTkTextbox)
        # Wire the "插入模板" toolbar button to open the historical-template dialog.
        self.l5_editor = RichTextFrame(
            self.l5_edit_frame,
            insert_callback=self._open_template_dialog,
        )
        self.l5_editor.pack(fill="both", expand=True, padx=15, pady=(0, 8))

        # ── 变体模板选择 ──
        self.variant_info_frame = ctk.CTkFrame(self.l5_edit_frame, fg_color="transparent")
        self.variant_info_frame.pack(fill="x", padx=15, pady=(4, 2))

        self.variant_status_label = ctk.CTkLabel(
            self.variant_info_frame, text="当前模板：默认",
            font=get_ctk_font("caption"), text_color=COLORS["text_primary"],
        )
        self.variant_status_label.pack(side="left", padx=(0, 10))

        self.select_variant_btn = ctk.CTkButton(
            self.variant_info_frame, text="选择模板", font=get_ctk_font("caption"),
            fg_color=COLORS["accent_blue"], text_color="white",
            corner_radius=6, width=70, height=28, command=self._open_variant_selector,
        )
        self.select_variant_btn.pack(side="left", padx=(0, 4))

        self.clear_variant_btn = ctk.CTkButton(
            self.variant_info_frame, text="清除选择", font=get_ctk_font("caption"),
            fg_color=COLORS["accent_red"], text_color="white",
            corner_radius=6, width=70, height=28, command=self._on_clear_variant,
        )
        self.clear_variant_btn.pack(side="left")

        # Variant metadata label
        self.variant_meta_label = ctk.CTkLabel(
            self.l5_edit_frame, text="", font=get_ctk_font("caption"),
            text_color=COLORS["text_tertiary"],
        )
        self.variant_meta_label.pack(anchor="w", padx=15, pady=(0, 4))

        # Save buttons row
        l5_btn_row = ctk.CTkFrame(self.l5_edit_frame, fg_color="transparent")
        l5_btn_row.pack(fill="x", padx=15, pady=(8, 15))

        self.l5_save_btn = ctk.CTkButton(
            l5_btn_row, text="💾 保存为模板", font=get_ctk_font("body"),
            fg_color=COLORS["accent_green"], text_color="white",
            corner_radius=8, height=36, command=self._save_submodule,
        )
        self.l5_save_btn.pack(side="left", expand=True, fill="x", padx=(0, 4))

        self.l5_save_as_btn = ctk.CTkButton(
            l5_btn_row, text="💾 另存为", font=get_ctk_font("body"),
            fg_color=COLORS["accent_blue"], text_color="white",
            corner_radius=8, height=36, command=self._save_submodule_as,
        )
        self.l5_save_as_btn.pack(side="left", expand=True, fill="x", padx=(4, 0))

        # -- Read-only info frame for L1-L4 in 模板内容 tab --
        self.template_info_frame = ctk.CTkFrame(self.tab_template, fg_color="transparent")
        self.template_info_label = ctk.CTkLabel(
            self.template_info_frame, text="请选择一个条目或子模块以编辑模板内容",
            font=get_ctk_font("body"), text_color=COLORS["text_secondary"],
            wraplength=360,
        )
        self.template_info_label.pack(anchor="w", padx=15, pady=(15, 5))
        self.template_info_list = ctk.CTkTextbox(
            self.template_info_frame, font=get_ctk_font("body"),
            fg_color=COLORS["bg_primary"], corner_radius=8, wrap="word",
        )
        self.template_info_list.pack(fill="both", expand=True, padx=15, pady=(5, 15))
        self.template_info_list.configure(state="disabled")

        # ── Tab 3: 子模块 ──
        self.tab_submodule = self.edit_tabview.add("子模块")

        self.submodule_frame = ctk.CTkFrame(self.tab_submodule, fg_color="transparent")
        self.submodule_title = ctk.CTkLabel(
            self.submodule_frame, text="子模块管理", font=get_ctk_font("section"),
            text_color=COLORS["text_primary"],
        )
        self.submodule_title.pack(anchor="w", padx=15, pady=(15, 5))
        self.submodule_info_text = ctk.CTkTextbox(
            self.submodule_frame, font=get_ctk_font("body"),
            fg_color=COLORS["bg_primary"], corner_radius=8, wrap="word",
        )
        self.submodule_info_text.pack(fill="both", expand=True, padx=15, pady=(5, 8))
        self.submodule_info_text.configure(state="disabled")

        # Submodule add/delete buttons for L4
        self.submodule_btn_row = ctk.CTkFrame(self.submodule_frame, fg_color="transparent")
        self.submodule_btn_row.pack(fill="x", padx=15, pady=(0, 15))
        self.submodule_add_btn = ctk.CTkButton(
            self.submodule_btn_row, text="➕ 添加子模块", font=get_ctk_font("body"),
            fg_color=COLORS["accent_green"], text_color="white",
            corner_radius=8, height=36, command=self._add_submodule,
        )
        self.submodule_add_btn.pack(side="left", expand=True, fill="x", padx=(0, 4))
        self.submodule_del_btn = ctk.CTkButton(
            self.submodule_btn_row, text="🗑 删除子模块", font=get_ctk_font("body"),
            fg_color=COLORS["accent_red"], text_color="white",
            corner_radius=8, height=36, command=self._delete_submodule,
        )
        self.submodule_del_btn.pack(side="left", expand=True, fill="x", padx=(4, 0))

        # Hide all edit sub-frames initially
        self.l4_edit_frame.pack_forget()
        self.l5_edit_frame.pack_forget()
        self.summary_edit_frame.pack_forget()
        self.template_info_frame.pack_forget()
        self.submodule_frame.pack_forget()
        self.edit_frame.pack_forget()

    # ── Edit mode toggle ──

    def _toggle_edit_mode(self) -> None:
        # Check for unsaved changes before leaving edit mode
        if self.edit_mode and self._is_l5_editing():
            current_content = self.l5_editor.get_content()
            if current_content != self._l5_original_content:
                answer = messagebox.askyesno(
                    "未保存的修改",
                    "当前子模块内容已修改，是否保存？",
                )
                if answer:
                    self._save_submodule()

        self.edit_mode = not self.edit_mode
        if self.edit_mode:
            self.edit_btn.configure(text="🔍 浏览模式")
            self.browse_frame.pack_forget()
            self.edit_frame.pack(fill="both", expand=True)
        else:
            self.edit_btn.configure(text="✏️ 编辑模式")
            self.edit_frame.pack_forget()
            self.browse_frame.pack(fill="both", expand=True)
        self._show_selected()

    # ── Toolbar actions ──────────────────────────────────────────────────

    def _open_template_dialog(self) -> None:
        """Open the TemplateInsertDialog modal for the L5 editor toolbar.

        Triggered by the "插入模板" button on the RichTextFrame toolbar. Loads
        the HistoricalTemplateImporter, opens a TemplateInsertDialog, and wires
        its insert callback to ``self.l5_editor.insert_at_cursor`` so the chosen
        template body lands at the current caret position.
        """
        # Build importer on demand — same pattern used by HistoricalTemplatePage.
        try:
            importer = HistoricalTemplateImporter()
        except Exception as e:
            messagebox.showerror("错误", f"无法加载历史模板库: {e}")
            return

        def insert_content(content: str) -> None:
            """Insert content at cursor position in the L5 editor."""
            try:
                self.l5_editor.insert_at_cursor(content)
            except Exception as e:
                messagebox.showerror("错误", f"插入失败: {e}")

        # Anchor the dialog to the top-level window so it floats above the main
        # app even though we're inside a nested page frame.
        toplevel = self.winfo_toplevel()
        dialog = TemplateInsertDialog(
            toplevel,
            importer=importer,
            insert_callback=insert_content,
        )
        # TemplateInsertDialog schedules grab_set() internally via after(50).

    # ── Selection display ──

    def _show_selected(self) -> None:
        """Show selected node in current mode (browse or edit)."""
        if not self._sel_iid:
            return
        meta = self._node_meta.get(self._sel_iid)
        if not meta:
            return

        level = meta["level"]

        if self.edit_mode:
            self.browse_frame.pack_forget()
            self.edit_frame.pack(fill="both", expand=True)

            # Populate all 3 tabs with level-appropriate content
            self._populate_tab_basic(meta)
            self._populate_tab_template(meta)
            self._populate_tab_submodule(meta)

            # Switch to the default tab for each level
            if level == self.L4_ITEM:
                self.edit_tabview.set("基本信息")
            elif level == self.L5_SUBMODULE:
                self.edit_tabview.set("模板内容")
            else:
                self.edit_tabview.set("基本信息")
        else:
            self.edit_frame.pack_forget()
            self.browse_frame.pack(fill="both", expand=True)
            self._populate_browse(meta)

    def _build_breadcrumb(self, meta: Dict[str, Any]) -> str:
        """Build breadcrumb string from metadata."""
        parts = []
        level = meta["level"]
        if level >= 1:
            parts.append(meta.get("label", meta.get("code", "")))
        if level >= 2:
            parts.append(meta.get("stage", ""))
        if level >= 3:
            parts.append(meta.get("cat_title", meta.get("cat_key", "")))
        if level >= 4:
            item = meta.get("item", {})
            parts.append(item.get("title", ""))
        if level >= 5:
            sub = meta.get("sub_module", {})
            parts.append(sub.get("name", ""))
        return "  ›  ".join(parts)

    def _populate_browse(self, meta: Dict[str, Any]) -> None:
        """Fill browse panel based on selected level."""
        level = meta["level"]
        breadcrumb = self._build_breadcrumb(meta)
        self.breadcrumb_label.configure(text=breadcrumb)

        if level == self.L1_PROJECT:
            self._browse_level1(meta)
        elif level == self.L2_STAGE:
            self._browse_level2(meta)
        elif level == self.L3_CATEGORY:
            self._browse_level3(meta)
        elif level == self.L4_ITEM:
            self._browse_level4(meta)
        elif level == self.L5_SUBMODULE:
            self._browse_level5(meta)

    def _browse_level1(self, meta: Dict[str, Any]) -> None:
        code = meta["code"]
        rule = self.ctx.engine.load_rule(code)
        if not rule:
            return
        self.detail_title.configure(text=f'{meta.get("label", "")} — 工程概览')
        stages = list(rule.get("design_stages", {}).keys())
        reg_ref = rule.get("regulation_ref", "")
        total_items = 0
        for stage_data in rule.get("design_stages", {}).values():
            categories = stage_data.get("sections", stage_data) if isinstance(stage_data, dict) else {}
            for _, cat_data in self._iter_cats_with_path(categories):
                total_items += len(cat_data.get("items", []))
        detail = (
            f'🏗️ 工程类型：{meta.get("label", "")}\n'
            f'📖 规范依据：{reg_ref}\n'
            f'📋 设计阶段：{", ".join(stages)}\n'
            f'📊 条目总数：{total_items}\n'
        )
        self._set_detail_text(detail)

    def _browse_level2(self, meta: Dict[str, Any]) -> None:
        code = meta["code"]
        stage = meta["stage"]
        rule = self.ctx.engine.load_rule(code)
        if not rule:
            return
        stage_data = rule.get("design_stages", {}).get(stage, {})
        categories = stage_data.get("sections", stage_data) if isinstance(stage_data, dict) else {}
        cat_count = 0
        item_count = 0
        cat_names = []
        for _, cat_data, _ in self._iter_cats_with_path(categories):
            cat_count += 1
            items = cat_data.get("items", [])
            item_count += len(items)
            cat_names.append(cat_data.get("title", ""))
        self.detail_title.configure(text=f'{meta.get("label", "")} — {stage}')
        detail = (
            f'📋 设计阶段：{stage}\n'
            f'📑 章节数量：{cat_count}\n'
            f'📊 条目总数：{item_count}\n'
            f'📂 包含章节：{", ".join(cat_names)}\n'
        )
        self._set_detail_text(detail)

    def _browse_level3(self, meta: Dict[str, Any]) -> None:
        code = meta["code"]
        stage = meta["stage"]
        cat_path = meta["cat_path"]
        rule = self.ctx.engine.load_rule(code)
        if not rule:
            return
        stage_data = rule.get("design_stages", {}).get(stage, {})
        categories = stage_data.get("sections", stage_data) if isinstance(stage_data, dict) else {}
        cat_data = self._navigate_categories(categories, cat_path)
        if not cat_data:
            return
        items = cat_data.get("items", [])
        item_titles = [it.get("title", "") for it in items]
        section_id = cat_data.get("section_id", "")
        self.detail_title.configure(text=meta.get("cat_title", ""))
        detail = (
            f'📖 章节编号：{section_id}\n'
            f'📑 章节标题：{cat_data.get("title", "")}\n'
            f'📊 条目数量：{len(items)}\n'
            f'📝 条目列表：\n'
        )
        for i, title in enumerate(item_titles, 1):
            detail += f'   {i}. {title}\n'
        self._set_detail_text(detail)

    def _browse_level4(self, meta: Dict[str, Any]) -> None:
        item = meta["item"]
        cat_title = meta.get("cat_title", "")
        self.detail_title.configure(
            text=f'{cat_title} → {item.get("title", "")}')
        requirement = item.get("requirement", "无具体要求")
        calc = "是" if item.get("has_calculation") else "否"
        table = "需要" if item.get("table_required") else "不需要"
        calc_from = "来自Excel" if item.get("calc_from_excel") else "直接生成"
        section_id = meta.get("section_id", "")
        sub_modules = item.get("sub_modules", [])
        sub_names = [sm.get("name", "") for sm in sub_modules]
        detail = (
            f'📋 要求：{requirement}\n\n'
            f'📌 含计算：{calc}\n'
            f'📊 需要表格：{table}\n'
            f'🔄 计算来源：{calc_from}\n'
            f'📖 章节编号：{section_id}\n'
        )
        if sub_names:
            detail += f'\n📦 子模块（{len(sub_names)}个）：\n'
            for i, name in enumerate(sub_names, 1):
                detail += f'   {i}. {name}\n'
        self._set_detail_text(detail)

    def _browse_level5(self, meta: Dict[str, Any]) -> None:
        sub = meta["sub_module"]
        item = meta["item"]
        cat_title = meta.get("cat_title", "")
        self.detail_title.configure(
            text=f'{cat_title} → {item.get("title", "")} → {sub.get("name", "")}')
        template = sub.get("template_content", "")
        if not template:
            template = "（暂无模板内容，可在编辑模式下添加）"
        detail = (
            f'📝 子模块名称：{sub.get("name", "")}\n'
            f'📄 所属条目：{item.get("title", "")}\n\n'
            f'📋 模板内容：\n{template}\n'
        )

        # Show saved custom templates
        try:
            saved_templates = self.ctx.engine.list_custom_templates(
                project_type=meta["code"],
                design_stage=meta["stage"],
                category=meta.get("cat_key", ""),
                item_title=item.get("title", ""),
                sub_module_name=sub.get("name", ""),
            )
        except Exception:
            saved_templates = []

        if saved_templates:
            detail += f'\n{"─" * 40}\n'
            detail += f'📦 已保存的自定义模板（{len(saved_templates)}个）：\n'
            for i, tpl in enumerate(saved_templates, 1):
                tpl_name = tpl.get("template_name", "未命名")
                updated = tpl.get("updated_at", "")
                if updated and "T" in updated:
                    updated = updated.split("T")[0]
                detail += f'   {i}. {tpl_name}'
                if updated:
                    detail += f'  （更新于 {updated}）'
                detail += '\n'
        else:
            detail += f'\n{"─" * 40}\n'
            detail += '📦 已保存的自定义模板：（无）\n'

        self._set_detail_text(detail)

    def _set_detail_text(self, text: str) -> None:
        self.detail_text.configure(state="normal")
        self.detail_text.delete("1.0", "end")
        self.detail_text.insert("1.0", text)
        self.detail_text.configure(state="disabled")

    # ── Edit mode population ──

    def _populate_l4_edit(self, meta: Dict[str, Any]) -> None:
        item = meta["item"]
        self.edit_title.delete(0, "end")
        self.edit_title.insert(0, item.get("title", ""))
        self.edit_requirement.delete("1.0", "end")
        self.edit_requirement.insert("1.0", item.get("requirement", ""))
        self.edit_calc_var.set(item.get("has_calculation", False))
        self.edit_table_var.set(item.get("table_required", False))
        self.edit_excel_var.set(item.get("calc_from_excel", False))

    def _populate_l5_edit(self, meta: Dict[str, Any]) -> None:
        sub = meta["sub_module"]
        item = meta["item"]
        self.l5_name_label.configure(
            text=f'📝 {sub.get("name", "")}  （所属：{item.get("title", "")}）')
        content = sub.get("template_content", "")
        self.l5_editor.set_content(content)
        self.l5_editor.configure_readonly(False)
        # Store original content for unsaved changes detection
        self._l5_original_content = content
        # Populate custom template combo
        self._update_variant_info(meta)

    def _is_l5_editing(self) -> bool:
        """Return True if currently showing Level 5 edit frame."""
        if not self._sel_iid:
            return False
        meta = self._node_meta.get(self._sel_iid)
        return meta is not None and meta.get("level") == self.L5_SUBMODULE

    def _populate_summary_edit(self, meta: Dict[str, Any]) -> None:
        breadcrumb = self._build_breadcrumb(meta)
        self.summary_edit_title.configure(text=f"📌 {breadcrumb}")
        level = meta["level"]
        info = ""
        if level == self.L1_PROJECT:
            rule = self.ctx.engine.load_rule(meta["code"])
            if rule:
                info = f'规范依据：{rule.get("regulation_ref", "")}\n请选择具体的章节条目或子模块进行编辑。'
        elif level == self.L2_STAGE:
            info = f'设计阶段：{meta.get("stage", "")}\n请选择具体的章节条目或子模块进行编辑。'
        elif level == self.L3_CATEGORY:
            info = f'章节：{meta.get("cat_title", "")}\n请选择具体的条目或子模块进行编辑。'
        self.summary_edit_text.configure(state="normal")
        self.summary_edit_text.delete("1.0", "end")
        self.summary_edit_text.insert("1.0", info)
        self.summary_edit_text.configure(state="disabled")

    # ── Tab-populate helpers for the 3-tab CTkTabview ─────────────────────

    def _populate_tab_basic(self, meta: Dict[str, Any]) -> None:
        """Populate 基本信息 tab based on selected node level."""
        level = meta["level"]
        # Hide all sub-frames in this tab
        self.l4_edit_frame.pack_forget()
        self.summary_edit_frame.pack_forget()

        if level == self.L4_ITEM:
            self._populate_l4_edit(meta)
            self.l4_edit_frame.pack(fill="both", expand=True)
        elif level == self.L5_SUBMODULE:
            sub = meta.get("sub_module", {})
            item = meta.get("item", {})
            self.summary_edit_title.configure(
                text=f'📝 子模块: {sub.get("name", "")}')
            info = (
                f'所属条目: {item.get("title", "")}\n'
                f'所属章节: {meta.get("cat_title", "")}\n'
                f'设计阶段: {meta.get("stage", "")}\n'
                f'工程类型: {meta.get("label", meta.get("code", ""))}'
            )
            self.summary_edit_text.configure(state="normal")
            self.summary_edit_text.delete("1.0", "end")
            self.summary_edit_text.insert("1.0", info)
            self.summary_edit_text.configure(state="disabled")
            self.summary_edit_frame.pack(fill="both", expand=True)
        else:
            # L1, L2, L3 — reuse existing summary logic
            self._populate_summary_edit(meta)
            self.summary_edit_frame.pack(fill="both", expand=True)

    def _populate_tab_template(self, meta: Dict[str, Any]) -> None:
        """Populate 模板内容 tab based on selected node level."""
        level = meta["level"]
        self.l5_edit_frame.pack_forget()
        self.template_info_frame.pack_forget()

        if level == self.L5_SUBMODULE:
            self._populate_l5_edit(meta)
            self.l5_edit_frame.pack(fill="both", expand=True)
        elif level == self.L4_ITEM:
            item = meta.get("item", {})
            sub_modules = item.get("sub_modules", [])
            self.template_info_label.configure(
                text="请选择该条目下的子模块以编辑模板内容")
            info = f'该条目包含 {len(sub_modules)} 个子模块:\n\n'
            if sub_modules:
                for i, sm in enumerate(sub_modules, 1):
                    info += f'  {i}. {sm.get("name", "")}\n'
            else:
                info += '  （暂无子模块，请在"子模块"标签页中添加）\n'
            self.template_info_list.configure(state="normal")
            self.template_info_list.delete("1.0", "end")
            self.template_info_list.insert("1.0", info)
            self.template_info_list.configure(state="disabled")
            self.template_info_frame.pack(fill="both", expand=True)
        else:
            # L1, L2, L3
            self.template_info_label.configure(
                text="请选择一个条目或子模块以编辑模板内容")
            info = (
                "模板编辑器仅在选中子模块（第5级节点）时可用。\n\n"
                "请先在左侧树形结构中选择一个条目（第4级节点），\n"
                "再展开该条目选择其下的子模块进行编辑。"
            )
            self.template_info_list.configure(state="normal")
            self.template_info_list.delete("1.0", "end")
            self.template_info_list.insert("1.0", info)
            self.template_info_list.configure(state="disabled")
            self.template_info_frame.pack(fill="both", expand=True)

    def _populate_tab_submodule(self, meta: Dict[str, Any]) -> None:
        """Populate 子模块 tab based on selected node level."""
        level = meta["level"]
        self.submodule_frame.pack(fill="both", expand=True)

        if level == self.L4_ITEM:
            item = meta.get("item", {})
            sub_modules = item.get("sub_modules", [])
            self.submodule_title.configure(
                text=f'📦 子模块管理 — {item.get("title", "")}')
            info = f'条目: {item.get("title", "")}\n子模块数量: {len(sub_modules)}\n\n'
            if sub_modules:
                info += '子模块列表:\n'
                for i, sm in enumerate(sub_modules, 1):
                    info += f'  {i}. {sm.get("name", "")}\n'
            else:
                info += '（暂无子模块，请使用下方按钮添加）'
            self.submodule_info_text.configure(state="normal")
            self.submodule_info_text.delete("1.0", "end")
            self.submodule_info_text.insert("1.0", info)
            self.submodule_info_text.configure(state="disabled")
            self.submodule_add_btn.pack(
                side="left", expand=True, fill="x", padx=(0, 4))
            self.submodule_del_btn.pack(
                side="left", expand=True, fill="x", padx=(4, 0))
        elif level == self.L5_SUBMODULE:
            sub = meta.get("sub_module", {})
            item = meta.get("item", {})
            self.submodule_title.configure(text="📝 子模块详情")
            info = (
                f'子模块名称: {sub.get("name", "")}\n'
                f'所属条目: {item.get("title", "")}\n'
                f'所属章节: {meta.get("cat_title", "")}\n'
                f'设计阶段: {meta.get("stage", "")}\n'
                f'工程类型: {meta.get("label", meta.get("code", ""))}'
            )
            self.submodule_info_text.configure(state="normal")
            self.submodule_info_text.delete("1.0", "end")
            self.submodule_info_text.insert("1.0", info)
            self.submodule_info_text.configure(state="disabled")
            self.submodule_add_btn.pack_forget()
            self.submodule_del_btn.pack(
                side="left", expand=True, fill="x", padx=(0, 4))
        else:
            # L1, L2, L3
            breadcrumb = self._build_breadcrumb(meta)
            self.submodule_title.configure(text="📦 子模块管理")
            info = (
                f'当前路径: {breadcrumb}\n\n'
                '请在左侧树形结构中选择一个条目（第4级节点），\n'
                '即可在此管理该条目下的子模块。'
            )
            self.submodule_info_text.configure(state="normal")
            self.submodule_info_text.delete("1.0", "end")
            self.submodule_info_text.insert("1.0", info)
            self.submodule_info_text.configure(state="disabled")
            self.submodule_add_btn.pack_forget()
            self.submodule_del_btn.pack_forget()

    # ── Numbering toggle ────────────────────────────────────────────────────

    def _on_numbering_toggle(self) -> None:
        """Save numbering config and refresh tree display."""
        config = self.ctx.engine.get_numbering_config()
        config["show_in_tree"] = self.show_numbering_var.get()
        config["updated_at"] = datetime.now().isoformat()
        self.ctx.engine.set_numbering_config(config)
        self.refresh()

    # ── Data loading ───────────────────────────────────────────────────────

    @staticmethod
    def _iter_cats_with_path(categories: dict):
        """Iterate categories yielding (key, data, path).

        For container categories (road 自控), path = [parent_key, sub_key].
        For standard categories, path = [key].
        """
        for key, data in categories.items():
            if not isinstance(data, dict):
                continue
            if 'items' in data:
                yield key, data, [key]
            else:
                for sub_key, sub_data in data.items():
                    if isinstance(sub_data, dict) and 'items' in sub_data:
                        yield sub_key, sub_data, [key, sub_key]

    @staticmethod
    def _navigate_categories(categories: dict, cat_path: list) -> Optional[dict]:
        """Navigate to a category by its path."""
        node = categories
        for key in cat_path:
            if isinstance(node, dict) and key in node:
                node = node[key]
            else:
                return None
        return node if isinstance(node, dict) and 'items' in node else None

    def refresh(self) -> None:
        """Rebuild the entire tree from JSON data."""
        for item in self.tree.get_children():
            self.tree.delete(item)
        self._node_meta.clear()
        self._sel_iid = None

        filter_val = self.filter_var.get()
        numbering_cfg = self.ctx.engine.get_numbering_config()
        show_numbers = numbering_cfg.get("enabled", True) and numbering_cfg.get("show_in_tree", True)

        for code, label in PROJECT_TYPES:
            if filter_val != "全部" and label != filter_val:
                continue
            rule = self.ctx.engine.load_rule(code)
            if not rule:
                continue

            icon = GenerateEngine.PROJECT_TYPES.get(code, {}).get('icon', '📁')

            # Level 1: project type
            l1_iid = code
            self.tree.insert("", "end", iid=l1_iid,
                             text=f"{icon} {label}", open=True)
            self._node_meta[l1_iid] = {
                "level": self.L1_PROJECT, "code": code, "label": label,
            }

            for stage_name, stage_data in rule.get("design_stages", {}).items():
                # Level 2: design stage
                l2_iid = f"{code}_{stage_name}"
                l2_text = f"📋 {stage_name}"
                if show_numbers:
                    l2_text = self.ctx.engine.get_numbered_title(
                        stage_name, 0, 0, l2_text)
                self.tree.insert(l1_iid, "end", iid=l2_iid,
                                 text=l2_text, open=False)
                self._node_meta[l2_iid] = {
                    "level": self.L2_STAGE, "code": code,
                    "stage": stage_name, "label": label,
                }

                categories = stage_data.get("sections", stage_data) if isinstance(stage_data, dict) else {}

                cat_idx = 0
                for cat_key, cat_data, cat_path in self._iter_cats_with_path(categories):
                    cat_idx += 1
                    cat_title = cat_data.get("title", cat_key)
                    section_id = cat_data.get("section_id", "")

                    # Determine icon for category
                    if "电气" in cat_key or "电气" in cat_title:
                        cat_icon = "⚡"
                    elif "自控" in cat_key or "自控" in cat_title or "智能" in cat_title:
                        cat_icon = "🔧"
                    else:
                        cat_icon = "📑"

                    # Level 3: category
                    l3_iid = f"{code}_{stage_name}_{'_'.join(cat_path)}"
                    l3_text = f"{cat_icon} {cat_title}"
                    if show_numbers:
                        l3_text = self.ctx.engine.get_numbered_title(
                            stage_name, cat_idx, 0, l3_text)
                    self.tree.insert(l2_iid, "end", iid=l3_iid,
                                     text=l3_text, open=False)
                    self._node_meta[l3_iid] = {
                        "level": self.L3_CATEGORY, "code": code,
                        "stage": stage_name, "cat_key": cat_key,
                        "cat_path": cat_path, "cat_title": cat_title,
                        "section_id": section_id,
                    }

                    for item in cat_data.get("items", []):
                        order = item.get("order", 0)
                        # Level 4: item
                        l4_iid = f"{l3_iid}_{order}"
                        l4_text = item.get("title", "")
                        if show_numbers:
                            l4_text = self.ctx.engine.get_numbered_title(
                                stage_name, cat_idx, order, l4_text)
                        self.tree.insert(l3_iid, "end", iid=l4_iid,
                                         text=l4_text, open=False)
                        self._node_meta[l4_iid] = {
                            "level": self.L4_ITEM, "code": code,
                            "stage": stage_name, "cat_key": cat_key,
                            "cat_path": cat_path, "cat_title": cat_title,
                            "section_id": section_id, "order": order,
                            "item": item,
                        }

                        for sub_idx, sub_mod in enumerate(item.get("sub_modules", [])):
                            # Level 5: sub-module
                            l5_iid = f"{l4_iid}_{sub_idx}"
                            self.tree.insert(l4_iid, "end", iid=l5_iid,
                                             text=f"📝 {sub_mod.get('name', '')}",
                                             open=False)
                            self._node_meta[l5_iid] = {
                                "level": self.L5_SUBMODULE, "code": code,
                                "stage": stage_name, "cat_key": cat_key,
                                "cat_path": cat_path, "cat_title": cat_title,
                                "section_id": section_id, "order": order,
                                "sub_idx": sub_idx, "sub_module": sub_mod,
                                "item": item,
                            }

    def _on_select(self, event) -> None:
        sel = self.tree.selection()
        if not sel:
            return
        new_iid = sel[0]
        # Unsaved changes detection for Level 5 edit mode
        if (self.edit_mode and self._sel_iid
                and self._sel_iid != new_iid
                and self._is_l5_editing()):
            current_content = self.l5_editor.get_content()
            if current_content != self._l5_original_content:
                answer = messagebox.askyesno(
                    "未保存的修改",
                    "当前子模块内容已修改，是否保存？",
                )
                if answer:
                    self._save_submodule()
        self._sel_iid = new_iid
        self._show_selected()

    # ── Edit operations ──

    def _save_changes(self) -> None:
        """Save edited Level 4 item back to JSON."""
        if not self._sel_iid:
            return
        meta = self._node_meta.get(self._sel_iid)
        if not meta or meta["level"] != self.L4_ITEM:
            return

        rule = self.ctx.engine.load_rule(meta["code"])
        if not rule:
            return
        stage_data = rule.get("design_stages", {}).get(meta["stage"], {})
        categories = stage_data.get("sections", stage_data) if isinstance(stage_data, dict) else {}
        cat_data = self._navigate_categories(categories, meta["cat_path"])
        if not cat_data:
            return

        for item in cat_data.get("items", []):
            if item.get("order") == meta["order"]:
                item["title"] = self.edit_title.get().strip()
                item["requirement"] = self.edit_requirement.get("1.0", "end-1c").strip()
                item["has_calculation"] = self.edit_calc_var.get()
                item["table_required"] = self.edit_table_var.get()
                item["calc_from_excel"] = self.edit_excel_var.get()
                break

        if self.ctx.engine.save_rules_to_json(meta["code"], rule):
            # Update tree node text directly (avoids full refresh)
            new_title = self.edit_title.get().strip()
            self.tree.item(self._sel_iid, text=new_title)
            # Also update cached meta data
            meta["item"]["title"] = new_title
            # Refresh right panel to reflect changes
            self._show_selected()
        else:
            self.edit_title.configure(placeholder_text="保存失败!")

    def _save_submodule(self) -> None:
        """Save edited submodule content as a custom template via CustomTemplateManager."""
        if not self._sel_iid:
            return
        meta = self._node_meta.get(self._sel_iid)
        if not meta or meta["level"] != self.L5_SUBMODULE:
            return

        content = self.l5_editor.get_content()
        item = meta.get("item", {})
        sub = meta.get("sub_module", {})

        template = {
            "project_type": meta["code"],
            "design_stage": meta["stage"],
            "category": meta.get("cat_key", ""),
            "item_title": item.get("title", ""),
            "sub_module_name": sub.get("name", ""),
            "template_name": "自定义_" + datetime.now().strftime("%m%d_%H%M"),
            "content": content,
            "is_rich_text": True,
        }

        try:
            saved = self.ctx.engine.save_custom_template(template)
            self._l5_original_content = content
            messagebox.showinfo("保存成功", f"已保存为自定义模板：{saved['template_name']}")
            # Refresh template combo
            self._update_variant_info(meta)
        except Exception as e:
            messagebox.showerror("保存失败", str(e))

    def _save_submodule_as(self) -> None:
        """Save with a user-provided template name."""
        if not self._sel_iid:
            return
        meta = self._node_meta.get(self._sel_iid)
        if not meta or meta["level"] != self.L5_SUBMODULE:
            return

        dialog = ctk.CTkInputDialog(
            text="请输入模板名称：",
            title="另存为自定义模板",
        )
        name = dialog.get_input()
        if not name or not name.strip():
            return
        name = name.strip()

        content = self.l5_editor.get_content()
        item = meta.get("item", {})
        sub = meta.get("sub_module", {})

        template = {
            "project_type": meta["code"],
            "design_stage": meta["stage"],
            "category": meta.get("cat_key", ""),
            "item_title": item.get("title", ""),
            "sub_module_name": sub.get("name", ""),
            "template_name": name,
            "content": content,
            "is_rich_text": True,
        }

        try:
            saved = self.ctx.engine.save_custom_template(template)
            self._l5_original_content = content
            messagebox.showinfo("保存成功", f"已另存为自定义模板：{saved['template_name']}")
            self._update_variant_info(meta)
        except Exception as e:
            messagebox.showerror("保存失败", str(e))

    def _update_variant_info(self, meta=None) -> None:
        """Update the variant status label and meta label for the current sub-module."""
        if meta is None:
            if not self._sel_iid:
                return
            meta = self._node_meta.get(self._sel_iid)
        if not meta or meta.get("level") != self.L5_SUBMODULE:
            return

        sub = meta.get("sub_module", {})
        variant_id = sub.get("selected_variant_id", "")
        if variant_id:
            try:
                tpl = self.ctx.engine.template_manager.get_template(variant_id)
                if tpl:
                    name = tpl.get("template_name", "未命名")
                    self.variant_status_label.configure(text=f"当前模板：{name}")
                    updated = tpl.get("updated_at", "")
                    if updated and "T" in updated:
                        updated = updated.split("T")[0]
                    self.variant_meta_label.configure(
                        text=f"最后修改：{updated}" if updated else ""
                    )
                    return
            except Exception:
                pass
            # Template record not found (was deleted)
            self.variant_status_label.configure(
                text=f"当前模板：{variant_id[:8]}…（已删除）")
            self.variant_meta_label.configure(text="")
        else:
            self.variant_status_label.configure(text="当前模板：默认")
            self.variant_meta_label.configure(text="")

    def _open_variant_selector(self) -> None:
        """Open a searchable dialog to select a template variant for the current sub-module."""
        meta = self._node_meta.get(self._sel_iid)
        if not meta or meta.get("level") != self.L5_SUBMODULE:
            return

        dialog = ctk.CTkToplevel(self)
        dialog.title("选择模板变体")
        dialog.geometry("560x500")
        dialog.transient(self)
        dialog.grab_set()

        result = {"variant_id": None}

        # ── Search entry ──
        search_var = ctk.StringVar()
        search_entry = ctk.CTkEntry(
            dialog, placeholder_text="🔍 搜索模板名称或内容...",
            textvariable=search_var,
        )
        search_entry.pack(fill="x", padx=15, pady=(15, 8))

        # ── Template list (Treeview) ──
        frame = ctk.CTkFrame(dialog)
        frame.pack(fill="both", expand=True, padx=15, pady=(0, 8))

        columns = ("name", "tags", "updated")
        tree = ttk.Treeview(frame, columns=columns, show="headings", height=14)
        tree.heading("name", text="模板名称")
        tree.heading("tags", text="标签")
        tree.heading("updated", text="最后修改")
        tree.column("name", width=220, minwidth=150)
        tree.column("tags", width=120, minwidth=80)
        tree.column("updated", width=130, minwidth=100)
        tree.pack(fill="both", expand=True, side="left")

        scrollbar = ttk.Scrollbar(frame, orient="vertical", command=tree.yview)
        scrollbar.pack(side="right", fill="y")
        tree.configure(yscrollcommand=scrollbar.set)

        def refresh_list(keyword=""):
            for item in tree.get_children():
                tree.delete(item)
            try:
                templates = self.ctx.engine.template_manager.search_templates(
                    project_type=meta["code"],
                    sub_module_name=meta["sub_module"].get("name", ""),
                    keyword=keyword if keyword else None,
                )
            except Exception:
                templates = []
            # Sort by updated_at descending
            templates.sort(key=lambda t: t.get("updated_at", ""), reverse=True)
            self._variant_search_results = templates
            for tpl in templates:
                tags = tpl.get("tags", "")
                if isinstance(tags, (list, tuple)):
                    tags = ", ".join(tags)
                updated = tpl.get("updated_at", "")
                if updated and "T" in updated:
                    updated = updated.split("T")[0]
                tree.insert("", "end", values=(
                    tpl.get("template_name", ""),
                    tags or "",
                    updated,
                ), iid=tpl.get("id", ""))

        def on_search(*_args):
            refresh_list(search_var.get().strip())

        search_var.trace_add("write", on_search)

        # Initial load
        refresh_list()

        # ── Buttons ──
        btn_frame = ctk.CTkFrame(dialog, fg_color="transparent")
        btn_frame.pack(fill="x", padx=15, pady=(0, 15))

        def _save_variant_id(variant_id):
            """Persist selected_variant_id on the sub_module via load_rule + save_rules_to_json."""
            if variant_id == "__clear__":
                meta["sub_module"].pop("selected_variant_id", None)
            else:
                meta["sub_module"]["selected_variant_id"] = variant_id
            # Save the modified rule to disk
            rule = self.ctx.engine.load_rule(meta["code"])
            if rule:
                self.ctx.engine.save_rules_to_json(meta["code"], rule)
            self._update_variant_info(meta)

        def on_select():
            sel = tree.selection()
            if sel:
                vid = sel[0]
                # Load template content into editor
                for tpl in self._variant_search_results:
                    if tpl.get("id") == vid:
                        content = tpl.get("content", "")
                        self.l5_editor.set_content(content)
                        self._l5_original_content = content
                        break
                _save_variant_id(vid)
                result["variant_id"] = vid
                dialog.destroy()

        def on_clear():
            _save_variant_id("__clear__")
            result["variant_id"] = "__clear__"
            dialog.destroy()

        def on_new():
            dialog.destroy()
            self._new_template_from_selector(meta)

        def on_delete():
            sel = tree.selection()
            if not sel:
                messagebox.showwarning("提示", "请先选中一个模板")
                return
            tpl_id = sel[0]
            tpl_name = tree.item(sel[0], "values")[0]
            answer = messagebox.askyesno(
                "确认删除", f"确定要删除模板「{tpl_name}」吗？",
            )
            if not answer:
                return
            try:
                ok = self.ctx.engine.delete_custom_template(tpl_id)
                if ok:
                    messagebox.showinfo("删除成功", f"已删除模板：{tpl_name}")
                else:
                    messagebox.showwarning("删除失败", "未找到要删除的模板")
            except Exception as e:
                messagebox.showerror("删除失败", str(e))
                return
            refresh_list(search_var.get().strip())

        select_btn = ctk.CTkButton(
            btn_frame, text="选择", command=on_select,
            fg_color=COLORS["accent_blue"], text_color="white",
            corner_radius=6, height=32,
        )
        select_btn.pack(side="left", padx=(0, 4), fill="x", expand=True)

        clear_btn = ctk.CTkButton(
            btn_frame, text="清除选择", command=on_clear,
            fg_color=COLORS["accent_red"], text_color="white",
            corner_radius=6, height=32,
        )
        clear_btn.pack(side="left", padx=(4, 0), fill="x", expand=True)

        new_btn = ctk.CTkButton(
            btn_frame, text="新建模板", command=on_new,
            fg_color=COLORS["accent_green"], text_color="white",
            corner_radius=6, height=32,
        )
        new_btn.pack(side="left", padx=(4, 0), fill="x", expand=True)

        del_btn = ctk.CTkButton(
            btn_frame, text="删除", command=on_delete,
            fg_color=COLORS["accent_red"], text_color="white",
            corner_radius=6, height=32,
        )
        del_btn.pack(side="left", padx=(4, 0), fill="x", expand=True)

        cancel_btn = ctk.CTkButton(
            btn_frame, text="取消", command=dialog.destroy,
            fg_color=COLORS["bg_tertiary"], text_color=COLORS["text_primary"],
            corner_radius=6, height=32,
        )
        cancel_btn.pack(side="left", padx=(4, 0), fill="x", expand=True)

        dialog.wait_window()

    def _new_template_from_selector(self, meta) -> None:
        """Create a blank template from the selector dialog context and auto-select it."""
        dialog = ctk.CTkInputDialog(
            text="请输入模板名称：",
            title="新建自定义模板",
        )
        name = dialog.get_input()
        if not name or not name.strip():
            return
        name = name.strip()

        item = meta.get("item", {})
        sub = meta.get("sub_module", {})

        template = {
            "project_type": meta["code"],
            "design_stage": meta["stage"],
            "category": meta.get("cat_key", ""),
            "item_title": item.get("title", ""),
            "sub_module_name": sub.get("name", ""),
            "template_name": name,
            "content": "",
            "is_rich_text": True,
        }

        try:
            saved = self.ctx.engine.save_custom_template(template)
            self._l5_original_content = ""
            self.l5_editor.set_content("")
            messagebox.showinfo("新建成功", f"已创建空白模板：{saved['template_name']}")
            # Auto-select the newly created template as the variant
            vid = saved.get("id", "")
            if vid:
                meta["sub_module"]["selected_variant_id"] = vid
                rule = self.ctx.engine.load_rule(meta["code"])
                if rule:
                    self.ctx.engine.save_rules_to_json(meta["code"], rule)
            self._update_variant_info(meta)
        except Exception as e:
            messagebox.showerror("新建失败", str(e))

    def _on_clear_variant(self) -> None:
        """Clear the selected variant for the current sub-module."""
        if not self._sel_iid:
            return
        meta = self._node_meta.get(self._sel_iid)
        if not meta or meta.get("level") != self.L5_SUBMODULE:
            return

        meta["sub_module"].pop("selected_variant_id", None)
        rule = self.ctx.engine.load_rule(meta["code"])
        if rule:
            self.ctx.engine.save_rules_to_json(meta["code"], rule)
        self._update_variant_info(meta)
        messagebox.showinfo("已清除", "已恢复为默认模板。")

    def _add_item(self) -> None:
        """Add a new item to the current category (if a Level 3/4 node is selected)."""
        if not self._sel_iid:
            return
        meta = self._node_meta.get(self._sel_iid)
        if not meta:
            return

        # Determine the Level 3 context from the selected node
        code = meta.get("code")
        stage = meta.get("stage")
        cat_path = meta.get("cat_path")
        if not all([code, stage, cat_path]):
            # If Level 1 or 2 selected, cannot add
            return

        rule = self.ctx.engine.load_rule(code)
        if not rule:
            return
        stage_data = rule.get("design_stages", {}).get(stage, {})
        categories = stage_data.get("sections", stage_data) if isinstance(stage_data, dict) else {}
        cat_data = self._navigate_categories(categories, cat_path)
        if not cat_data:
            return

        existing = cat_data.get("items", [])
        selected_level = meta.get("level")

        # Determine insertion index: after selected L4 item, or at end for L3
        if selected_level == self.L4_ITEM and meta.get("cat_path") == cat_path:
            # Insert after the selected L4 item
            sel_order = meta.get("order", 0)
            insert_idx = None
            for i, it in enumerate(existing):
                if it.get("order") == sel_order:
                    insert_idx = i + 1
                    break
            if insert_idx is None or insert_idx > len(existing):
                insert_idx = len(existing)
        else:
            # Append at end
            insert_idx = len(existing)

        max_order = max((it.get("order", 0) for it in existing), default=0)
        new_item = {
            "order": max_order + 1,  # temporary; will be renumbered
            "title": "新条目",
            "requirement": "请输入要求描述",
            "has_calculation": False,
            "table_required": False,
            "calc_from_excel": False,
            "sub_modules": [],
        }
        existing.insert(insert_idx, new_item)
        self._renumber_items_inplace(existing)
        if self.ctx.engine.save_rules_to_json(code, rule):
            self.refresh()

    def _delete_item(self) -> None:
        """Delete the selected Level 4 item."""
        if not self._sel_iid:
            return
        meta = self._node_meta.get(self._sel_iid)
        if not meta or meta["level"] != self.L4_ITEM:
            return

        rule = self.ctx.engine.load_rule(meta["code"])
        if not rule:
            return
        stage_data = rule.get("design_stages", {}).get(meta["stage"], {})
        categories = stage_data.get("sections", stage_data) if isinstance(stage_data, dict) else {}
        cat_data = self._navigate_categories(categories, meta["cat_path"])
        if not cat_data:
            return

        items = cat_data.get("items", [])
        before = len(items)
        cat_data["items"] = [it for it in items if it.get("order") != meta["order"]]
        if len(cat_data["items"]) < before:
            # Renumber remaining items sequentially after deletion
            self._renumber_items_inplace(cat_data["items"])
            if self.ctx.engine.save_rules_to_json(meta["code"], rule):
                self._sel_iid = None
                self.refresh()

    def _add_submodule(self) -> None:
        """Add a new sub-module to the selected Level 4 item."""
        if not self._sel_iid:
            return
        meta = self._node_meta.get(self._sel_iid)
        if not meta:
            return

        # Determine parent L4 iid and meta
        if meta["level"] == self.L4_ITEM:
            parent_iid = self._sel_iid
            parent_meta = meta
        elif meta["level"] == self.L5_SUBMODULE:
            # Derive parent L4 iid from L5 iid: strip last _{sub_idx}
            parent_iid = "_".join(self._sel_iid.split("_")[:-1])
            parent_meta = self._node_meta.get(parent_iid)
            if not parent_meta:
                return
        else:
            messagebox.showwarning("提示", "请先选择一个条目（第4级节点）或子模块（第5级节点）")
            return

        # Get live reference to the item dict
        item = parent_meta.get("item")
        if item is None:
            return

        # Append new sub_module via live dict mutation
        item.setdefault("sub_modules", []).append(
            {"name": "新子模块", "template_content": ""}
        )

        # Persist to JSON (same pattern as _add_item)
        rule = self.ctx.engine.load_rule(parent_meta["code"])
        if rule and self.ctx.engine.save_rules_to_json(parent_meta["code"], rule):
            self.refresh()
            # Expand parent node after rebuild (iid is deterministic)
            if self.tree.exists(parent_iid):
                self.tree.item(parent_iid, open=True)

    def _delete_submodule(self) -> None:
        """Delete the selected Level 5 sub-module."""
        if not self._sel_iid:
            return
        meta = self._node_meta.get(self._sel_iid)
        if not meta:
            return

        level = meta["level"]
        if level == self.L4_ITEM:
            messagebox.showinfo("提示", "请选择一个具体的子模块（第5级节点）进行删除。")
            return
        if level != self.L5_SUBMODULE:
            messagebox.showwarning("提示", "请先选择一个条目或子模块")
            return

        sub_module = meta.get("sub_module", {})
        if not messagebox.askyesno(
            "确认删除", f"确认删除子模块「{sub_module.get('name', '')}」？"
        ):
            return

        # Get parent item (live reference) and remove sub_module by index
        item = meta.get("item")
        sub_idx = meta.get("sub_idx")
        if item is None or sub_idx is None:
            return

        item.get("sub_modules", []).pop(sub_idx)

        # Derive parent L4 iid for post-refresh selection
        parent_iid = "_".join(self._sel_iid.split("_")[:-1])

        # Persist to JSON
        rule = self.ctx.engine.load_rule(meta["code"])
        if rule and self.ctx.engine.save_rules_to_json(meta["code"], rule):
            self.refresh()
            # Select parent L4 node after rebuild
            if self.tree.exists(parent_iid):
                self.tree.selection_set(parent_iid)
                self.tree.see(parent_iid)

    # ── Drag-drop reordering (Level 4 items) ──

    def _on_tree_drag_start(self, event) -> None:
        """Record the potential drag source on button press."""
        iid = self.tree.identify_row(event.y)
        self._drag_iid = None
        self._drag_started = False
        self._drag_x = event.x
        self._drag_y = event.y
        if not iid:
            return
        meta = self._node_meta.get(iid)
        if meta and meta.get("level") == self.L4_ITEM:
            self._drag_iid = iid

    def _on_tree_drag_motion(self, event) -> None:
        """Track drag motion; mark drag as started once threshold exceeded."""
        if self._drag_iid is None:
            return
        if self._drag_started:
            return
        dx = abs(event.x - self._drag_x)
        dy = abs(event.y - self._drag_y)
        if dx >= self._drag_threshold or dy >= self._drag_threshold:
            self._drag_started = True

    def _on_tree_drag_release(self, event) -> None:
        """Handle drop when a dragged L4 item is released."""
        drag_iid = self._drag_iid
        was_drag = self._drag_started
        # Reset drag state immediately
        self._drag_iid = None
        self._drag_started = False

        if not drag_iid or not was_drag:
            return  # Not a drag — just a click

        # Find the drop target
        target_iid = self.tree.identify_row(event.y)
        if not target_iid or target_iid == drag_iid:
            return

        drag_meta = self._node_meta.get(drag_iid)
        target_meta = self._node_meta.get(target_iid)
        if not drag_meta or not target_meta:
            return
        if drag_meta.get("level") != self.L4_ITEM or target_meta.get("level") != self.L4_ITEM:
            return

        # Must be in the same category (same code, stage, cat_path)
        if (drag_meta.get("code") != target_meta.get("code")
                or drag_meta.get("stage") != target_meta.get("stage")
                or drag_meta.get("cat_path") != target_meta.get("cat_path")):
            return

        self._reorder_items(drag_meta, target_meta)

    def _reorder_items(self, drag_meta: Dict[str, Any], target_meta: Dict[str, Any]) -> None:
        """Move drag item to target item's position within the same category, then renumber."""
        code = drag_meta["code"]
        rule = self.ctx.engine.load_rule(code)
        if not rule:
            return
        stage_data = rule.get("design_stages", {}).get(drag_meta["stage"], {})
        categories = stage_data.get("sections", stage_data) if isinstance(stage_data, dict) else {}
        cat_data = self._navigate_categories(categories, drag_meta["cat_path"])
        if not cat_data:
            return

        items = cat_data.get("items", [])
        if len(items) < 2:
            return

        drag_order = drag_meta["order"]
        target_order = target_meta["order"]
        if drag_order == target_order:
            return

        # Find indices
        drag_idx: Optional[int] = None
        target_idx: Optional[int] = None
        for i, it in enumerate(items):
            o = it.get("order", 0)
            if o == drag_order:
                drag_idx = i
            if o == target_order:
                target_idx = i
            if drag_idx is not None and target_idx is not None:
                break

        if drag_idx is None or target_idx is None:
            return

        # Remove dragged item and re-insert at target position
        drag_item = items.pop(drag_idx)
        if drag_idx < target_idx:
            target_idx -= 1
        items.insert(target_idx, drag_item)

        # Renumber and save
        self._renumber_items_inplace(items)
        if self.ctx.engine.save_rules_to_json(code, rule):
            self.refresh()

    # ── Move-up / move-down operations (Level 4 items) ──

    def _move_item_up(self) -> None:
        """Move selected L4 item up one position within its category."""
        if not self._sel_iid:
            return
        meta = self._node_meta.get(self._sel_iid)
        if not meta or meta.get("level") != self.L4_ITEM:
            return

        code = meta["code"]
        rule = self.ctx.engine.load_rule(code)
        if not rule:
            return
        stage_data = rule.get("design_stages", {}).get(meta["stage"], {})
        categories = stage_data.get("sections", stage_data) if isinstance(stage_data, dict) else {}
        cat_data = self._navigate_categories(categories, meta["cat_path"])
        if not cat_data:
            return

        items = cat_data.get("items", [])
        cur_order = meta["order"]
        cur_idx: Optional[int] = None
        for i, it in enumerate(items):
            if it.get("order") == cur_order:
                cur_idx = i
                break
        if cur_idx is None or cur_idx == 0:
            return  # Already at top

        # Swap with previous item
        prev_item = items[cur_idx - 1]
        items[cur_idx - 1] = items[cur_idx]
        items[cur_idx] = prev_item

        self._renumber_items_inplace(items)
        if self.ctx.engine.save_rules_to_json(code, rule):
            self.refresh()

    def _move_item_down(self) -> None:
        """Move selected L4 item down one position within its category."""
        if not self._sel_iid:
            return
        meta = self._node_meta.get(self._sel_iid)
        if not meta or meta.get("level") != self.L4_ITEM:
            return

        code = meta["code"]
        rule = self.ctx.engine.load_rule(code)
        if not rule:
            return
        stage_data = rule.get("design_stages", {}).get(meta["stage"], {})
        categories = stage_data.get("sections", stage_data) if isinstance(stage_data, dict) else {}
        cat_data = self._navigate_categories(categories, meta["cat_path"])
        if not cat_data:
            return

        items = cat_data.get("items", [])
        cur_order = meta["order"]
        cur_idx: Optional[int] = None
        for i, it in enumerate(items):
            if it.get("order") == cur_order:
                cur_idx = i
                break
        if cur_idx is None or cur_idx >= len(items) - 1:
            return  # Already at bottom

        # Swap with next item
        next_item = items[cur_idx + 1]
        items[cur_idx + 1] = items[cur_idx]
        items[cur_idx] = next_item

        self._renumber_items_inplace(items)
        if self.ctx.engine.save_rules_to_json(code, rule):
            self.refresh()

    @staticmethod
    def _renumber_items_inplace(items: list) -> None:
        """Renumber items sequentially (1, 2, 3…) in-place."""
        for idx, it in enumerate(items):
            it["order"] = idx + 1

    # ── Inline tree editing (Level 4 only) ──

    def _on_double_click(self, event) -> None:
        """Double-click on tree → start inline editing for Level 4 nodes."""
        region = self.tree.identify_region(event.x, event.y)
        if region != "tree":
            return
        iid = self.tree.identify_row(event.y)
        if not iid:
            return
        meta = self._node_meta.get(iid)
        if not meta or meta.get("level") != self.L4_ITEM:
            return
        # Select the node first
        self.tree.selection_set(iid)
        self._start_tree_inline_edit(iid)

    def _on_right_click(self, event) -> None:
        """Right-click on tree → show context menu."""
        iid = self.tree.identify_row(event.y)
        if iid:
            self.tree.selection_set(iid)
        meta = self._node_meta.get(iid) if iid else None
        is_l4 = meta is not None and meta.get("level") == self.L4_ITEM

        menu = tk.Menu(self.tree, tearoff=0)

        if is_l4:
            menu.add_command(
                label="✏️ 重命名",
                command=lambda: self._start_tree_inline_edit(iid),
            )
        else:
            menu.add_command(label="✏️ 重命名", state="disabled")

        menu.add_separator()

        if iid and meta and meta.get("level") in (self.L3_CATEGORY, self.L4_ITEM):
            menu.add_command(label="➕ 添加细则", command=self._add_item)
        else:
            menu.add_command(label="➕ 添加细则", state="disabled")

        if is_l4:
            menu.add_command(label="🗑 删除细则", command=self._delete_item)
        else:
            menu.add_command(label="🗑 删除细则", state="disabled")

        menu.add_separator()
        if is_l4:
            menu.add_command(label="⬆ 上移", command=self._move_item_up)
            menu.add_command(label="⬇ 下移", command=self._move_item_down)
        else:
            menu.add_command(label="⬆ 上移", state="disabled")
            menu.add_command(label="⬇ 下移", state="disabled")

        try:
            menu.tk_popup(event.x_root, event.y_root)
        finally:
            menu.grab_release()

    def _cancel_tree_edit(self) -> None:
        """Cancel any ongoing inline tree edit without saving."""
        if self._edit_widget is not None:
            self._edit_widget.destroy()
            self._edit_widget = None
        self._editing_iid = None
        self._original_title = ""

    def _start_tree_inline_edit(self, iid: str) -> None:
        """Start inline editing on a Level 4 tree node."""
        meta = self._node_meta.get(iid)
        if not meta or meta.get("level") != self.L4_ITEM:
            return

        # Cancel any existing edit first
        self._cancel_tree_edit()

        # Ensure the item is visible
        self.tree.see(iid)

        # Get bounding box of the tree cell (column #0 text area)
        bbox = self.tree.bbox(iid, column="#0")
        if not bbox:
            return
        x, y, width, height = bbox

        # Get current title
        current_title = self.tree.item(iid, "text")
        if not current_title:
            current_title = ""

        self._original_title = current_title
        self._editing_iid = iid

        # Create a tk.Entry placed over the tree cell
        entry = tk.Entry(
            self.tree,
            font=("微软雅黑", 11),
            relief="solid",
            borderwidth=1,
        )
        entry.insert(0, current_title)
        entry.select_range(0, "end")
        entry.focus_set()

        entry.bind("<Return>", lambda e: self._finish_tree_inline_edit(e, save=True))
        entry.bind("<Escape>", lambda e: self._finish_tree_inline_edit(e, save=False))
        entry.bind("<FocusOut>", self._on_tree_edit_focus_out)

        entry.place(x=x, y=y, width=width, height=height)
        self._edit_widget = entry

    def _on_tree_edit_focus_out(self, event) -> None:
        """Handle FocusOut on the tree inline edit widget with a delay
        to avoid conflicts with menu clicks."""
        # Use after() to avoid saving before a menu item click is processed
        self.tree.after(100, lambda: self._finish_tree_inline_edit(save=True))

    def _finish_tree_inline_edit(self, event=None, save: bool = True) -> None:
        """Finish inline tree editing, optionally saving the new title."""
        if self._edit_widget is None:
            return

        entry = self._edit_widget
        new_title = entry.get().strip()

        # Cancel the delayed FocusOut handler if Enter/Escape triggered
        self._edit_widget = None
        iid = self._editing_iid
        self._editing_iid = None

        # Destroy the entry widget
        try:
            entry.destroy()
        except Exception:
            pass

        if save and iid and new_title and new_title != self._original_title:
            self._save_tree_title_edit(iid, new_title)
        elif not save:
            # Restore original text in tree display (title unchanged in data)
            if iid:
                self.tree.item(iid, text=self._original_title)
        self._original_title = ""

    def _save_tree_title_edit(self, iid: str, new_title: str) -> bool:
        """Save a title change for a tree node directly to JSON and update tree.

        Returns True on success, False on failure.
        """
        meta = self._node_meta.get(iid)
        if not meta or meta.get("level") != self.L4_ITEM:
            return False

        rule = self.ctx.engine.load_rule(meta["code"])
        if not rule:
            return False
        stage_data = rule.get("design_stages", {}).get(meta["stage"], {})
        categories = stage_data.get("sections", stage_data) if isinstance(stage_data, dict) else {}
        cat_data = self._navigate_categories(categories, meta["cat_path"])
        if not cat_data:
            return False

        for item in cat_data.get("items", []):
            if item.get("order") == meta["order"]:
                item["title"] = new_title
                break

        if self.ctx.engine.save_rules_to_json(meta["code"], rule):
            # Update the tree node text directly (avoid full refresh)
            self.tree.item(iid, text=new_title)
            # Also update cached meta data
            meta["item"]["title"] = new_title
            # Refresh right panel if showing the same item
            if self._sel_iid == iid:
                self._show_selected()
            return True
        return False


# ═══════════════════════════════════════════════════════════════════════════
# 7. PAGE: EquipmentPage
# ═══════════════════════════════════════════════════════════════════════════

class EquipmentPage(ctk.CTkFrame):
    """设备材料表管理页面 — 选择设备类型、预览、导出"""

    def __init__(self, master, ctx: EngineContext, **kwargs) -> None:
        super().__init__(master, fg_color="transparent", **kwargs)
        self.ctx = ctx
        self.eq_gen = EquipmentTableGenerator()
        self._check_vars: List[ctk.BooleanVar] = []
        self._build()

    def _build(self) -> None:
        # Header
        hdr = ctk.CTkFrame(self, fg_color="transparent")
        hdr.pack(fill="x", pady=(10, 15))
        ctk.CTkLabel(hdr, text="📦 设备材料表", font=get_ctk_font("header"),
                     text_color=COLORS["text_primary"]).pack(side="left")

        # Config row
        cfg = ctk.CTkFrame(self, fg_color=COLORS["bg_card"], corner_radius=12)
        cfg.pack(fill="x", padx=0, pady=(0, 10))

        row1 = ctk.CTkFrame(cfg, fg_color="transparent")
        row1.pack(fill="x", padx=15, pady=12)

        ctk.CTkLabel(row1, text="工程类型：", font=get_ctk_font("caption"),
                     text_color=COLORS["text_secondary"]).pack(side="left")
        self.pt_var = ctk.StringVar(value="给水工程")
        self.pt_menu = ctk.CTkOptionMenu(
            row1, values=[lbl for _, lbl in PROJECT_TYPES],
            variable=self.pt_var, font=get_ctk_font("body"),
            corner_radius=8, fg_color=COLORS["bg_primary"],
            button_color=COLORS["accent_blue"],
            command=lambda _: self._refresh_types(),
        )
        self.pt_menu.pack(side="left", padx=(0, 20))

        ctk.CTkLabel(row1, text="设计阶段：", font=get_ctk_font("caption"),
                     text_color=COLORS["text_secondary"]).pack(side="left")
        self.ds_var = ctk.StringVar(value="可行性研究")
        ctk.CTkOptionMenu(
            row1, values=["可行性研究", "初步设计"],
            variable=self.ds_var, font=get_ctk_font("body"),
            corner_radius=8, fg_color=COLORS["bg_primary"],
            button_color=COLORS["accent_blue"],
        ).pack(side="left", padx=(0, 20))

        ctk.CTkButton(
            row1, text="🔄 预览", font=get_ctk_font("body"),
            fg_color=COLORS["accent_blue"], text_color="white",
            corner_radius=8, height=32, width=100, command=self._preview,
        ).pack(side="left")

        ctk.CTkButton(
            row1, text="📄 导出Word", font=get_ctk_font("body"),
            fg_color=COLORS["accent_green"], text_color="white",
            corner_radius=8, height=32, width=120, command=self._export,
        ).pack(side="left", padx=(10, 0))

        # Equipment checklist
        check_card = ctk.CTkFrame(self, fg_color=COLORS["bg_card"], corner_radius=12)
        check_card.pack(fill="x", padx=0, pady=(0, 10))

        ctk.CTkLabel(check_card, text="选择设备类型（可多选）：",
                     font=get_ctk_font("body"),
                     text_color=COLORS["text_primary"]).pack(anchor="w", padx=15, pady=(12, 4))

        self.check_container = ctk.CTkFrame(check_card, fg_color="transparent")
        self.check_container.pack(fill="x", padx=15, pady=(0, 10))

        # Preview area
        prev_card = ctk.CTkFrame(self, fg_color=COLORS["bg_card"], corner_radius=12)
        prev_card.pack(fill="both", expand=True)

        ctk.CTkLabel(prev_card, text="预览：", font=get_ctk_font("body"),
                     text_color=COLORS["text_primary"]).pack(anchor="w", padx=15, pady=(12, 4))

        self.preview_text = ctk.CTkTextbox(
            prev_card, font=("Consolas", 11),
            fg_color=COLORS["bg_primary"], corner_radius=8,
            wrap="none",
        )
        self.preview_text.pack(fill="both", expand=True, padx=15, pady=(0, 15))
        self.preview_text.insert("1.0", "选择工程类型和设备类型后，点击「预览」查看设备材料表")
        self.preview_text.configure(state="disabled")

        self._refresh_types()

    def _get_pt_code(self) -> str:
        lbl = self.pt_var.get()
        for code, label in PROJECT_TYPES:
            if label == lbl:
                return code
        return 'water_supply'

    def _refresh_types(self) -> None:
        """根据所选工程类型刷新设备类型复选框"""
        for w in self.check_container.winfo_children():
            w.destroy()
        self._check_vars.clear()

        pt = self._get_pt_code()
        types = self.eq_gen.get_available_types(pt)
        if not types:
            ctk.CTkLabel(self.check_container, text="暂无可用设备类型",
                         font=get_ctk_font("caption"),
                         text_color=COLORS["text_secondary"]).pack(anchor="w")
            return

        for idx, eq in enumerate(types):
            var = ctk.BooleanVar(value=True)  # 默认全选
            self._check_vars.append(var)
            cb = ctk.CTkCheckBox(
                self.check_container, text=f'{eq["name"]} ({eq["unit"]})',
                variable=var, font=get_ctk_font("caption"),
                text_color=COLORS["text_primary"],
                checkbox_width=18, checkbox_height=18,
            )
            cb.pack(anchor="w", pady=2)

    def _preview(self) -> None:
        """生成设备材料表预览"""
        pt = self._get_pt_code()
        ds = self.ds_var.get()

        # 获取选中的设备类型索引
        selected = [i for i, v in enumerate(self._check_vars) if v.get()]

        # Use empty summary for preview (no load data loaded)
        result = self.eq_gen.generate(pt, {}, [], ds, selected)

        self.preview_text.configure(state="normal")
        self.preview_text.delete("1.0", "end")

        if not result['has_data']:
            self.preview_text.insert("1.0", "未选择任何设备类型，请勾选后重试。")
        else:
            self.preview_text.insert("1.0", self.eq_gen.to_text(result))

        self.preview_text.configure(state="disabled")

    def _export(self) -> None:
        """导出设备材料表为独立的 Word 文档"""
        pt = self._get_pt_code()
        ds = self.ds_var.get()
        selected = [i for i, v in enumerate(self._check_vars) if v.get()]
        result = self.eq_gen.generate(pt, {}, [], ds, selected)

        if not result['has_data']:
            return

        from docx import Document
        from docx.shared import Pt, Cm

        doc = Document()
        # Title
        p = doc.add_paragraph()
        run = p.add_run(f'{self.pt_var.get()} {ds}  主要设备材料表')
        run.bold = True
        run.font.size = Pt(16)

        doc.add_paragraph(result['note'])

        # Table
        headers = result['headers']
        rows = result['rows']
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = 'Table Grid'

        # Header row
        for j, h in enumerate(headers):
            cell = table.rows[0].cells[j]
            cell.text = h

        # Data rows
        for i, row in enumerate(rows):
            for j, val in enumerate(row):
                table.rows[i + 1].cells[j].text = val

        # Save
        ts = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f'设备材料表_{self.pt_var.get()}_{ds}_{ts}.docx'
        output_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'output')
        os.makedirs(output_dir, exist_ok=True)
        path = os.path.join(output_dir, filename)
        doc.save(path)

        self.preview_text.configure(state="normal")
        self.preview_text.delete("1.0", "end")
        self.preview_text.insert("1.0", f'✅ 已导出: {path}')
        self.preview_text.configure(state="disabled")


# ═══════════════════════════════════════════════════════════════════════════
# 7b. PAGE: HistoricalTemplatePage
# ═══════════════════════════════════════════════════════════════════════════

class HistoricalTemplatePage(ctk.CTkFrame):
    """历史模板库页面 — 导入、分类、浏览历史设计文档"""

    # 工程类型显示映射
    _PT_LABELS: Dict[str, str] = {
        "water_supply": "💧 给水工程",
        "drainage":     "🌊 排水工程",
        "road":         "🛣 道路工程",
        "sanitation":   "🗑 环卫工程",
        "unclassified": "❓ 未分类",
    }
    _DS_LABELS: Dict[str, str] = {
        "可行性研究": "📋 可行性研究",
        "初步设计":   "📋 初步设计",
        "unclassified": "❓ 未分类",
    }
    _CAT_LABELS: Dict[str, str] = {
        "电气": "⚡ 电气",
        "自控": "🔧 自控",
        "unclassified": "❓ 未分类",
    }

    def __init__(self, master, ctx: EngineContext, **kwargs) -> None:
        super().__init__(master, fg_color="transparent", **kwargs)
        self.ctx = ctx
        self.importer = HistoricalTemplateImporter()
        self._current_template_id: Optional[str] = None
        self._build()

    def _build(self) -> None:
        # ── Header ──
        hdr = ctk.CTkFrame(self, fg_color="transparent")
        hdr.pack(fill="x", pady=(10, 8))

        ctk.CTkLabel(hdr, text="📚 历史模板库", font=get_ctk_font("header"),
                     text_color=COLORS["text_primary"]).pack(side="left")

        # Search box
        self.search_var = ctk.StringVar()
        self.search_entry = ctk.CTkEntry(
            hdr, placeholder_text="🔍 搜索模板内容...",
            textvariable=self.search_var,
            font=get_ctk_font("body"),
            corner_radius=8, height=34, width=260,
            fg_color=COLORS["bg_card"],
            border_color=COLORS["separator"],
        )
        self.search_entry.pack(side="right", padx=(10, 0))
        self.search_entry.bind("<Return>", lambda _: self._on_search())
        self.search_entry.bind("<KP_Enter>", lambda _: self._on_search())

        ctk.CTkButton(
            hdr, text="搜索", font=get_ctk_font("caption"),
            fg_color=COLORS["accent_blue"], text_color="white",
            corner_radius=8, height=34, width=60,
            command=self._on_search,
        ).pack(side="right")

        # ── Main content: left tree + right detail ──
        body = ctk.CTkFrame(self, fg_color="transparent")
        body.pack(fill="both", expand=True, pady=(0, 10))
        body.grid_columnconfigure(0, weight=1)
        body.grid_columnconfigure(1, weight=2)
        body.grid_rowconfigure(0, weight=1)

        # Left panel: tree
        left_card = ctk.CTkFrame(body, fg_color=COLORS["bg_card"], corner_radius=12)
        left_card.grid(row=0, column=0, sticky="nsew", padx=(0, 5))

        ctk.CTkLabel(left_card, text="分类浏览", font=get_ctk_font("section"),
                     text_color=COLORS["text_primary"]).pack(anchor="w", padx=15, pady=(12, 4))

        # Treeview with scrollbar
        tree_frame = ctk.CTkFrame(left_card, fg_color="transparent")
        tree_frame.pack(fill="both", expand=True, padx=10, pady=(0, 5))

        self.tree = ttk.Treeview(tree_frame, show="tree", selectmode="browse")
        tree_scroll = ttk.Scrollbar(tree_frame, orient="vertical", command=self.tree.yview)
        self.tree.configure(yscrollcommand=tree_scroll.set)
        self.tree.pack(side="left", fill="both", expand=True)
        tree_scroll.pack(side="right", fill="y")
        self.tree.bind("<<TreeviewSelect>>", self._on_tree_select)

        # Import button
        btn_row = ctk.CTkFrame(left_card, fg_color="transparent")
        btn_row.pack(fill="x", padx=10, pady=(0, 10))
        ctk.CTkButton(
            btn_row, text="📂 导入文件", font=get_ctk_font("body"),
            fg_color=COLORS["accent_blue"], text_color="white",
            corner_radius=8, height=34,
            command=self._on_import,
        ).pack(side="left")
        ctk.CTkButton(
            btn_row, text="🔄 刷新", font=get_ctk_font("caption"),
            fg_color=COLORS["bg_primary"], text_color=COLORS["text_secondary"],
            corner_radius=8, height=34, width=70,
            command=self._populate_tree,
        ).pack(side="left", padx=(8, 0))

        # Right panel: detail
        right_card = ctk.CTkFrame(body, fg_color=COLORS["bg_card"], corner_radius=12)
        right_card.grid(row=0, column=1, sticky="nsew", padx=(5, 0))

        # Detail header
        detail_hdr = ctk.CTkFrame(right_card, fg_color="transparent")
        detail_hdr.pack(fill="x", padx=15, pady=(12, 4))
        self.detail_title = ctk.CTkLabel(
            detail_hdr, text="选择一个模板查看详情",
            font=get_ctk_font("section"),
            text_color=COLORS["text_primary"], anchor="w",
        )
        self.detail_title.pack(anchor="w")

        # Meta info
        meta_frame = ctk.CTkFrame(right_card, fg_color="transparent")
        meta_frame.pack(fill="x", padx=15, pady=(0, 4))

        self.meta_source = ctk.CTkLabel(
            meta_frame, text="", font=get_ctk_font("caption"),
            text_color=COLORS["text_secondary"], anchor="w",
        )
        self.meta_source.pack(anchor="w")
        self.meta_classify = ctk.CTkLabel(
            meta_frame, text="", font=get_ctk_font("caption"),
            text_color=COLORS["text_secondary"], anchor="w",
        )
        self.meta_classify.pack(anchor="w")
        self.meta_confidence = ctk.CTkLabel(
            meta_frame, text="", font=get_ctk_font("caption"),
            anchor="w",
        )
        self.meta_confidence.pack(anchor="w")

        # Content preview
        ctk.CTkLabel(right_card, text="内容预览：", font=get_ctk_font("caption"),
                     text_color=COLORS["text_secondary"]).pack(anchor="w", padx=15, pady=(6, 2))
        self.content_preview = ctk.CTkTextbox(
            right_card, font=get_ctk_font("mono"),
            fg_color=COLORS["bg_primary"], corner_radius=8,
            wrap="word",
        )
        self.content_preview.pack(fill="both", expand=True, padx=15, pady=(0, 4))
        self.content_preview.configure(state="disabled")

        # Notes section
        notes_frame = ctk.CTkFrame(right_card, fg_color="transparent")
        notes_frame.pack(fill="x", padx=15, pady=(0, 4))
        ctk.CTkLabel(notes_frame, text="备注：", font=get_ctk_font("caption"),
                     text_color=COLORS["text_secondary"]).pack(side="left")
        self.notes_entry = ctk.CTkEntry(
            notes_frame, placeholder_text="添加备注...",
            font=get_ctk_font("caption"),
            corner_radius=6, height=30,
            fg_color=COLORS["bg_primary"],
        )
        self.notes_entry.pack(side="left", fill="x", expand=True, padx=(4, 4))
        ctk.CTkButton(
            notes_frame, text="保存备注", font=get_ctk_font("caption"),
            fg_color=COLORS["accent_green"], text_color="white",
            corner_radius=6, height=30, width=70,
            command=self._on_save_notes,
        ).pack(side="left")

        # Bottom buttons
        bottom_btns = ctk.CTkFrame(right_card, fg_color="transparent")
        bottom_btns.pack(fill="x", padx=15, pady=(0, 12))
        self.delete_btn = ctk.CTkButton(
            bottom_btns, text="🗑 删除", font=get_ctk_font("caption"),
            fg_color=COLORS["accent_red"], text_color="white",
            corner_radius=8, height=32, width=80,
            command=self._on_delete, state="disabled",
        )
        self.delete_btn.pack(side="right")

        # Initial tree population
        self._populate_tree()

    # ── Tree population ───────────────────────────────────────────────

    def _populate_tree(self) -> None:
        """从导入器加载数据并构建分类树。"""
        self.tree.delete(*self.tree.get_children())
        templates = self.importer.list_templates()

        # 构建层级: project_type → design_stage → category → items
        hierarchy: Dict[str, Dict[str, Dict[str, List[dict]]]] = {}
        for tpl in templates:
            pt = tpl.get("project_type", "unclassified")
            ds = tpl.get("design_stage", "unclassified")
            cat = tpl.get("category", "unclassified")
            hierarchy.setdefault(pt, {}).setdefault(ds, {}).setdefault(cat, []).append(tpl)

        # 按固定顺序显示工程类型
        pt_order = ["water_supply", "drainage", "road", "sanitation", "unclassified"]
        for pt in pt_order:
            if pt not in hierarchy:
                continue
            pt_label = self._PT_LABELS.get(pt, pt)
            pt_count = sum(
                len(items)
                for ds_dict in hierarchy[pt].values()
                for items in ds_dict.values()
            )
            pt_iid = self.tree.insert("", "end", text=f"{pt_label} ({pt_count})", open=True)

            ds_order = ["可行性研究", "初步设计", "unclassified"]
            for ds in ds_order:
                if ds not in hierarchy[pt]:
                    continue
                ds_label = self._DS_LABELS.get(ds, ds)
                ds_count = sum(len(v) for v in hierarchy[pt][ds].values())
                ds_iid = self.tree.insert(pt_iid, "end", text=f"{ds_label} ({ds_count})", open=True)

                cat_order = ["电气", "自控", "unclassified"]
                for cat in cat_order:
                    if cat not in hierarchy[pt][ds]:
                        continue
                    cat_label = self._CAT_LABELS.get(cat, cat)
                    items = hierarchy[pt][ds][cat]
                    cat_iid = self.tree.insert(
                        ds_iid, "end",
                        text=f"{cat_label} ({len(items)})",
                        open=True,
                    )
                    for tpl in items:
                        title = tpl.get("title", "未命名")
                        if len(title) > 30:
                            title = title[:30] + "..."
                        self.tree.insert(
                            cat_iid, "end",
                            text=f"📄 {title}",
                            values=(tpl["id"],),
                        )

    # ── Event handlers ────────────────────────────────────────────────

    def _on_tree_select(self, _event: Any) -> None:
        """树节点选中时显示详情。"""
        sel = self.tree.selection()
        if not sel:
            return
        item = self.tree.item(sel[0])
        values = item.get("values", [])
        if not values:
            # 非叶子节点（分类节点）
            self._clear_detail()
            return

        template_id = values[0]
        tpl = self.importer.get_template(template_id)
        if not tpl:
            return

        self._current_template_id = template_id
        self._show_detail(tpl)

    def _show_detail(self, tpl: dict) -> None:
        """在右侧面板显示模板详情。"""
        self.detail_title.configure(text=tpl.get("title", "未命名"))

        source = tpl.get("source_file", "")
        imported = tpl.get("imported_at", "")
        if imported:
            try:
                dt = datetime.fromisoformat(imported)
                imported = dt.strftime("%Y-%m-%d %H:%M")
            except ValueError:
                pass
        self.meta_source.configure(text=f"来源: {source}  |  导入: {imported}")

        pt = self._PT_LABELS.get(tpl.get("project_type", ""), tpl.get("project_type", ""))
        ds = tpl.get("design_stage", "")
        cat = self._CAT_LABELS.get(tpl.get("category", ""), tpl.get("category", ""))
        item = tpl.get("item_title", "")
        self.meta_classify.configure(text=f"分类: {pt} > {ds} > {cat} > {item}")

        confidence = tpl.get("_confidence", 0.0)
        if confidence >= 0.7:
            conf_color = COLORS["accent_green"]
            conf_text = "高"
        elif confidence >= 0.4:
            conf_color = COLORS["accent_orange"]
            conf_text = "中"
        else:
            conf_color = COLORS["accent_red"]
            conf_text = "低"
        self.meta_confidence.configure(
            text=f"置信度: {conf_text} ({confidence:.0%})",
            text_color=conf_color,
        )

        # Content preview
        content = tpl.get("content", "")
        self.content_preview.configure(state="normal")
        self.content_preview.delete("1.0", "end")
        self.content_preview.insert("1.0", content[:5000])
        self.content_preview.configure(state="disabled")

        # Notes
        self.notes_entry.delete(0, "end")
        self.notes_entry.insert(0, tpl.get("notes", ""))

        self.delete_btn.configure(state="normal")

    def _clear_detail(self) -> None:
        """清空右侧详情面板。"""
        self._current_template_id = None
        self.detail_title.configure(text="选择一个模板查看详情")
        self.meta_source.configure(text="")
        self.meta_classify.configure(text="")
        self.meta_confidence.configure(text="")
        self.content_preview.configure(state="normal")
        self.content_preview.delete("1.0", "end")
        self.content_preview.configure(state="disabled")
        self.notes_entry.delete(0, "end")
        self.delete_btn.configure(state="disabled")

    def _on_import(self) -> None:
        """打开文件对话框导入文件。"""
        file_paths = filedialog.askopenfilenames(
            title="选择要导入的设计文档",
            filetypes=[
                ("支持的格式", "*.docx *.txt *.md"),
                ("Word 文档", "*.docx"),
                ("文本文件", "*.txt"),
                ("Markdown", "*.md"),
                ("所有文件", "*.*"),
            ],
        )
        if not file_paths:
            return

        results = self.importer.import_files(list(file_paths))
        count = len(results)
        self._populate_tree()
        messagebox.showinfo("导入完成", f"成功导入 {count} 个模板记录")

    def _on_search(self) -> None:
        """搜索模板。"""
        query = self.search_var.get().strip()
        if not query:
            self._populate_tree()
            return

        results = self.importer.list_templates(search=query)
        self.tree.delete(*self.tree.get_children())

        for tpl in results:
            title = tpl.get("title", "未命名")
            if len(title) > 40:
                title = title[:40] + "..."
            self.tree.insert(
                "", "end",
                text=f"📄 {title}",
                values=(tpl["id"],),
            )

    def _on_save_notes(self) -> None:
        """保存当前模板的备注。"""
        if not self._current_template_id:
            return
        notes = self.notes_entry.get()
        ok = self.importer.update_notes(self._current_template_id, notes)
        if ok:
            self._populate_tree()

    def _on_delete(self) -> None:
        """删除当前模板。"""
        if not self._current_template_id:
            return
        if not messagebox.askyesno("确认删除", "确定要删除此模板吗？"):
            return
        ok = self.importer.delete_template(self._current_template_id)
        if ok:
            self._clear_detail()
            self._populate_tree()

    def refresh(self) -> None:
        """刷新页面数据。"""
        self._populate_tree()


# ═══════════════════════════════════════════════════════════════════════════
# 8. PAGE: AboutPage
# ═══════════════════════════════════════════════════════════════════════════

class AboutPage(ctk.CTkFrame):
    """About / system status page with project save/load."""

    def __init__(self, master, ctx: EngineContext, **kwargs) -> None:
        super().__init__(master, fg_color="transparent", **kwargs)
        self.ctx = ctx
        self._build()

    # ── UI ─────────────────────────────────────────────────────────────────

    def _build(self) -> None:
        # Centered content
        content = ctk.CTkFrame(self, fg_color="transparent")
        content.pack(expand=True, fill="both", padx=80)

        # App info
        ctk.CTkLabel(content, text="⚡", font=get_ctk_font("title"),
                     text_color=COLORS["accent_blue"]).pack(pady=(30, 0))
        ctk.CTkLabel(content, text=APP_TITLE, font=get_ctk_font("title"),
                     text_color=COLORS["text_primary"]).pack(pady=(5, 2))
        ctk.CTkLabel(content, text=f"版本 {APP_VERSION}", font=get_ctk_font("body"),
                     text_color=COLORS["text_secondary"]).pack()
        desc = (
            "基于《市政公用工程设计文件编制深度规定》（2025年版）\n"
            "自动生成工程电气及自控设计说明书Word文档"
        )
        ctk.CTkLabel(content, text=desc, font=get_ctk_font("caption"),
                     text_color=COLORS["text_secondary"],
                     justify="center").pack(pady=(10, 20))

        # System status card
        status_card = ctk.CTkFrame(content, fg_color=COLORS["bg_card"],
                                   corner_radius=12)
        status_card.pack(fill="x", pady=(0, 15))

        ctk.CTkLabel(status_card, text="系统状态", font=get_ctk_font("section"),
                     text_color=COLORS["text_primary"]).pack(anchor="w", padx=20, pady=(15, 10))

        self.status_labels: Dict[str, ctk.CTkLabel] = {}
        checks = [
            ("规则文件", "rules"),
            ("Excel解析器", "excel"),
            ("Word生成器", "docx"),
            ("python-docx", "pydocx"),
            ("输出目录", "output"),
        ]
        for label, key in checks:
            row = ctk.CTkFrame(status_card, fg_color="transparent")
            row.pack(fill="x", padx=20, pady=3)
            ctk.CTkLabel(row, text=label, font=get_ctk_font("body"),
                         text_color=COLORS["text_primary"],
                         width=120, anchor="w").pack(side="left")
            lbl = ctk.CTkLabel(row, text="⚪ 检查中...", font=get_ctk_font("caption"),
                               text_color=COLORS["text_secondary"])
            lbl.pack(side="right")
            self.status_labels[key] = lbl

        # Project management buttons
        mgmt_card = ctk.CTkFrame(content, fg_color=COLORS["bg_card"],
                                 corner_radius=12)
        mgmt_card.pack(fill="x")

        ctk.CTkLabel(mgmt_card, text="项目管理", font=get_ctk_font("section"),
                     text_color=COLORS["text_primary"]).pack(anchor="w", padx=20, pady=(15, 10))

        btn_row = ctk.CTkFrame(mgmt_card, fg_color="transparent")
        btn_row.pack(fill="x", padx=20, pady=(0, 10))
        ctk.CTkButton(
            btn_row, text="💾 保存项目", font=get_ctk_font("body"),
            fg_color=COLORS["accent_blue"], corner_radius=8,
            command=self._save_project,
        ).pack(side="left", padx=(0, 8))
        ctk.CTkButton(
            btn_row, text="📂 加载项目", font=get_ctk_font("body"),
            fg_color=COLORS["bg_selected"],
            text_color=COLORS["text_primary"],
            corner_radius=8, command=self._load_project,
        ).pack(side="left")

        # Recent projects
        self.recent_frame = ctk.CTkFrame(mgmt_card, fg_color="transparent")
        self.recent_frame.pack(fill="x", padx=20, pady=(5, 15))
        self._load_recent_list()

    # ── Health Check ───────────────────────────────────────────────────────

    def refresh(self) -> None:
        status = self.ctx.health_status()
        mapping = {
            "rules":  "规则文件",
            "excel":  "Excel解析器",
            "docx":   "Word生成器",
            "pydocx": "python-docx",
            "output": "输出目录",
        }
        for item_name, msg in status.get("checks", []):
            for key, label in mapping.items():
                if label in item_name:
                    ok = "[OK]" in msg
                    self.status_labels[key].configure(
                        text=f'{"✅" if ok else "❌"} {msg}',
                        text_color=(COLORS["accent_green"]
                                    if ok else COLORS["accent_red"]))
                    break

    # ── Project Save / Load ────────────────────────────────────────────────

    def _save_project(self) -> None:
        path = filedialog.asksaveasfilename(
            title="保存项目",
            defaultextension=".json",
            filetypes=[("JSON文件", "*.json"), ("所有文件", "*.*")],
        )
        if not path:
            return
        data = {
            "params": self.ctx.project_params,
            "excel_path": self.ctx.excel_path,
            "saved_at": datetime.now().isoformat(),
        }
        try:
            with open(path, "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
            self._add_recent(path)
            messagebox.showinfo("保存成功", f"项目已保存至:\n{path}")
        except OSError as exc:
            messagebox.showerror("保存失败", str(exc))

    def _load_project(self) -> None:
        path = filedialog.askopenfilename(
            title="加载项目",
            filetypes=[("JSON文件", "*.json"), ("所有文件", "*.*")],
        )
        if not path:
            return
        try:
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
            self.ctx.project_params = data.get("params", dict(DEFAULT_PARAMS))
            self.ctx.excel_path = data.get("excel_path", "")
            if self.ctx.excel_path and os.path.exists(self.ctx.excel_path):
                self.ctx.excel_data = None
                threading.Thread(
                    target=lambda: self.ctx.engine.parse_excel(self.ctx.excel_path),
                    daemon=True,
                ).start()
            self._add_recent(path)
            messagebox.showinfo("加载成功", f"项目已加载:\n{path}")
        except (OSError, json.JSONDecodeError) as exc:
            messagebox.showerror("加载失败", str(exc))

    def _add_recent(self, path: str) -> None:
        os.makedirs(_RECENT_DIR, exist_ok=True)
        recent: list = []
        if os.path.exists(_RECENT_FILE):
            try:
                with open(_RECENT_FILE, "r", encoding="utf-8") as f:
                    recent = json.load(f)
            except (OSError, json.JSONDecodeError):
                pass
        # Deduplicate and keep last 5
        recent = [r for r in recent if r.get("path") != path]
        recent.insert(0, {"path": path, "name": os.path.basename(path),
                          "time": datetime.now().isoformat()})
        recent = recent[:5]
        with open(_RECENT_FILE, "w", encoding="utf-8") as f:
            json.dump(recent, f, ensure_ascii=False, indent=2)
        self._load_recent_list()

    def _load_recent_list(self) -> None:
        for w in self.recent_frame.winfo_children():
            w.destroy()
        if not os.path.exists(_RECENT_FILE):
            ctk.CTkLabel(self.recent_frame, text="暂无最近项目",
                         font=get_ctk_font("caption"),
                         text_color=COLORS["text_tertiary"]).pack(anchor="w")
            return
        try:
            with open(_RECENT_FILE, "r", encoding="utf-8") as f:
                recent = json.load(f)
        except (OSError, json.JSONDecodeError):
            recent = []
        if not recent:
            ctk.CTkLabel(self.recent_frame, text="暂无最近项目",
                         font=get_ctk_font("caption"),
                         text_color=COLORS["text_tertiary"]).pack(anchor="w")
            return
        ctk.CTkLabel(self.recent_frame, text="最近项目:", font=get_ctk_font("caption"),
                     text_color=COLORS["text_secondary"]).pack(anchor="w", pady=(0, 4))
        for entry in recent:
            name = entry.get("name", "未知")
            path = entry.get("path", "")
            btn = ctk.CTkButton(
                self.recent_frame, text=f"📄 {name}", font=get_ctk_font("caption"),
                fg_color="transparent", text_color=COLORS["accent_blue"],
                hover_color=COLORS["bg_hover"], corner_radius=6,
                anchor="w",
                command=lambda p=path: self._load_project_by_path(p),
            )
            btn.pack(fill="x", pady=1)

    def _load_project_by_path(self, path: str) -> None:
        if not os.path.exists(path):
            messagebox.showerror("错误", f"文件不存在:\n{path}")
            return
        try:
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
            self.ctx.project_params = data.get("params", dict(DEFAULT_PARAMS))
            self.ctx.excel_path = data.get("excel_path", "")
            if self.ctx.excel_path and os.path.exists(self.ctx.excel_path):
                self.ctx.excel_data = None
            self._add_recent(path)
            messagebox.showinfo("加载成功", f"项目已加载:\n{path}")
        except Exception as exc:
            messagebox.showerror("加载失败", str(exc))


# ═══════════════════════════════════════════════════════════════════════════
# 7c. PAGE: TemplateLibraryPage
# ═══════════════════════════════════════════════════════════════════════════

class TemplateLibraryPage(ctk.CTkFrame):
    """模板库 — 浏览、搜索和管理 CustomTemplateManager 中的模板变体"""

    def __init__(self, master, ctx: EngineContext, **kwargs) -> None:
        super().__init__(master, fg_color="transparent", **kwargs)
        self.ctx = ctx
        self._templates: List[dict] = []
        self._build()

    def _build(self) -> None:
        # ── Title ──
        ctk.CTkLabel(self, text="📚 模板库", font=get_ctk_font("header"),
                     text_color=COLORS["text_primary"]).pack(anchor="w", pady=(10, 15))

        # ── Filter row ──
        filter_frame = ctk.CTkFrame(self, fg_color="transparent")
        filter_frame.pack(fill="x", pady=(0, 10))

        ctk.CTkLabel(filter_frame, text="工程类型：", font=get_ctk_font("caption"),
                     text_color=COLORS["text_secondary"]).pack(side="left")
        self.filter_type = ctk.CTkComboBox(
            filter_frame,
            values=["全部", "给水工程", "排水工程", "道路交通工程", "环境卫生工程"],
            width=130,
            state="readonly",
        )
        self.filter_type.pack(side="left", padx=(4, 12))
        self.filter_type.set("全部")
        self.filter_type.configure(command=lambda _: self._refresh())

        ctk.CTkLabel(filter_frame, text="子模块：", font=get_ctk_font("caption"),
                     text_color=COLORS["text_secondary"]).pack(side="left")
        self.filter_submodule = ctk.CTkComboBox(
            filter_frame,
            values=["全部"],
            width=200,
            state="readonly",
        )
        self.filter_submodule.pack(side="left", padx=(4, 12))
        self.filter_submodule.set("全部")
        self.filter_submodule.configure(command=lambda _: self._refresh())

        ctk.CTkLabel(filter_frame, text="关键词：", font=get_ctk_font("caption"),
                     text_color=COLORS["text_secondary"]).pack(side="left")
        self.filter_keyword = ctk.CTkEntry(filter_frame, width=150)
        self.filter_keyword.pack(side="left", padx=(4, 12))
        self.filter_keyword.bind("<KeyRelease>", lambda e: self._refresh())

        ctk.CTkButton(
            filter_frame, text="🔄 刷新", font=get_ctk_font("caption"),
            fg_color=COLORS["bg_primary"], text_color=COLORS["text_secondary"],
            corner_radius=8, height=30, width=70,
            command=self._refresh,
        ).pack(side="left")

        # ── Main area: list + preview ──
        main = ctk.CTkFrame(self, fg_color="transparent")
        main.pack(fill="both", expand=True)

        # Left: Tree list
        left_card = ctk.CTkFrame(main, fg_color=COLORS["bg_card"], corner_radius=12)
        left_card.pack(side="left", fill="both", expand=True, padx=(0, 6))

        columns = ("name", "source", "submodule", "updated")
        tree_frame = ctk.CTkFrame(left_card, fg_color="transparent")
        tree_frame.pack(fill="both", expand=True, padx=8, pady=6)
        self.tree = ttk.Treeview(tree_frame, columns=columns, show="headings", height=20)
        self.tree.heading("name", text="模板名称")
        self.tree.heading("source", text="来源文档")
        self.tree.heading("submodule", text="子模块")
        self.tree.heading("updated", text="更新日期")
        self.tree.column("name", width=180)
        self.tree.column("source", width=140)
        self.tree.column("submodule", width=120)
        self.tree.column("updated", width=100)
        self.tree.pack(side="left", fill="both", expand=True)
        self.tree.bind("<<TreeviewSelect>>", self._on_select)

        vsb = ttk.Scrollbar(tree_frame, orient="vertical", command=self.tree.yview)
        self.tree.configure(yscrollcommand=vsb.set)
        vsb.pack(side="right", fill="y")

        # Right: preview
        right_card = ctk.CTkFrame(main, fg_color=COLORS["bg_card"], corner_radius=12)
        right_card.pack(side="right", fill="both", expand=True, padx=(6, 0))

        ctk.CTkLabel(right_card, text="内容预览", font=get_ctk_font("section"),
                     text_color=COLORS["text_primary"]).pack(anchor="w", pady=(10, 5), padx=12)
        self.preview_text = ctk.CTkTextbox(right_card, wrap="word", state="disabled")
        self.preview_text.pack(fill="both", expand=True, padx=10, pady=(0, 10))

        # ── Bottom: action buttons ──
        btn_frame = ctk.CTkFrame(self, fg_color="transparent")
        btn_frame.pack(fill="x", pady=(8, 0))
        ctk.CTkButton(
            btn_frame, text="🗑 删除选中", font=get_ctk_font("body"),
            fg_color=COLORS["accent_red"], text_color="white",
            corner_radius=8, height=34,
            command=self._delete_selected,
        ).pack(side="left", padx=(0, 10))
        ctk.CTkButton(
            btn_frame, text="📥 导入文档", font=get_ctk_font("body"),
            fg_color=COLORS["accent_blue"], text_color="white",
            corner_radius=8, height=34,
            command=self._open_import_wizard,
        ).pack(side="left")

        self._load_submodule_names()

    # ── Helpers ──────────────────────────────────────────────────────────

    def _load_submodule_names(self) -> None:
        """从所有规则文件中提取子模块名称，填充筛选下拉框。"""
        names: set = set()
        for code in ["water_supply", "drainage", "road", "sanitation"]:
            rule = self.ctx.engine.load_rule(code)
            if not rule:
                continue
            for _stage_name, stage_data in rule.get("design_stages", {}).items():
                categories = (
                    stage_data.get("sections", stage_data)
                    if isinstance(stage_data, dict)
                    else {}
                )
                for _cat_key, cat_data in categories.items():
                    if isinstance(cat_data, dict) and "items" in cat_data:
                        for item in cat_data.get("items", []):
                            for sm in item.get("sub_modules", []):
                                n = sm.get("name", "")
                                if n:
                                    names.add(n)
        sorted_names = sorted(names)
        self.filter_submodule.configure(values=["全部"] + sorted_names)

    def refresh(self) -> None:
        """公开的刷新入口（供 MacOSApp 导航时调用）。"""
        self._refresh()

    def _refresh(self) -> None:
        """从 CustomTemplateManager 重新加载模板并刷新树形列表。"""
        for row in self.tree.get_children():
            self.tree.delete(row)
        self.preview_text.configure(state="normal")
        self.preview_text.delete("1.0", "end")
        self.preview_text.configure(state="disabled")

        try:
            from app.services.custom_template_manager import CustomTemplateManager

            ctm_dir = os.path.join(self.ctx.engine.project_root,
                                   "backend", "data", "custom_templates")
            mgr = CustomTemplateManager(ctm_dir)

            project_type = self.filter_type.get()
            if project_type != "全部":
                code_map = {
                    "给水工程": "water_supply",
                    "排水工程": "drainage",
                    "道路交通工程": "road",
                    "环境卫生工程": "sanitation",
                }
                pt: Optional[str] = code_map.get(project_type, "")
            else:
                pt = None

            sm_name = self.filter_submodule.get()
            if sm_name == "全部":
                sm_name = None

            keyword = self.filter_keyword.get().strip()
            if not keyword:
                keyword = None

            self._templates = mgr.search_templates(
                project_type=pt,
                sub_module_name=sm_name,
                keyword=keyword,
            )

            for tmpl in self._templates:
                name = tmpl.get("template_name", "")
                source = tmpl.get("source_doc", "")
                submodule = tmpl.get("sub_module_name", "")
                updated = (tmpl.get("updated_at", "") or "")[:10]
                self.tree.insert("", "end", values=(name, source, submodule, updated))
        except Exception as e:
            self.preview_text.configure(state="normal")
            self.preview_text.insert("1.0", f"加载模板库失败：{e}")
            self.preview_text.configure(state="disabled")

    def _on_select(self, _event: object = None) -> None:
        """在预览区域显示选中模板的详细信息。"""
        sel = self.tree.selection()
        if not sel:
            return
        idx = self.tree.index(sel[0])
        if idx >= len(self._templates):
            return
        tmpl = self._templates[idx]
        content = tmpl.get("content", "")
        source = tmpl.get("source_doc", "")
        name = tmpl.get("template_name", "")
        submodule = tmpl.get("sub_module_name", "")
        created = (tmpl.get("created_at", "") or "")[:19]
        updated = (tmpl.get("updated_at", "") or "")[:19]

        preview = (
            f"名称：{name}\n"
            f"来源文档：{source}\n"
            f"子模块：{submodule}\n"
            f"创建：{created}\n"
            f"更新：{updated}\n"
            f"\n{'-' * 40}\n\n{content}"
        )
        self.preview_text.configure(state="normal")
        self.preview_text.delete("1.0", "end")
        self.preview_text.insert("1.0", preview)
        self.preview_text.configure(state="disabled")

    def _delete_selected(self) -> None:
        """删除当前选中的模板变体。"""
        sel = self.tree.selection()
        if not sel:
            return
        idx = self.tree.index(sel[0])
        if idx >= len(self._templates):
            return
        tmpl = self._templates[idx]
        answer = messagebox.askyesno(
            "确认删除",
            f"确认删除模板「{tmpl.get('template_name', '')}」？"
        )
        if not answer:
            return

        try:
            from app.services.custom_template_manager import CustomTemplateManager

            ctm_dir = os.path.join(self.ctx.engine.project_root,
                                   "backend", "data", "custom_templates")
            mgr = CustomTemplateManager(ctm_dir)
            mgr.delete_template(tmpl["id"])
            self._refresh()
        except Exception as e:
            messagebox.showerror("删除失败", str(e))

    def _open_import_wizard(self) -> None:
        """打开模板导入向导对话框。由独立 task 实现。"""
        try:
            from import_wizard import ImportWizard  # noqa: F401 (T6)
            ImportWizard(self, self.ctx)
        except ImportError:
            messagebox.showinfo("提示", "导入向导功能即将上线，敬请期待。")
        except Exception as e:
            messagebox.showerror("错误", f"无法打开导入向导：{e}")


# ═══════════════════════════════════════════════════════════════════════════
# 8. MAIN APPLICATION WINDOW
# ═══════════════════════════════════════════════════════════════════════════

class MacOSApp(ctk.CTk):
    """macOS Sonoma-style main application window."""

    def __init__(self) -> None:
        super().__init__()

        # ── Appearance ──
        ctk.set_appearance_mode("light")
        ctk.set_default_color_theme("blue")

        # ── Window config ──
        self.title(APP_TITLE)
        self.geometry("1280x820")
        self.minsize(1050, 700)
        self.attributes("-alpha", 0.97)
        self.configure(fg_color=COLORS["bg_primary"])

        # Center on screen
        self.update_idletasks()
        sw = self.winfo_screenwidth()
        sh = self.winfo_screenheight()
        self.geometry(f"+{(sw - 1280) // 2}+{(sh - 820) // 2}")

        # ── Context ──
        self.ctx = EngineContext()
        self.status_var = tk.StringVar(value="就绪")
        self.health_var = tk.StringVar(value="⚪ 检查中...")

        # ── State ──
        self._current_page: Optional[str] = None
        self._pages: Dict[str, ctk.CTkFrame] = {}
        self._nav_buttons: Dict[str, ctk.CTkButton] = {}
        self._nav_selectors: Dict[str, ctk.CTkFrame] = {}
        self._fade_after_id: Optional[str] = None

        # ── Build ──
        self._build_layout()
        self._switch_page("project")
        self._check_health()

    # ── Layout ─────────────────────────────────────────────────────────────

    def _build_layout(self) -> None:
        """Full window layout: sidebar + content area + status bar."""
        # Container grid
        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(0, weight=1)

        # ── Content container ──
        self.content_frame = ctk.CTkFrame(self, fg_color="transparent")
        self.content_frame.grid(row=0, column=1, sticky="nsew", padx=(0, 0))

        # ── Pages (must exist before sidebar references them) ──
        self._pages["project"]  = ProjectPage(self.content_frame, self.ctx)
        self._pages["excel"]    = ExcelPage(self.content_frame, self.ctx)
        self._pages["generate"] = GeneratePage(
            self.content_frame, self.ctx,
            get_params_cb=lambda: (
                self._pages["project"].get_params()
                if "project" in self._pages
                else dict(DEFAULT_PARAMS)
            ),
        )
        self._pages["rules"]    = RulesPage(self.content_frame, self.ctx)
        self._pages["equipment"] = EquipmentPage(self.content_frame, self.ctx)
        self._pages["history"] = HistoricalTemplatePage(self.content_frame, self.ctx)
        self._pages["templates"] = TemplateLibraryPage(self.content_frame, self.ctx)
        self._pages["about"]    = AboutPage(self.content_frame, self.ctx)

        # ── Sidebar (safe now — pages exist) ──
        self._build_sidebar()

        # ── Status bar ──
        self._build_statusbar()

    def _build_sidebar(self) -> None:
        """220px sidebar with nav items."""
        sidebar = ctk.CTkFrame(
            self, fg_color=COLORS["bg_sidebar"],
            corner_radius=0, width=220,
        )
        sidebar.grid(row=0, column=0, sticky="ns")
        sidebar.grid_propagate(False)

        # App title
        title_area = ctk.CTkFrame(sidebar, fg_color="transparent")
        title_area.pack(fill="x", padx=16, pady=(20, 20))
        ctk.CTkLabel(title_area, text="⚡ 电气自控生成器",
                     font=get_ctk_font("sidebar_item"),
                     text_color=COLORS["text_primary"]).pack(anchor="w")
        ctk.CTkLabel(title_area, text=f"v{APP_VERSION}",
                     font=get_ctk_font("caption"),
                     text_color=COLORS["text_secondary"]).pack(anchor="w")

        # Separator
        ctk.CTkFrame(sidebar, fg_color=COLORS["separator"],
                     height=1).pack(fill="x", padx=12)

        # Nav items
        nav_items = [
            ("project",  "📋 项目信息"),
            ("excel",    "📊 Excel导入"),
            ("generate", "📄 生成文档"),
            ("rules",    "📚 知识库"),
            ("equipment","📦 设备材料表"),
            ("history",   "📚 历史模板库"),
            ("templates","📚 模板库"),
            ("about",    "ℹ️ 关于"),
        ]

        nav_container = ctk.CTkFrame(sidebar, fg_color="transparent")
        nav_container.pack(fill="x", padx=8, pady=(12, 0))

        for key, label in nav_items:
            # Row container for selector strip + button
            row = ctk.CTkFrame(nav_container, fg_color="transparent", height=38)
            row.pack(fill="x", pady=2)
            row.pack_propagate(False)

            # Selection indicator strip (hidden by default)
            sel_strip = ctk.CTkFrame(row, fg_color="transparent",
                                     width=3, corner_radius=0)
            sel_strip.place(x=0, rely=0, relheight=1.0)
            self._nav_selectors[key] = sel_strip

            btn = ctk.CTkButton(
                row, text=label,
                font=get_ctk_font("sidebar_item"),
                fg_color="transparent",
                hover_color=COLORS["bg_selected"],
                text_color=COLORS["text_primary"],
                anchor="w", corner_radius=8, height=34,
                command=lambda k=key: self._switch_page(k),
            )
            btn.pack(fill="both", padx=(4, 4), pady=0)
            self._nav_buttons[key] = btn

        # Bottom: project management + version
        ctk.CTkFrame(sidebar, fg_color=COLORS["separator"],
                     height=1).pack(fill="x", padx=12, pady=(10, 5))

        # Save/Load buttons
        proj_mgmt = ctk.CTkFrame(sidebar, fg_color="transparent")
        proj_mgmt.pack(fill="x", padx=8, pady=(0, 2))
        ctk.CTkButton(
            proj_mgmt, text="💾 保存项目", font=get_ctk_font("caption"),
            fg_color="transparent", text_color=COLORS["text_secondary"],
            hover_color=COLORS["bg_selected"], corner_radius=6,
            anchor="w", height=30,
            command=self._pages["about"]._save_project,
        ).pack(fill="x", pady=1)
        ctk.CTkButton(
            proj_mgmt, text="📂 加载项目", font=get_ctk_font("caption"),
            fg_color="transparent", text_color=COLORS["text_secondary"],
            hover_color=COLORS["bg_selected"], corner_radius=6,
            anchor="w", height=30,
            command=self._pages["about"]._load_project,
        ).pack(fill="x", pady=1)

        # Bottom spacer + status
        bottom = ctk.CTkFrame(sidebar, fg_color="transparent")
        bottom.pack(side="bottom", fill="x", padx=12, pady=(0, 15))
        self.sidebar_status = ctk.CTkLabel(
            bottom, textvariable=self.health_var, font=get_ctk_font("caption"),
            text_color=COLORS["text_secondary"],
        )
        self.sidebar_status.pack(anchor="w")

    def _build_statusbar(self) -> None:
        bar = ctk.CTkFrame(self, fg_color=COLORS["bg_sidebar"], height=28,
                           corner_radius=0)
        bar.grid(row=1, column=1, sticky="ew")
        bar.pack_propagate(False)
        ctk.CTkLabel(bar, textvariable=self.status_var, font=get_ctk_font("caption"),
                     text_color=COLORS["text_secondary"]).pack(
            side="left", padx=12, pady=2)

    # ── Navigation ─────────────────────────────────────────────────────────

    def _switch_page(self, page_name: str) -> None:
        """Switch content area to the target page with smooth transition."""
        if self._current_page == page_name:
            return

        # Cancel any in-progress fade
        if self._fade_after_id:
            self.after_cancel(self._fade_after_id)
            self._fade_after_id = None

        # Hide current page
        if self._current_page and self._current_page in self._pages:
            self._pages[self._current_page].pack_forget()

        # Update nav button styling and selection strips
        for key, btn in self._nav_buttons.items():
            if key == page_name:
                btn.configure(fg_color=COLORS["bg_selected"])
            else:
                btn.configure(fg_color="transparent")
        for key, strip in self._nav_selectors.items():
            strip.configure(
                fg_color=COLORS["accent_blue"] if key == page_name else "transparent"
            )

        # Show target page with fade-in effect
        page = self._pages.get(page_name)
        if page:
            # Start hidden (low opacity via fg_color trick using a cover frame)
            page.pack(fill="both", expand=True)
            # Force immediate render then trigger data refresh
            self.update_idletasks()
            self._current_page = page_name

        # Refresh data on certain pages (deferred for smoothness)
        if page_name == "rules":
            self._pages["rules"].refresh()
        elif page_name == "about":
            self._pages["about"].refresh()
        elif page_name == "history":
            self._pages["history"].refresh()
        elif page_name == "templates":
            self._pages["templates"].refresh()
        elif page_name == "generate":
            self._pages["generate"]._refresh_config()

        self.status_var.set(f"就绪 — {self._nav_labels().get(page_name, '')}")

    @staticmethod
    def _nav_labels() -> Dict[str, str]:
        return {
            "project":  "项目信息",
            "excel":    "Excel导入",
            "generate": "生成文档",
            "rules":    "知识库",
            "templates":"模板库",
            "about":    "关于",
        }

    # ── Health Check ───────────────────────────────────────────────────────

    def _check_health(self) -> None:
        status = self.ctx.health_status()
        if status["all_pass"]:
            self.health_var.set("✅ 系统正常")
        else:
            fails = [m for _, m in status["checks"] if "[FAIL]" in m]
            msg = f'⚠️ {fails[0][:20]}' if fails else "⚠️ 部分异常"
            self.health_var.set(msg)
        self.status_var.set("就绪")

    # ── Close ──────────────────────────────────────────────────────────────

    def _on_close(self) -> None:
        self.destroy()


# ═══════════════════════════════════════════════════════════════════════════
# 9. MAIN ENTRY POINT
# ═══════════════════════════════════════════════════════════════════════════

def main() -> None:
    """Launch the macOS-style desktop application."""
    app = MacOSApp()
    app.protocol("WM_DELETE_WINDOW", app._on_close)
    # Initial data loads after UI is ready
    app.after(300, lambda: app._pages["rules"].refresh())
    app.after(300, lambda: app._pages["about"].refresh())
    app.after(500, lambda: (
        app._pages["generate"]._refresh_config()
        if "generate" in app._pages else None
    ))
    app.mainloop()


if __name__ == "__main__":
    main()
