"""Extract nanchang.doc using Word COM - save to working dir."""
import os, sys, ctypes

src = r"D:\qoderwork\可研初步设计电气文本TESTB\source\nanchang.doc"
out_docx = r"D:\qoderwork\可研初步设计电气文本TESTB\source\nanchang_converted.docx"
out_md = r"D:\qoderwork\可研初步设计电气文本TESTB\output\extracted\南昌取水口优化调整可研报告2021.md"

ole32 = ctypes.OleDLL('ole32')
ole32.CoInitialize(None)

word = None
try:
    import comtypes.client
    word = comtypes.client.CreateObject("Word.Application")
    word.Visible = False
    word.DisplayAlerts = False

    abs_path = os.path.abspath(src)
    print(f"Opening: {src}")
    print(f"Size: {os.path.getsize(src)/1024/1024:.1f} MB")
    doc = word.Documents.Open(abs_path)

    # Save as DOCX
    out_docx_abs = os.path.abspath(out_docx)
    doc.SaveAs(out_docx_abs, FileFormat=16)
    doc.Close()
    print(f"Converted -> {out_docx}")
    print(f"Size: {os.path.getsize(out_docx)/1024/1024:.1f} MB")

    # Now extract via python-docx
    from docx import Document
    doc2 = Document(out_docx)
    lines = []
    for p in doc2.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        is_heading = False
        if p.style and p.style.name and p.style.name.startswith("Heading"):
            is_heading = True
        if not is_heading and p.runs:
            is_heading = any(r.bold for r in p.runs if r.bold)
        prefix = "## " if is_heading else ""
        lines.append(prefix + text)

    table_text = []
    for ti, table in enumerate(doc2.tables, 1):
        table_text.append(f"### 表格 {ti}")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            table_text.append(" | ".join(cells))

    out_text = "\n\n".join(lines)
    if table_text:
        out_text += "\n\n---\n\n" + "\n\n".join(table_text)

    with open(out_md, "w", encoding="utf-8") as f:
        f.write(out_text)

    print(f"\nOutput: {out_md}")
    print(f"Paragraphs: {len(lines)}, Tables: {len(doc2.tables)}")

except Exception as e:
    print(f"Error: {e}")
    import traceback
    traceback.print_exc()
finally:
    if word:
        try: word.Quit()
        except: pass
    ole32.CoUninitialize()
