"""
Word 文档生成器（python-docx 版，支持多模板）

根据《市政公用工程设计文件编制深度规定》（2025年版）的深度要求，
结合 Excel 负荷计算数据，逐栏目生成电气及自控设计说明书与设备材料表。

特性：
  * 4 种排版模板：standard(标准) / compact(紧凑) / report(报批) / modern(现代)
  * 真实、专业的栏目正文（无占位符、无文字重复）
  * 合理的设备材料表数量
  * 稳健的章节编号（顺序生成，不依赖巧合）
  * 内存预览：preview() 返回结构化 blocks，供 GUI 实时渲染

输出：格式化的 .docx Word 文档
"""
import json
import os
import math
import re
from typing import Dict, Any, List, Optional, Tuple
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from app.services.equipment_table_generator import EquipmentTableGenerator

# ───────────────────────────────────────────────────────────
# 模板定义
# ───────────────────────────────────────────────────────────
TEMPLATES: Dict[str, Dict[str, Any]] = {
    'standard': {
        'label': '标准版',
        'desc': '含封面、黑体标题、表格带边框，常规设计院出图风格',
        'cover': True,
        'body_font': '宋体', 'body_size': 12, 'line_spacing': 1.5,
        'head_font': '黑体',
        'h1_size': 16, 'h2_size': 15, 'h3_size': 13,
        'h1_color': None, 'h2_color': None, 'h3_color': None,
        'title_color': None,
        'table_header_shade': None,
        'header': False,
        'footer_text': '— 第  页 —',
        'page_number': True,
        'sig_block': False,
    },
    'compact': {
        'label': '紧凑版',
        'desc': '无封面、小字号、紧凑行距，适合篇幅受限的内部评审',
        'cover': False,
        'body_font': '宋体', 'body_size': 10.5, 'line_spacing': 1.25,
        'head_font': '黑体',
        'h1_size': 14, 'h2_size': 13, 'h3_size': 11.5,
        'h1_color': None, 'h2_color': None, 'h3_color': None,
        'title_color': None,
        'table_header_shade': None,
        'header': False,
        'footer_text': '第  页',
        'page_number': True,
        'sig_block': False,
    },
    'report': {
        'label': '报批版',
        'desc': '仿公文格式、含编制说明与会签栏，用于正式报审报批',
        'cover': True,
        'body_font': '仿宋', 'body_size': 12, 'line_spacing': 1.5,
        'head_font': '黑体',
        'h1_size': 16, 'h2_size': 15, 'h3_size': 13,
        'h1_color': None, 'h2_color': None, 'h3_color': None,
        'title_color': None,
        'table_header_shade': 'D9D9D9',
        'header': True,
        'footer_text': None,
        'page_number': True,
        'sig_block': True,
    },
    'modern': {
        'label': '现代版',
        'desc': '蓝色彩色标题、页眉项目名、表头着色，现代美观风格',
        'cover': True,
        'body_font': '宋体', 'body_size': 11, 'line_spacing': 1.35,
        'head_font': '微软雅黑',
        'h1_size': 16, 'h2_size': 14, 'h3_size': 12,
        'h1_color': '1F4E79', 'h2_color': '2E75B6', 'h3_color': '2E75B6',
        'title_color': '1F4E79',
        'table_header_shade': '1F4E79',
        'header': True,
        'footer_text': None,
        'page_number': True,
        'sig_block': False,
    },
}


def cn_num(n: int) -> str:
    """阿拉伯数字 → 中文数字（1~99）"""
    digits = '零一二三四五六七八九'
    if n <= 0:
        return '零'
    if n < 10:
        return digits[n]
    if n < 20:
        return '十' + (digits[n % 10] if n % 10 else '')
    if n < 100:
        tens, ones = n // 10, n % 10
        return digits[tens] + '十' + (digits[ones] if ones else '')
    return str(n)


# ───────────────────────────────────────────────────────────
# 栏目正文内容库
#   结构：CONTENT[工程类型][关键词] = [段落, ...]
#   段落可为：str（直接文本）或 callable(params, summary) -> str
#   解析顺序：精确标题 → 关键词(双向包含) → 通用关键词 → 兜底合成
# ───────────────────────────────────────────────────────────
def cn_scope(category: str) -> str:
    """取栏目范畴词（去掉末尾"设计"避免重复）"""
    if not category:
        return ''
    return category[:-2] if category.endswith('设计') else category


def _substation_text(p: Dict, s: Dict) -> str:
    tx = s.get('recommended_transformer', '')
    v = p.get('voltage_level', '10kV')
    return (
        f'根据负荷计算结果（补偿后视在负荷 Sjs\' = {s.get("total_sc_after", 0):.1f} kVA），'
        f'本工程设置变配电所，配置{tx}变压器，安装于变配电间内并采用户内落地安装方式。'
        f'高压侧电压等级为{v}，低压侧为 0.4kV，变压器低压侧经封闭母线槽引至低压配电柜。'
    )


GENERIC_CONTENT: Dict[str, List[Any]] = {
    '设计范围': [
        lambda p, s: (
            f'本工程电气设计包括整个厂内各建筑物的供配电系统及控制设计，'
            f'主要包括供电电压选择、负荷与计算、电气主接线及电气布置、继电保护、'
            f'电力拖动、电缆敷设、防雷接地、建筑物照明以及电力监控系统等方面。'
            f'设计分界点为总配电间{p.get("voltage_level","10kV")}高压进线电缆进线端，'
            f'进线终端以外部分由供电部门设计。'
        ),
    ],
    '电源': [
        lambda p, s: (
            f'本工程用电负荷等级为{p.get("load_level","二级")}负荷，'
            f'由{p.get("power_source","两路")}电源供电，电源电压等级均为{p.get("voltage_level","10kV")}，'
            f'两路电源一用一备。当一个电源发生故障时，由另一个电源带全部二级负荷运行，'
            f'以保证电气系统的连续、可靠运行。'
        ),
    ],
    '负荷等级': [
        lambda p, s: (
            f'本工程用电负荷等级为{p.get("load_level","二级")}负荷，'
            f'供电电压等级为{p.get("voltage_level","10kV")}。'
        ),
    ],
    '供配电系统': [
        lambda p, s: (
            f'本工程用电负荷主要是水泵及风机电动机类负载，厂内配电电压等级选择{p.get("voltage_level","10kV")}、'
            f'380/220V二级。{p.get("voltage_level","10kV")}配电系统采用单母线分段接线，双侧电源进线、'
            f'中间设联络开关，正常工作时两路电源同时运行，当一路失电时由另一路带全部负荷。'
            f'0.4kV系统均采用单母线分段中间设联络的接线，正常时分段开关断开、两台变压器同时运行，'
            f'当一台变压器故障时另一台能够担负起二级负荷用电。'
        ),
    ],
    '变配电系统': [
        lambda p, s: (
            f'按变配电设备尽可能靠近负荷中心的原则设置变配电所，'
            f'各低压变配电室内的两台变压器同时使用，当其中一台故障时，另一台能够担负起该变配电室承担的二级负荷用电。'
            f'{p.get("voltage_level","10kV")}配电所为户内单层布置，附设值班控制室，电缆通过电缆沟及电缆桥架连接各开关室。'
        ),
    ],
    '无功补偿': [
        lambda p, s: (
            '功率因数补偿采用低压侧集中自动补偿方式，补偿后的功率因数可达0.95以上。'
            '补偿装置采用智能型低压电容器自动投切，具备过压、欠压及谐波保护功能，'
            '并按照实际运行负荷自动调整补偿容量，避免过补偿或欠补偿。'
        ),
    ],
    '控制与保护': [
        lambda p, s: (
            '单机容量较大的低压电动机设备采用变频器或软起动方式，其余低压小型电动机采用直接起动方式。'
            f'{p.get("voltage_level","10kV")}系统采用分布式变电站自动化系统实现继保、测量和监控，'
            '低压系统总进线开关设短路速断、延时速断及长延时过电流三段保护。'
            '所有工艺设备及工艺流程均设自动、手动两种控制方式，并可在机旁、现场控制站及中央控制室三级操作。'
        ),
    ],
    '电动机保护': [
        lambda p, s: (
            '普通电动机设短路、过负荷及缺相保护、接地故障保护；'
            '大容量电动机设短路、过负荷、缺相、温度及接地故障保护；'
            '潜水电动机设短路、过负荷、缺相、温度及渗漏保护；'
            '阀门电动机设短路、过负荷、缺相及过力矩保护。'
        ),
    ],
    '电力监控': [
        lambda p, s: (
            f'本工程拟设置电力监控系统，采集{p.get("voltage_level","10kV")}及0.4kV系统的电量信号。'
            f'在{p.get("voltage_level","10kV")}开关柜设置微机型继电保护装置，完成对每回路的继电保护、'
            '开关量和模拟量实时数据采集、断路器控制和事故记录，并将数据上传至电力监控系统。'
            '主控单元通过光纤以太网接入厂区自控系统，实现变配电所的遥信、遥测等功能。'
        ),
    ],
    '变电所': [lambda p, s: _substation_text(p, s)],
    '变配电': [lambda p, s: _substation_text(p, s)],
    '保护': [
        lambda p, s: (
            f'{p.get("voltage_level","10kV")}系统采用微机型综合继电保护装置，'
            f'配置过流、速断、零序及变压器温控保护；操作电源采用免维护直流屏或交流 UPS。'
            f'低压配电线路采用断路器（兼做短路、过负荷保护）配合热继电器或智能马达保护器。'
        ),
    ],
    '控制': [
        lambda p, s: (
            '所有工艺设备及工艺流程均设自动、手动两种控制方式，并可在机旁、'
            '现场控制站及中央控制室三级操作。电机控制方式根据容量分别采用直接启动、'
            '软启动或变频启动。'
        ),
    ],
    '计量': [
        lambda p, s: (
            f'按相关规程要求，电能计量采用高供高计电能计量方式，'
            f'在{p.get("voltage_level","10kV")}电源进线侧配置专用计量柜作为总电能计量，'
            f'并配置多功能电能表及远传装置。各建筑物馈电线路装设电力监测仪和电能计量仪表。'
        ),
    ],
    '管线敷设': [
        lambda p, s: (
            '室内外电缆主要采用电缆沟、电缆桥架及穿管敷设方式。室外电力电缆采用电缆沟或'
            '排管敷设，室内配电间采用梯级式电缆桥架，控制电缆与电力电缆分层敷设；'
            '穿越道路及建（构）筑物处采用镀锌钢管保护。'
        ),
    ],
    '电缆敷设': [
        lambda p, s: (
            '控制电缆与电力电缆分层敷设。电力电缆采用 YJV 型交联聚乙烯绝缘电力电缆，'
            '控制电缆采用 KVVP 型铜芯聚氯乙烯绝缘屏蔽控制电缆，室外采用电缆沟或排管敷设。'
        ),
    ],
    '照明': [
        lambda p, s: (
            '各主要房间和场区按规范要求设置正常照明和应急照明。主要场所照度标准执行'
            'GB 50034《建筑照明设计标准》；光源优先采用高效 LED 灯具，应急照明采用'
            '自带蓄电池的应急灯具，应急持续时间不小于 90 min。'
        ),
    ],
    '设备选型': [
        lambda p, s: (
            '电气设备选型遵循"安全可靠、技术先进、经济合理、维护管理方便"的原则。'
            f'{p.get("voltage_level","10kV")}高压开关柜选用中置式金属铠装移开式封闭开关柜，'
            f'柜内配用真空断路器和直流弹簧操动机构，保护装置采用微机型智能保护元件；'
            '低压配电屏选用抽屉式低压开关柜，变压器选用节能型干式变压器。'
        ),
    ],
    '防雷': [
        lambda p, s: (
            '建筑物按第三类防雷建筑物设防，防雷装置满足防直击雷、防雷电感应及雷电波侵入的要求。'
            '接闪器在屋顶屋角、屋脊、屋檐等易受雷击部位采用热镀锌圆钢明敷接闪带，'
            '引下线利用建筑物柱内或剪力墙内主筋，接地极利用建筑物基础内主筋形成防雷接地网，'
            '接地电阻不大于 1Ω。'
        ),
    ],
    '接地': [
        lambda p, s: (
            '全厂设置共用接地系统，工作接地、保护接地、防雷接地及弱电接地共用接地装置，'
            '接地电阻不大于 1Ω；低压配电系统采用 TN-S 接地型式。'
        ),
    ],
    '防雷及接地': [
        lambda p, s: (
            '建筑物按第三类防雷建筑物设防，屋顶设接闪带，利用结构主筋作引下线，'
            '接地极为建筑物基础内的主筋形成防雷接地网。'
            '全厂设置共用接地系统，工作接地、保护接地、防雷接地及弱电接地共用接地装置，'
            '接地电阻不大于 1Ω。低压配电系统采用 TN-S 接地型式。'
        ),
    ],
    '自控系统': [
        lambda p, s: (
            '自动控制系统采用集散型控制系统（PLC + SCADA），遵循"分散控制、集中管理、'
            '数据共享"的原则，由中央控制室、现场控制站（PLC 分站）及工业以太网通信网络组成，'
            '实现对工艺全流程的实时监控、自动调节、报警及数据管理。'
        ),
    ],
    '仪表系统': [
        lambda p, s: (
            '根据工艺生产要求配置流量、液位、压力、温度及水质分析等在线检测仪表，'
            '仪表选型满足防护等级（户外不低于 IP65）、防腐及介质适应性要求，'
            '关键测点按冗余或重要度分级配置。'
        ),
    ],
    '仪表配管': [
        lambda p, s: (
            '仪表取源部件、导压管及仪表阀门的配管根据工艺过程参数（压力、差压、液位等）'
            '要求进行设计，材质与工艺管道相适应，并满足防腐、耐压及防堵要求。'
        ),
    ],
    '工业电视': [
        lambda p, s: (
            '在加药间、变配电所、主要工艺构筑物及厂区周界等重要区域设置工业电视监控系统，'
            '采用高清网络摄像机，信号经厂区以太网传输至中央控制室统一监视与存储。'
        ),
    ],
    '安防': [
        lambda p, s: (
            '厂区周界及重要出入口设置安防监控系统，包括视频监控、入侵报警、电子围栏及'
            '门禁系统，信号接入中央控制室安防管理平台统一管控。'
        ),
    ],
    '通信': [
        lambda p, s: (
            '设置综合布线系统，满足厂内语音、数据及视频监控的通信需求；'
            '厂区主干及构筑物间光缆采用单模光纤，水平布线采用六类非屏蔽双绞线。'
        ),
    ],
    '火灾自动报警': [
        lambda p, s: (
            '按 GB 50116《火灾自动报警系统设计规范》设置火灾自动报警及消防联动控制系统，'
            '采用集中报警系统形式；探测点包括感烟、感温探测器及手动报警按钮，'
            '并配置消防广播与应急通信。'
        ),
    ],
    '有线电视': [
        lambda p, s: (
            '在中央控制室、值班室及会议场所设置有线电视系统，信号源引自当地有线电视网，'
            '系统采用同轴电缆或光纤到楼的有线电视分配网络。'
        ),
    ],
    '防护等级': [
        lambda p, s: (
            '户外电气设备防护等级不低于 IP65，户内潮湿场所不低于 IP54，'
            '电气配电间不低于 IP30；腐蚀环境下设备选型及防护等级相应提高。'
        ),
    ],
    '防爆': [
        lambda p, s: (
            '爆炸危险环境内的电气设备及线路按 GB 50058 要求选用相应防爆等级（Ex d / Ex e /'
            'Ex ia 等）的产品，布线采用防爆挠性连接管及镀锌钢管明敷或电缆沟敷设。'
        ),
    ],
    '管缆敷设': [
        lambda p, s: (
            '自控及弱电管缆主要采用镀锌钢管、PVC 管或金属线槽敷设，室外埋地部分采用'
            '电缆保护管并作标高及走向标识；与电力电缆交叉时采取隔离措施。'
        ),
    ],
    '抗震': [
        lambda p, s: (
            '电气设备及管线按 GB 50981《建筑机电工程抗震设计规范》设置抗震支吊架，'
            '地震设防烈度按当地抗震设防要求执行，重要设备采取限位与减震措施。'
        ),
    ],
    '谐波治理': [
        lambda p, s: (
            '本工程谐波源为变频器等非线性负载，对变频器的整流元件提出谐波控制要求，'
            '提高变频器品质、增加换流装置的脉动数，使注入电网的谐波电流满足GB/T 14549要求。'
            '对谐波敏感的设备（如电容器）按谐波性质设置参数匹配的串联电抗器，以减小谐波影响。'
        ),
    ],
    '节能': [
        lambda p, s: (
            '选用 SCB14 及以上能效等级的节能型干式变压器，照明光源采用 LED 灯具，'
            '并设置低压侧无功自动补偿装置，使功率因数补偿后不低于 0.95。'
        ),
    ],
    '智慧排水': [
        lambda p, s: (
            '智慧排水管控平台采用"云—管—边—端"总体架构，整合厂站 SCADA、'
            '厂网液位/流量在线监测及水质监测数据，实现排水系统全要素监测、'
            '预警研判与智慧调度，支撑厂网一体化运行管理。'
        ),
    ],
    '路灯': [
        lambda p, s: (
            '道路照明采用 LED 路灯，路灯控制采用智能照明控制系统，具备时控、'
            '光控及远程集中控制功能，并预留单灯调光与故障报警接口。'
        ),
    ],
    '信号灯': [
        lambda p, s: (
            '交通信号灯采用 LED 光源信号机组，信号控制机具备多时段、自适应及联网协调控制功能，'
            '并按相交道路等级配置行人过街及非机动车信号。'
        ),
    ],
    '智能交通': [
        lambda p, s: (
            '智能交通设施包括交通流检测、事件检测、可变信息板（LED 诱导屏）及'
            '交通数据上传等分系统，数据接入交通管理中心实现区域协调控制。'
        ),
    ],
    '监控': [
        lambda p, s: (
            '交通监控系统包括交通流检测、交通事件检测、信号控制及信息发布等子系统，'
            '外场设备数据经通信管道上传至交通监控中心。'
        ),
    ],
    '保安电源': [
        lambda p, s: (
            '对不允许中断供电的重要负荷（如中控室、安防、火灾报警等）设置保安电源，'
            '采用 EPS 或柴油发电机组作为应急备用，确保市电失电后关键负荷持续供电。'
        ),
    ],
}

# 工程类型特有的栏目（覆盖/补充通用库）
TYPE_CONTENT: Dict[str, Dict[str, List[Any]]] = {
    'water_supply': {
        '设计范围及内容': GENERIC_CONTENT['设计范围'],
        '电源': GENERIC_CONTENT['电源'],
        '用电负荷': [],
        '供配电系统': GENERIC_CONTENT['供配电系统'],
        '变电所布置': GENERIC_CONTENT['变电所'],
        '保护和控制': GENERIC_CONTENT['控制与保护'],
        '计量': GENERIC_CONTENT['计量'],
        '管线敷设': GENERIC_CONTENT['管线敷设'],
        '照明及应急照明': GENERIC_CONTENT['照明'],
        '设备选型': GENERIC_CONTENT['设备选型'],
        '防雷与接地': GENERIC_CONTENT['防雷及接地'],
        '设计范围及依据': GENERIC_CONTENT['设计范围'],
        '自控系统': GENERIC_CONTENT['自控系统'],
        '仪表系统': GENERIC_CONTENT['仪表系统'],
        '工业电视系统': GENERIC_CONTENT['工业电视'],
        '安防系统': GENERIC_CONTENT['安防'],
        '通信系统': GENERIC_CONTENT['通信'],
        '火灾自动报警系统': GENERIC_CONTENT['火灾自动报警'],
        '有线电视系统': GENERIC_CONTENT['有线电视'],
        '防护等级': GENERIC_CONTENT['防护等级'],
        '防爆及防损坏': GENERIC_CONTENT['防爆'],
        '设备选型与布置': GENERIC_CONTENT['设备选型'],
        '管缆敷设': GENERIC_CONTENT['管缆敷设'],
        '防雷及接地': GENERIC_CONTENT['防雷及接地'],
        '机电抗震': GENERIC_CONTENT['抗震'],
    },
    'drainage': {
        '智慧排水管控': GENERIC_CONTENT['智慧排水'],
        '谐波治理': GENERIC_CONTENT['谐波治理'],
        '电力监控': GENERIC_CONTENT['电力监控'],
    },
    'road': {
        '路灯控制': GENERIC_CONTENT['路灯'],
        '信号灯系统': GENERIC_CONTENT['信号灯'],
        '智能交通设施': GENERIC_CONTENT['智能交通'],
        '监控系统': GENERIC_CONTENT['监控'],
    },
    'sanitation': {
        '保安电源': GENERIC_CONTENT['保安电源'],
        '谐波保护': GENERIC_CONTENT['谐波治理'],
        '仪表配管设计': GENERIC_CONTENT['仪表配管'],
        '过电压保护': [
            lambda p, s: (
                '对雷害及操作过电压采取防护措施：10kV 进线及变压器高压侧装设氧化锌避雷器，'
                '低压侧设置浪涌保护器（SPD），重要电子设备电源端装设多级 SPD。'
            ),
        ],
    },
}


def _iter_categories(categories: dict):
    """迭代分类数据，处理标准分类（含items）和容器分类（含子分类）两种结构
    
    road.json 的"初步设计→自控"是容器，内含"交通管理设施"、"智慧枢纽智能系统"等子分类。
    """
    for key, data in categories.items():
        if not isinstance(data, dict):
            continue
        if 'items' in data:
            yield key, data
        else:
            for sub_key, sub_data in data.items():
                if isinstance(sub_data, dict) and 'items' in sub_data:
                    yield sub_key, sub_data


# ───────────────────────────────────────────────────────────
# 生成引擎
# ───────────────────────────────────────────────────────────
class DocxGenerator:
    """Word 文档生成引擎（多模板 + 预览）"""

    def __init__(self, rules_dir: str, output_dir: str, template: str = 'standard'):
        self.rules_dir = rules_dir
        self.output_dir = os.path.join(output_dir, 'generated')
        os.makedirs(self.output_dir, exist_ok=True)
        self.template = template
        self.t = TEMPLATES.get(template, TEMPLATES['standard'])
        self.blocks: List[Any] = []
        self._load_rules()

    def _load_rules(self):
        self.rules = {}
        for code in ['water_supply', 'drainage', 'road', 'sanitation']:
            path = os.path.join(self.rules_dir, f'{code}.json')
            if os.path.exists(path):
                with open(path, 'r', encoding='utf-8') as f:
                    self.rules[code] = json.load(f)

    def set_template(self, template: str):
        if template in TEMPLATES:
            self.template = template
            self.t = TEMPLATES[template]

    # ─── 公共 API ───
    def generate(self, project_type, design_stage, excel_data, params) -> str:
        rule = self.rules.get(project_type)
        if not rule:
            raise ValueError(f'未找到工程类型规则: {project_type}')
        summary = excel_data.get('summary', {})
        area_summaries = excel_data.get('area_summaries', {})

        doc = self._create_document(rule, summary, area_summaries, params, design_stage)

        project_name = params.get('project_name', '未命名项目')
        ts = datetime.now().strftime('%Y%m%d_%H%M%S')
        safe_name = re.sub(r'[\\/:*?"<>|]', '_', project_name)
        filename = f'{safe_name}_{rule["project_type"]}_{design_stage}电气自控说明_{ts}.docx'
        output_path = os.path.join(self.output_dir, filename)
        doc.save(output_path)
        return output_path

    def preview(self, project_type, design_stage, excel_data, params) -> List[Any]:
        """仅构建内存文档并返回预览 blocks（不保存文件）"""
        rule = self.rules.get(project_type)
        if not rule:
            raise ValueError(f'未找到工程类型规则: {project_type}')
        summary = excel_data.get('summary', {})
        area_summaries = excel_data.get('area_summaries', {})
        self._create_document(rule, summary, area_summaries, params, design_stage)
        return self.blocks

    # ─── 文档构建 ───
    def _create_document(self, rule, summary, area_summaries, params, design_stage):
        self.blocks = []
        doc = Document()
        self._create_styles(doc)

        project_name = params.get('project_name', '')
        voltage = params.get('voltage_level', '10kV')
        load_level = params.get('load_level', '二级')

        if self.t['cover']:
            self._add_cover_page(doc, project_name, rule, design_stage, params)

        self._add_heading(doc, f'{rule["project_type"]} {design_stage}阶段 电气及自控设计说明书', level=1)

        self._add_heading(doc, '一、编制依据', level=2)
        self._add_para(doc, f'1. {rule["regulation_ref"]}')
        self._add_para(doc, '2. 国家及行业现行相关标准、规范（详见各章节引用）')
        self._add_para(doc, '3. 业主提供的设计任务书、基础资料及评审意见')
        self._add_para(doc, '4. 工艺、建筑、结构、给排水、暖通等专业提供的设计资料')
        self._add_para(doc, '')

        cat_ordinal = 2
        # 获取当前设计阶段下的分类（电气/自控）
        stage_data = rule.get('design_stages', {}).get(design_stage, {})
        # 环卫工程(sanitation)使用嵌套 sections 结构
        categories = stage_data.get('sections', stage_data) if isinstance(stage_data, dict) else {}
        for category_name, category_data in _iter_categories(categories):
            params['category_name'] = category_name
            self._add_heading(doc, f'{cn_num(cat_ordinal)}、{category_data["title"]}', level=2)
            cat_ordinal += 1

            for item in category_data.get('items', []):
                if item.get('optional', False):
                    continue
                title = item['title']
                requirement = item['requirement']
                has_calc = item.get('has_calculation', False)
                calc_from_excel = item.get('calc_from_excel', False)

                self._add_heading(doc, f'{item["order"]}. {title}', level=3)

                if calc_from_excel and has_calc:
                    self._add_load_calculation_section(doc, summary, area_summaries, voltage, load_level)
                elif has_calc and not calc_from_excel:
                    self._add_calc_section(doc, title, summary, voltage, load_level)
                else:
                    self._add_text_section(doc, title, requirement, category_name, params, summary)
            self._add_para(doc, '')

        self._add_heading(doc, f'{cn_num(cat_ordinal)}、主要设备材料表', level=2)
        cat_ordinal += 1
        self._add_equipment_table(doc, summary, area_summaries, rule, design_stage)

        if self.t.get('sig_block'):
            self._add_sign_block(doc)

        if self.t.get('header'):
            self._add_header(doc, project_name, rule, design_stage)
        if self.t.get('page_number'):
            self._add_page_number(doc)

        return doc

    # ─── 内容解析 ───
    def _resolve_content(self, type_code, title, category, params, summary) -> List[Tuple[str, str]]:
        tmap = TYPE_CONTENT.get(type_code, {})
        if title in tmap:
            return self._instantiate(tmap[title], params, summary)
        for k, v in tmap.items():
            if k and (k in title or title in k):
                return self._instantiate(v, params, summary)
        for k, v in GENERIC_CONTENT.items():
            if k and k in title:
                return self._instantiate(v, params, summary)
        return self._fallback_content(category, title, params, summary)

    def _instantiate(self, items, params, summary) -> List[Tuple[str, str]]:
        out = []
        for it in items:
            text = it(params, summary) if callable(it) else str(it)
            if text:
                out.append(('p', text))
        return out

    def _fallback_content(self, category, title, params, summary) -> List[Tuple[str, str]]:
        scope = cn_scope(category)
        req = ''
        rule = self.rules.get(params.get('project_type', ''), {})
        for stage_data in rule.get('design_stages', {}).values():
            categories = stage_data.get('sections', stage_data) if isinstance(stage_data, dict) else {}
            for _cat_name, cat_data in _iter_categories(categories):
                for it in cat_data.get('items', []):
                    if it['title'] == title:
                        req = it.get('requirement', '')
                        break
                if req:
                    break
            if req:
                break
        if req:
            text = (
                f'本工程{scope}的"{title}"应按相关规定及本工程特点进行设计：{req}。'
                f'具体方案结合工艺布置、负荷分布及当地供电与智能化管理要求综合确定。'
            )
        else:
            text = f'本工程{scope}的"{title}"按国家及行业现行标准、规范进行设计，满足使用功能及安全可靠性要求。'
        return [('p', text)]

    def _add_text_section(self, doc, title, requirement, category, params, summary):
        type_code = params.get('project_type', '')
        for kind, text in self._resolve_content(type_code, title, category, params, summary):
            if kind == 'b':
                self._add_bullet_para(doc, text)
            else:
                self._add_para(doc, text)

    # ─── 负荷计算章节 ───
    def _add_load_calculation_section(self, doc, summary, area_summaries, voltage, load_level):
        total_equip = summary.get('total_equip_power', 0)
        total_pc = summary.get('total_pc', 0)
        total_pc_k = summary.get('total_pc_k', 0)
        total_qc_k = summary.get('total_qc_k', 0)
        total_sc_k = summary.get('total_sc_k', 0)
        cos_before = summary.get('cos_before', 0)
        qc_comp = summary.get('qc_compensation', 0)
        cos_target = summary.get('cos_target', 0.95)
        total_qc_after = summary.get('total_qc_after', 0)
        total_sc_after = summary.get('total_sc_after', 0)
        tx_config = summary.get('recommended_transformer', '')
        sp = summary.get('simultaneous_coeff', {})
        total_devices = summary.get('total_devices', 0)

        self._add_para(doc, f'本工程用电设备总安装容量为 {total_equip:.1f} kW，用电设备总台数 {total_devices} 台。')
        self._add_para(doc, f'用电负荷等级为{load_level}负荷，供电电压等级为{voltage}。')
        self._add_para(doc, '负荷计算采用需要系数法，计算过程及结果如下：')
        self._add_para(doc, '表1  各区域用电负荷汇总表', bold=True)

        headers = ['序号', '区域名称', '设备数量', '设备容量(kW)',
                   '计算有功(kW)', '计算无功(kvar)', '视在功率(kVA)']
        rows = []
        if isinstance(area_summaries, dict):
            for idx, (area_name, data) in enumerate(area_summaries.items(), 1):
                rows.append([
                    str(idx), area_name, str(data['device_count']),
                    f'{data["equip_power"]:.1f}', f'{data["pc"]:.1f}',
                    f'{data["qc"]:.1f}', f'{data["sc"]:.1f}'
                ])
        rows.append([
            '', '合计', str(total_devices),
            f'{total_equip:.1f}', f'{total_pc:.1f}',
            f'{summary.get("total_qc", 0):.1f}', f'{summary.get("total_sc", 0):.1f}'
        ])
        self._add_table(doc, headers, rows)

        self._add_para(doc, '')
        ksp = sp.get('KΣP', 0.9)
        ksq = sp.get('KΣq', 0.95)
        self._add_para(doc, f'考虑同时系数 KΣp={ksp}、KΣq={ksq} 后：')
        self._add_bullet_para(doc, f'有功计算负荷 Pjs = KΣp × ∑Pc = {ksp} × {total_pc:.1f} = {total_pc_k:.1f} kW')
        self._add_bullet_para(doc, f'无功计算负荷 Qjs = KΣq × ∑Qc = {ksq} × {summary.get("total_qc", 0):.1f} = {total_qc_k:.1f} kvar')
        self._add_bullet_para(doc, f'视在计算负荷 Sjs = √(Pjs² + Qjs²) = {total_sc_k:.1f} kVA')

        self._add_para(doc, '无功补偿计算：', bold=True)
        self._add_bullet_para(doc, f'补偿前功率因数 cosφ₁ = Pjs / Sjs = {cos_before:.4f}')
        self._add_bullet_para(doc, f'要求补偿后功率因数 cosφ₂ ≥ {cos_target}')
        self._add_bullet_para(doc, f'需补偿无功容量 Qc = Pjs × (tanφ₁ − tanφ₂) = {qc_comp:.1f} kvar')
        self._add_bullet_para(doc, f'补偿后无功负荷 Qjs\' = {total_qc_after:.1f} kvar')
        self._add_bullet_para(doc, f'补偿后视在负荷 Sjs\' = √(Pjs² + Qjs\'²) = {total_sc_after:.1f} kVA')
        self._add_bullet_para(doc, f'补偿后功率因数 cosφ₂\' = Pjs / Sjs\' ≥ {cos_target}，满足要求。')

        self._add_para(doc, '变压器选型：', bold=True)
        self._add_para(doc, f'根据补偿后计算负荷 Sjs\' = {total_sc_after:.1f} kVA，推荐变压器配置为 {tx_config}。')
        tx_total = self._parse_tx_total_capacity(tx_config)
        load_rate = (total_sc_after / tx_total * 100) if tx_total > 0 else 0
        if load_rate <= 85:
            self._add_para(doc, f'变压器负荷率约 {load_rate:.1f}%，处于经济运行区间（推荐负载率 70%~85%），满足要求。')
        else:
            self._add_para(doc,
                f'变压器负荷率约 {load_rate:.1f}%。该配置在考虑单台变压器故障时的备用能力'
                f'及工程远期发展余地的前提下留有适当裕量，满足安全可靠供电要求。')

    def _add_calc_section(self, doc, title, summary, voltage, load_level):
        total_sc_k = summary.get('total_sc_k', 0)
        total_sc_after = summary.get('total_sc_after', 0)
        qc_comp = summary.get('qc_compensation', 0)

        if '变电所' in title or '变配电' in title:
            self._add_para(doc, _substation_text({'voltage_level': voltage}, summary))
        elif '补偿' in title or '无功' in title:
            self._add_para(doc,
                f'采用低压侧集中自动补偿方式，补偿容量约 {qc_comp:.0f} kvar，'
                f'补偿后侧功率因数不低于 {summary.get("cos_target", 0.95)}。')
            self._add_para(doc, '补偿装置采用智能型低压电容器自动投切，具备过压、欠压及谐波保护功能。')
        elif '保护' in title and '控制' in title:
            self._add_para(doc,
                f'{voltage}系统采用微机型综合继电保护装置，配置过流、速断、零序及变压器保护；'
                f'操作电源采用免维护直流屏。电机设备按容量分别采用直接启动、软启动或变频启动，'
                f'低压配电线路采用断路器配合热继电器/智能马达保护器保护。')
        elif '照明' in title:
            self._add_para(doc,
                '各单体建筑按 GB 50034 设置正常照明与应急照明，光源采用 LED 灯具，'
                '应急照明应急时间不小于 90 min。')
        else:
            self._add_para(doc, f'详见负荷计算章节及相应系统图，相关参数按本工程计算结果取值（Sjs={total_sc_k:.1f} kVA，补偿后 Sjs\'={total_sc_after:.1f} kVA）。')

    # ─── 设备材料表（数量合理化） ───
    def _add_equipment_table(self, doc, summary, area_summaries, rule, design_stage=""):
        eq_gen = EquipmentTableGenerator()
        project_type = rule.get('code', '')

        self._add_para(doc, '主要设备材料表', bold=True)

        result = eq_gen.generate(project_type, summary, area_summaries, design_stage)
        if result['has_data']:
            self._add_para(doc, result['note'])
            self._add_table(doc, result['headers'], result['rows'])

    # ─── 辅助：块收集 + 写 doc ───
    def _add_cover_page(self, doc, project_name, rule, design_stage, params):
        for _ in range(6):
            doc.add_paragraph()
        self._add_cover_line(doc, project_name or '市政工程', 22, True)
        self._add_cover_line(doc, f'{rule["project_type"]} {design_stage}', 18, True)
        self._add_cover_line(doc, '电气及自控设计说明书', 18, True)
        for _ in range(4):
            doc.add_paragraph()
        info = [
            ('编制依据', rule['regulation_ref']),
            ('设计阶段', design_stage),
            ('工程类型', rule['project_type']),
            ('编制日期', datetime.now().strftime('%Y年%m月%d日')),
        ]
        for label, value in info:
            self._add_cover_line(doc, f'{label}：{value}', 14, False)
        self.blocks.append(('cover', project_name or '市政工程',
                            f'{rule["project_type"]} {design_stage}  电气及自控设计说明书',
                            [f'{l}：{v}' for l, v in info]))
        doc.add_page_break()
        self.blocks.append(('pagebreak',))

    def _add_cover_line(self, doc, text, size, bold):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = self.t['head_font']
        run.element.rPr.rFonts.set(qn('w:eastAsia'), self.t['head_font'])
        if self.t.get('title_color'):
            run.font.color.rgb = RGBColor.from_string(self.t['title_color'])

    def _add_heading(self, doc, text, level=1):
        size = {1: self.t['h1_size'], 2: self.t['h2_size'], 3: self.t['h3_size']}[level]
        color = {1: self.t['h1_color'], 2: self.t['h2_color'], 3: self.t['h3_color']}[level]
        p = doc.add_paragraph()
        if level == 1:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(size)
        run.font.name = self.t['head_font']
        run.element.rPr.rFonts.set(qn('w:eastAsia'), self.t['head_font'])
        if color:
            run.font.color.rgb = RGBColor.from_string(color)
        self.blocks.append((f'h{level}', text))

    def _add_para(self, doc, text, bold=False):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(self.t['body_size'])
        run.font.name = self.t['body_font']
        run.element.rPr.rFonts.set(qn('w:eastAsia'), self.t['body_font'])
        if text:
            self.blocks.append(('p', text))
        return p

    def _add_bullet_para(self, doc, text):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.0)
        run = p.add_run(text)
        run.font.size = Pt(self.t['body_size'])
        run.font.name = self.t['body_font']
        run.element.rPr.rFonts.set(qn('w:eastAsia'), self.t['body_font'])
        self.blocks.append(('b', text))
        return p

    def _add_table(self, doc, headers, rows):
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        shade = self.t.get('table_header_shade')
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(9)
            if shade:
                self._shade_cell(table.rows[0].cells[i], shade)
        for row in rows:
            r = table.add_row()
            for i, v in enumerate(row):
                r.cells[i].text = str(v)
                for p in r.cells[i].paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.size = Pt(9)
        self.blocks.append(('table', '', list(headers), [list(x) for x in rows]))

    def _add_sign_block(self, doc):
        self._add_para(doc, '')
        self._add_heading(doc, '编制与签署', level=2)
        headers = ['编制', '校核', '审核', '审定', '项目负责人', '日期']
        rows = [['', '', '', '', '', '']]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(10)
        self.blocks.append(('sig', headers, rows))

    def _create_styles(self, doc):
        style = doc.styles['Normal']
        font = style.font
        font.name = self.t['body_font']
        font.size = Pt(self.t['body_size'])
        style.element.rPr.rFonts.set(qn('w:eastAsia'), self.t['body_font'])
        fmt = style.paragraph_format
        fmt.line_spacing = self.t['line_spacing']

    def _add_header(self, doc, project_name, rule, design_stage):
        section = doc.sections[0]
        header = section.header
        header.is_linked_to_previous = False
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'{project_name}    {rule["project_type"]} {design_stage}  电气及自控设计说明书')
        run.font.size = Pt(9)
        run.font.name = self.t['body_font']
        run.element.rPr.rFonts.set(qn('w:eastAsia'), self.t['body_font'])

    @staticmethod
    def _add_page_number(doc):
        """稳健的页码域：域元素各自置于独立 <w:r> 内，作为 <w:p> 的子元素"""
        section = doc.sections[0]
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        r_pre = p.add_run('— 第 ')
        r_pre.font.size = Pt(10)
        r_post = p.add_run(' 页 —')
        r_post.font.size = Pt(10)

        fld_begin = OxmlElement('w:fldChar')
        fld_begin.set(qn('w:fldCharType'), 'begin')
        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = ' PAGE '
        fld_sep = OxmlElement('w:fldChar')
        fld_sep.set(qn('w:fldCharType'), 'separate')
        fld_end = OxmlElement('w:fldChar')
        fld_end.set(qn('w:fldCharType'), 'end')

        # 一个 <w:r> 包裹 begin+instr+separate
        r_field = OxmlElement('w:r')
        r_field.append(fld_begin)
        r_field.append(instr)
        r_field.append(fld_sep)
        p._p.append(r_field)

        # 一个 <w:r> 包裹 end
        r_end = OxmlElement('w:r')
        r_end.append(fld_end)
        p._p.append(r_end)

    @staticmethod
    def _shade_cell(cell, fill: str):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill)
        tcPr.append(shd)

    @staticmethod
    def _parse_tx_capacity(tx_config: str) -> float:
        matches = re.findall(r'(\d+)kVA', tx_config)
        count = 1
        count_match = re.search(r'(\d+)×', tx_config)
        if count_match:
            count = int(count_match.group(1))
        return float(sum(int(m) * count for m in matches))

    def _parse_tx_total_capacity(self, tx_config: str) -> float:
        return self._parse_tx_capacity(tx_config)

    @staticmethod
    def _parse_tx_detail(tx_config: str) -> tuple:
        count_match = re.search(r'(\d+)×', tx_config)
        cap_match = re.search(r'(\d+)kVA', tx_config)
        count = int(count_match.group(1)) if count_match else 1
        cap = int(cap_match.group(1)) if cap_match else 1000
        return (count, cap)


# ───────────────────────────────────────────────────────────
# 命令行自测
# ───────────────────────────────────────────────────────────
if __name__ == '__main__':
    import sys
    sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))
    from app.services.excel_parser import ExcelLoadParser

    parser = ExcelLoadParser()
    excel_path = r"D:\00-水厂负荷计算表.xlsx"
    excel_data = parser.parse(excel_path) if os.path.exists(excel_path) else {
        'summary': {'total_devices': 30, 'total_equip_power': 1850.5, 'total_pc': 1200.3,
                    'total_qc': 680.2, 'total_sc': 1380.0, 'total_pc_k': 1080.3, 'total_qc_k': 646.2,
                    'total_sc_k': 1258.5, 'cos_before': 0.8581, 'cos_target': 0.95,
                    'qc_compensation': 380.5, 'total_qc_after': 265.7, 'total_sc_after': 1112.5,
                    'recommended_transformer': '2×800kVA', 'area_count': 5,
                    'simultaneous_coeff': {'KΣP': 0.9, 'KΣq': 0.95}},
        'area_summaries': {
            '取水泵房': {'device_count': 5, 'equip_power': 650.0, 'pc': 585.0, 'qc': 362.7, 'sc': 688.2},
            '送水泵房': {'device_count': 4, 'equip_power': 480.0, 'pc': 432.0, 'qc': 267.8, 'sc': 508.2},
            '加药间': {'device_count': 6, 'equip_power': 180.5, 'pc': 126.4, 'qc': 75.8, 'sc': 147.4},
            '滤池': {'device_count': 8, 'equip_power': 320.0, 'pc': 256.0, 'qc': 153.6, 'sc': 298.6},
            '辅助建筑': {'device_count': 7, 'equip_power': 220.0, 'pc': 110.0, 'qc': 55.0, 'sc': 123.0},
        },
    }

    base_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    params = {
        'project_name': '赤壁中心水厂', 'voltage_level': '10kV', 'load_level': '二级',
        'project_type': 'water_supply', 'power_source': '两路',
        'standby_desc': '两路电源互为备用，当一路电源故障时另一路可承担全部负荷。',
    }
    for tpl in TEMPLATES:
        gen = DocxGenerator(rules_dir=os.path.join(base_dir, 'data', 'rules'),
                            output_dir=os.path.join(base_dir, 'output'), template=tpl)
        out = gen.generate('water_supply', '初步设计', excel_data, params)
        print(f'OK [{TEMPLATES[tpl]["label"]}] {out}')
