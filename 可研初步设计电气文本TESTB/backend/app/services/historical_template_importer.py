"""
历史模板导入器 - 文件导入与自动分类服务

管理用户导入的历史设计文档（.docx, .txt, .md），
自动按工程类型/设计阶段/专业分类，并存储为独立 JSON 文件。

特性：
  * 支持 .docx / .txt / .md 三种格式
  * 基于关键词的自动分类（工程类型、设计阶段、专业、条目）
  * 文档按标题自动拆分为多个模板记录
  * 扁平 JSON 文件存储 + index.json 快速索引
  * 支持搜索、过滤、备注编辑
"""
from __future__ import annotations

import json
import logging
import os
import re
import uuid
from datetime import datetime
from typing import Dict, List, Optional

logger = logging.getLogger(__name__)

# ── 分类关键词映射 ─────────────────────────────────────────────────────────

_PROJECT_TYPE_KEYWORDS: Dict[str, List[str]] = {
    "water_supply": ["给水", "供水", "自来水", "净水", "水厂", "取水泵站", "送水泵站"],
    "drainage": ["排水", "污水", "处理厂", "污水处理", "雨水", "泵站", "管网", "废水处理"],
    "road": ["道路", "公路", "市政路", "城市道路", "交通", "桥梁", "隧道", "立交"],
    "sanitation": ["环卫", "垃圾", "焚烧", "填埋", "固废", "废物", "垃圾处理"],
}

_DESIGN_STAGE_KEYWORDS: Dict[str, List[str]] = {
    "可行性研究": ["可研", "可行性", "可行性研究", "项目建议书"],
    "初步设计": ["初设", "初步设计", "初步"],
}

_CATEGORY_KEYWORDS: Dict[str, List[str]] = {
    "电气": ["电气", "供配电", "供电", "照明", "防雷", "接地", "配电", "变压器",
             "开关柜", "电缆", "负荷", "用电", "变电", "电压"],
    "自控": ["自控", "仪表", "监控", "通信", "安防", "自动化", "控制", "PLC",
             "SCADA", "弱电", "信号", "传感器", "DCS"],
}

# ── 标题拆分正则 ───────────────────────────────────────────────────────────

_TXT_HEADING_RE = re.compile(
    r"^(?:第[一二三四五六七八九十百零\d]+[章节篇部]"
    r"|\d+[\.\)、]"
    r"|[一二三四五六七八九十]+[、\.]"
    r"|\(\d+\)"
    r"|\（\d+\）)",
    re.MULTILINE,
)

_MD_HEADING_RE = re.compile(r"^#{1,4}\s+.+$", re.MULTILINE)

_DOCX_HEADING_STYLES = {
    "Heading 1", "Heading 2", "Heading 3", "Heading 4",
    "标题 1", "标题 2", "标题 3", "标题 4",
}


class HistoricalTemplateImporter:
    """历史模板导入器 — 文件导入、自动分类、存储管理"""

    def __init__(self, data_dir: Optional[str] = None):
        """初始化导入器。

        Args:
            data_dir: 模板存储根目录。默认 backend/data/historical_templates/
        """
        if data_dir is None:
            data_dir = os.path.join(
                os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))),
                "data", "historical_templates",
            )
        self.data_dir = data_dir
        os.makedirs(self.data_dir, exist_ok=True)

        self._rules_dir = os.path.join(
            os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))),
            "data", "rules",
        )
        self._item_keywords: Optional[Dict[str, List[str]]] = None

    # ── 内部辅助 ──────────────────────────────────────────────────────

    @staticmethod
    def _generate_id() -> str:
        return str(uuid.uuid4())

    @staticmethod
    def _now_iso() -> str:
        return datetime.now().isoformat()

    def _get_index_path(self) -> str:
        return os.path.join(self.data_dir, "index.json")

    def _get_template_path(self, template_id: str) -> str:
        return os.path.join(self.data_dir, f"{template_id}.json")

    def _load_index(self) -> List[dict]:
        path = self._get_index_path()
        if not os.path.exists(path):
            return []
        try:
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
            if isinstance(data, list):
                return data
            return []
        except (json.JSONDecodeError, OSError) as exc:
            logger.warning("Failed to load index: %s", exc)
            return []

    def _save_index(self, index: List[dict]) -> None:
        path = self._get_index_path()
        tmp_path = path + ".tmp"
        with open(tmp_path, "w", encoding="utf-8") as f:
            json.dump(index, f, ensure_ascii=False, indent=2)
        os.replace(tmp_path, path)

    def _load_item_keywords(self) -> Dict[str, List[str]]:
        """从 rules JSON 文件加载条目标题，构建关键词→标题映射。"""
        if self._item_keywords is not None:
            return self._item_keywords

        mapping: Dict[str, List[str]] = {}
        if not os.path.isdir(self._rules_dir):
            self._item_keywords = mapping
            return mapping

        for fname in os.listdir(self._rules_dir):
            if not fname.endswith(".json"):
                continue
            fpath = os.path.join(self._rules_dir, fname)
            try:
                with open(fpath, "r", encoding="utf-8") as f:
                    data = json.load(f)
            except (json.JSONDecodeError, OSError):
                continue

            for stage_name, stage_data in data.get("design_stages", {}).items():
                if not isinstance(stage_data, dict):
                    continue
                for cat_name, cat_data in stage_data.items():
                    if not isinstance(cat_data, dict):
                        continue
                    for item in cat_data.get("items", []):
                        if not isinstance(item, dict):
                            continue
                        title = item.get("title", "")
                        if title:
                            # 用标题本身作为关键词
                            mapping.setdefault(title, [])
                            # 也提取标题中的核心词
                            for kw in re.split(r"[及与、和/]", title):
                                kw = kw.strip()
                                if len(kw) >= 2:
                                    mapping.setdefault(kw, [])
                                    if title not in mapping[kw]:
                                        mapping[kw].append(title)

        self._item_keywords = mapping
        return mapping

    # ── 文件读取 ──────────────────────────────────────────────────────

    def _read_txt(self, file_path: str) -> str:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()

    def _read_docx(self, file_path: str) -> str:
        try:
            from docx import Document
        except ImportError:
            logger.error("python-docx not installed")
            return ""
        try:
            doc = Document(file_path)
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            return "\n\n".join(paragraphs)
        except Exception as exc:
            logger.error("Failed to read docx %s: %s", file_path, exc)
            return ""

    # ── 文档拆分 ──────────────────────────────────────────────────────

    def split_document(self, content: str, content_type: str = "txt") -> List[dict]:
        """将文档内容按标题拆分为多个章节。

        Args:
            content: 文档文本内容
            content_type: "txt", "md", 或 "docx"

        Returns:
            [{title, content}, ...]
        """
        if not content or not content.strip():
            return []

        if content_type == "md":
            return self._split_by_md(content)
        else:
            return self._split_by_txt(content)

    def _split_by_txt(self, content: str) -> List[dict]:
        """按中文标题模式拆分纯文本。"""
        lines = content.split("\n")
        sections: List[dict] = []
        current_title = ""
        current_lines: List[str] = []

        for line in lines:
            stripped = line.strip()
            if stripped and _TXT_HEADING_RE.match(stripped):
                # 保存前一个章节
                if current_lines or current_title:
                    sections.append({
                        "title": current_title or "未命名章节",
                        "content": "\n".join(current_lines).strip(),
                    })
                current_title = stripped
                current_lines = []
            else:
                current_lines.append(line)

        # 最后一个章节
        if current_lines or current_title:
            sections.append({
                "title": current_title or "未命名章节",
                "content": "\n".join(current_lines).strip(),
            })

        # 如果没有拆分出任何章节，返回整体
        if not sections:
            sections.append({"title": "全文", "content": content.strip()})

        return sections

    def _split_by_md(self, content: str) -> List[dict]:
        """按 Markdown 标题拆分。"""
        lines = content.split("\n")
        sections: List[dict] = []
        current_title = ""
        current_lines: List[str] = []

        for line in lines:
            stripped = line.strip()
            if stripped and _MD_HEADING_RE.match(stripped):
                if current_lines or current_title:
                    sections.append({
                        "title": current_title or "未命名章节",
                        "content": "\n".join(current_lines).strip(),
                    })
                # 去掉 # 前缀
                current_title = re.sub(r"^#+\s*", "", stripped).strip()
                current_lines = []
            else:
                current_lines.append(line)

        if current_lines or current_title:
            sections.append({
                "title": current_title or "未命名章节",
                "content": "\n".join(current_lines).strip(),
            })

        if not sections:
            sections.append({"title": "全文", "content": content.strip()})

        return sections

    def _split_docx_by_styles(self, file_path: str) -> List[dict]:
        """按 docx 段落样式拆分文档。"""
        try:
            from docx import Document
        except ImportError:
            return []

        try:
            doc = Document(file_path)
        except Exception as exc:
            logger.error("Failed to open docx %s: %s", file_path, exc)
            return []

        sections: List[dict] = []
        current_title = ""
        current_lines: List[str] = []

        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ""
            text = para.text.strip()

            if style_name in _DOCX_HEADING_STYLES and text:
                if current_lines or current_title:
                    sections.append({
                        "title": current_title or "未命名章节",
                        "content": "\n".join(current_lines).strip(),
                    })
                current_title = text
                current_lines = []
            elif text:
                current_lines.append(text)

        if current_lines or current_title:
            sections.append({
                "title": current_title or "未命名章节",
                "content": "\n".join(current_lines).strip(),
            })

        if not sections:
            # 回退到纯文本拆分
            full_text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
            if full_text:
                sections.append({"title": "全文", "content": full_text})

        return sections

    # ── 分类逻辑 ──────────────────────────────────────────────────────

    def classify_content(self, title: str, content: str) -> dict:
        """基于关键词自动分类。

        Args:
            title: 章节标题
            content: 文本内容

        Returns:
            {project_type, design_stage, category, item_title, confidence}
        """
        text = f"{title} {content[:2000]}".lower()
        scores: Dict[str, float] = {}

        # 工程类型
        project_type = "unclassified"
        pt_scores: Dict[str, float] = {}
        for pt, keywords in _PROJECT_TYPE_KEYWORDS.items():
            count = sum(1 for kw in keywords if kw.lower() in text)
            if count > 0:
                pt_scores[pt] = count
        if pt_scores:
            project_type = max(pt_scores, key=pt_scores.get)  # type: ignore[arg-type]

        # 设计阶段
        design_stage = "unclassified"
        ds_scores: Dict[str, float] = {}
        for ds, keywords in _DESIGN_STAGE_KEYWORDS.items():
            count = sum(1 for kw in keywords if kw.lower() in text)
            if count > 0:
                ds_scores[ds] = count
        if ds_scores:
            design_stage = max(ds_scores, key=ds_scores.get)  # type: ignore[arg-type]

        # 专业分类
        category = "unclassified"
        cat_scores: Dict[str, float] = {}
        for cat, keywords in _CATEGORY_KEYWORDS.items():
            count = sum(1 for kw in keywords if kw.lower() in text)
            if count > 0:
                cat_scores[cat] = count
        if cat_scores:
            category = max(cat_scores, key=cat_scores.get)  # type: ignore[arg-type]

        # 条目匹配
        item_title = "unclassified"
        item_keywords = self._load_item_keywords()
        best_item_score = 0
        for kw, titles in item_keywords.items():
            if kw.lower() in text:
                for t in titles:
                    score = len(kw)
                    if score > best_item_score:
                        best_item_score = score
                        item_title = t

        # 置信度计算
        total_hits = sum(pt_scores.values()) + sum(ds_scores.values()) + sum(cat_scores.values())
        if best_item_score > 0:
            total_hits += 1
        confidence = min(1.0, total_hits / 10.0) if total_hits > 0 else 0.0

        # 低置信度标记为 unclassified
        if confidence < 0.3:
            project_type = "unclassified"
            design_stage = "unclassified"
            category = "unclassified"
            item_title = "unclassified"

        return {
            "project_type": project_type,
            "design_stage": design_stage,
            "category": category,
            "item_title": item_title,
            "confidence": round(confidence, 2),
        }

    # ── 导入 ──────────────────────────────────────────────────────────

    def import_file(self, file_path: str) -> dict:
        """导入单个文件，自动分类并存储。

        Args:
            file_path: 文件路径

        Returns:
            保存的模板记录
        """
        if not os.path.isfile(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = os.path.splitext(file_path)[1].lower()
        if ext not in (".docx", ".txt", ".md"):
            raise ValueError(f"Unsupported file type: {ext}")

        filename = os.path.basename(file_path)
        content_type = ext.lstrip(".")

        # 读取内容
        if ext == ".docx":
            sections = self._split_docx_by_styles(file_path)
        else:
            raw_content = self._read_txt(file_path)
            sections = self.split_document(raw_content, content_type)

        # 为每个章节创建模板记录
        records: List[dict] = []
        for section in sections:
            title = section["title"]
            content = section["content"]

            classification = self.classify_content(title, content)

            record = {
                "id": self._generate_id(),
                "source_file": filename,
                "source_path": os.path.abspath(file_path),
                "imported_at": self._now_iso(),
                "project_type": classification["project_type"],
                "design_stage": classification["design_stage"],
                "category": classification["category"],
                "item_title": classification["item_title"],
                "title": title,
                "content": content[:5000],  # 限制预览长度
                "content_type": content_type,
                "notes": "",
                "_confidence": classification["confidence"],
            }

            # 保存单个文件
            tpl_path = self._get_template_path(record["id"])
            with open(tpl_path, "w", encoding="utf-8") as f:
                json.dump(record, f, ensure_ascii=False, indent=2)

            records.append(record)

        # 更新索引
        index = self._load_index()
        index.extend(records)
        self._save_index(index)

        logger.info("Imported %d section(s) from %s", len(records), filename)
        return records[0] if len(records) == 1 else records  # type: ignore[return-value]

    def import_files(self, file_paths: List[str]) -> List[dict]:
        """批量导入多个文件。"""
        results: List[dict] = []
        for fp in file_paths:
            try:
                result = self.import_file(fp)
                if isinstance(result, list):
                    results.extend(result)
                else:
                    results.append(result)
            except Exception as exc:
                logger.error("Failed to import %s: %s", fp, exc)
        return results

    # ── 查询 ──────────────────────────────────────────────────────────

    def list_templates(
        self,
        project_type: Optional[str] = None,
        design_stage: Optional[str] = None,
        category: Optional[str] = None,
        search: Optional[str] = None,
    ) -> List[dict]:
        """列出模板，支持过滤和搜索。"""
        index = self._load_index()
        results: List[dict] = []

        for record in index:
            if project_type and record.get("project_type") != project_type:
                continue
            if design_stage and record.get("design_stage") != design_stage:
                continue
            if category and record.get("category") != category:
                continue
            if search:
                query = search.lower()
                title = record.get("title", "").lower()
                content = record.get("content", "").lower()
                notes = record.get("notes", "").lower()
                if query not in title and query not in content and query not in notes:
                    continue
            results.append(record)

        return results

    def get_template(self, template_id: str) -> Optional[dict]:
        """按 ID 获取单个模板。"""
        path = self._get_template_path(template_id)
        if not os.path.exists(path):
            return None
        try:
            with open(path, "r", encoding="utf-8") as f:
                return json.load(f)
        except (json.JSONDecodeError, OSError) as exc:
            logger.warning("Failed to load template %s: %s", template_id, exc)
            return None

    def delete_template(self, template_id: str) -> bool:
        """删除模板。"""
        path = self._get_template_path(template_id)
        if os.path.exists(path):
            try:
                os.remove(path)
            except OSError as exc:
                logger.error("Failed to delete file %s: %s", path, exc)
                return False

        index = self._load_index()
        new_index = [r for r in index if r.get("id") != template_id]
        if len(new_index) == len(index):
            return False
        self._save_index(new_index)
        return True

    def update_notes(self, template_id: str, notes: str) -> bool:
        """更新模板备注。"""
        path = self._get_template_path(template_id)
        if not os.path.exists(path):
            return False

        try:
            with open(path, "r", encoding="utf-8") as f:
                record = json.load(f)
            record["notes"] = notes
            with open(path, "w", encoding="utf-8") as f:
                json.dump(record, f, ensure_ascii=False, indent=2)
        except (json.JSONDecodeError, OSError) as exc:
            logger.error("Failed to update notes for %s: %s", template_id, exc)
            return False

        # 同步更新索引中的 notes
        index = self._load_index()
        for record in index:
            if record.get("id") == template_id:
                record["notes"] = notes
                break
        self._save_index(index)
        return True
