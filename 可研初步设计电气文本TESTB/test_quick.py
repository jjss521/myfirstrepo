# -*- coding: utf-8 -*-
import sys, os, json

os.environ['PYTHONIOENCODING'] = 'utf-8'

PROJ = r"D:\qoderwork\可研初步设计电气自控文本编写"
os.chdir(PROJ)

sys.path.insert(0, os.path.join(PROJ, 'backend'))

# 1
from app.services.excel_parser import ExcelLoadParser
print("[OK] ExcelLoadParser")

# 2
from app.services.docx_generator import DocxGenerator
print("[OK] DocxGenerator")

# 3
sys.path.insert(0, os.path.join(PROJ, 'src'))
from core.engine import GenerateEngine
eng = GenerateEngine(project_root=PROJ)
print("[OK] GenerateEngine")

# 4
status = eng.health_check()
for name, msg in status['checks']:
    m = msg.replace('[OK]', 'PASS').replace('OK', 'PASS')
    print(f"  {name}: {m}")

# 5
summ = eng.get_rule_summary('water_supply')
print(f"\nWater supply rules: {summ['project_type']}")
for cat in summ['categories']:
    print(f"  [{cat['section_id']}] {cat['title']}: {cat['item_count']} items")

# 6 - generate doc
gen = DocxGenerator(
    rules_dir=os.path.join(PROJ, 'backend', 'data', 'rules'),
    output_dir=os.path.join(PROJ, 'output'),
)

excel_data = {
    'summary': {
        'total_devices': 30, 'total_equip_power': 1850.5,
        'total_pc': 1200.3, 'total_qc': 680.2, 'total_sc': 1380.0,
        'total_pc_k': 1080.3, 'total_qc_k': 646.2, 'total_sc_k': 1258.5,
        'cos_before': 0.8581, 'cos_target': 0.95,
        'qc_compensation': 380.5,
        'total_qc_after': 265.7, 'total_sc_after': 1112.5,
        'recommended_transformer': '2x800kVA',
        'simultaneous_coeff': {'KP': 0.9, 'Kq': 0.95},
    },
    'area_summaries': {
        'Qushui Pump': {'device_count': 5, 'equip_power': 650.0, 'pc': 585.0, 'qc': 362.7, 'sc': 688.2},
        'Songshui Pump': {'device_count': 4, 'equip_power': 480.0, 'pc': 432.0, 'qc': 267.8, 'sc': 508.2},
        'Chemical': {'device_count': 6, 'equip_power': 180.5, 'pc': 126.4, 'qc': 75.8, 'sc': 147.4},
        'Filter': {'device_count': 8, 'equip_power': 320.0, 'pc': 256.0, 'qc': 153.6, 'sc': 298.6},
        'Auxiliary': {'device_count': 7, 'equip_power': 220.0, 'pc': 110.0, 'qc': 55.0, 'sc': 123.0},
    },
    'area_count': 5,
}

params = {
    'project_name': 'Chibi Water Plant',
    'voltage_level': '10kV',
    'load_level': 'Level 2',
    'project_type': 'water_supply',
    'power_source': 'Two circuits',
    'standby_desc': 'Standby desc',
    'tx_config': '2x800kVA',
    'tx_count': '2',
    'tx_location': 'Switchgear room',
}

out = gen.generate('water_supply', 'Preliminary Design', excel_data, params)
fsize = os.path.getsize(out)
print(f"\n[OK] DOCX generated: {out}")
print(f"    Size: {fsize} bytes ({fsize/1024:.1f} KB)")

# verify
from docx import Document
doc = Document(out)
para_count = len(doc.paragraphs)
table_count = len(doc.tables)
print(f"    Paragraphs: {para_count}")
print(f"    Tables: {table_count}")
headings = [p.text for p in doc.paragraphs if p.text.strip() and any(r.bold for r in p.runs)]
for h in headings[:15]:
    print(f"    - {h}")
