# -*- coding: utf-8 -*-
"""Verify generated DOCX quality"""
import os, sys
from docx import Document

os.chdir(r"D:\qoderwork\可研初步设计电气自控文本编写")

out_dir = r"output\generated"
files = [f for f in os.listdir(out_dir) if f.endswith('.docx')]
if not files:
    print("No DOCX files found")
    sys.exit(1)

latest = sorted(files)[-1]
path = os.path.join(out_dir, latest)
print(f"Verifying: {latest}")
print(f"Size: {os.path.getsize(path)} bytes")

doc = Document(path)

# Count structure
headings_bold = [(p.text, len([r for r in p.runs if r.bold]))
                 for p in doc.paragraphs if p.text.strip()]
print(f"\nDocument structure ({len(headings_bold)} non-empty paragraphs):")
for text, bcount in headings_bold:
    prefix = "## BOLD" if bcount > 0 else "   "
    text_trunc = text[:70] + ("..." if len(text) > 70 else "")
    print(f"  {prefix} {text_trunc}")

print(f"\nTables: {len(doc.tables)}")
for i, table in enumerate(doc.tables):
    rows = len(table.rows)
    cols = len(table.columns)
    header = [cell.text for cell in table.rows[0].cells]
    print(f"  Table {i+1}: {rows} rows x {cols} cols")
    print(f"    Headers: {header}")

print("\nSections:", len(doc.sections))
for i, section in enumerate(doc.sections):
    print(f"  Section {i+1}: {section.orientation}")

print("\n[OK] Verification complete")
